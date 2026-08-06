# -*- coding: utf-8 -*-
"""
POLTECH - Sistema de Control de Nomina
Version 1 (cimiento): Login + Roles, Obras, Catalogo de sueldos,
Personal (con validacion de NSS y CURP) y Cuentas bancarias (sin duplicados).

Hecho para correr en Flask + gunicorn (Render). Base de datos: SQLite.
NOTA: en el plan gratis de Render los datos de SQLite NO son permanentes.
Usa datos de prueba hasta conectar una base de datos permanente.
"""
import os
import re
import io
import json
import sqlite3
import smtplib
import urllib.request
import urllib.error
import urllib.parse
from datetime import date, datetime, timedelta
from functools import wraps
from email.message import EmailMessage

import openpyxl
from flask import (Flask, request, redirect, url_for, session,
                   flash, render_template, g, abort, send_file, jsonify)
from werkzeug.security import generate_password_hash, check_password_hash

APP_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.environ.get("DB_PATH", os.path.join(APP_DIR, "poltech.db"))

APP_VERSION = "1.22"   # version del sistema (visible en el menu)

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-cambia-esta-clave-en-render")
app.config["MAX_CONTENT_LENGTH"] = 6 * 1024 * 1024  # 6 MB por archivo subido

# ---------------------------------------------------------------------------
# Roles y jerarquia (a mayor numero, mas permisos)
# ---------------------------------------------------------------------------
ROLES = {
    "admin":           ("Administrador", 100),
    "superintendente": ("Superintendente de obra", 70),
    "residente":       ("Residente", 30),
}
GERENTE_RANK = 70  # de gerente de obra hacia arriba
ADMIN_RANK = 100   # solo administrador

def role_label(r): return ROLES.get(r, (r, 0))[0]
def role_rank(r):  return ROLES.get(r, (r, 0))[1]

def obras_del_usuario(db=None):
    """None = ve todas las obras (administrador).
    Lista de obra_ids = solo esas obras (superintendente/residente). [] = ninguna."""
    if role_rank(session.get("role", "")) >= ADMIN_RANK:
        return None
    db = db or get_db()
    rows = db.execute("SELECT obra_id FROM user_obras WHERE user_id=?",
                      (session.get("user_id"),)).fetchall()
    return [r["obra_id"] for r in rows]

def enviar_correo(asunto, cuerpo, destinatarios, adjuntos=None):
    """Envia un correo a una lista de destinatarios, con adjuntos opcionales.
    adjuntos: lista de (nombre_archivo, bytes, mimetype). Nunca detiene el sistema."""
    host = os.environ.get("SMTP_HOST")
    user = os.environ.get("SMTP_USER")
    pw   = os.environ.get("SMTP_PASS")
    dest = [d for d in (destinatarios or []) if d and "@" in d]
    if not (host and user and pw and dest):
        app.logger.warning("CORREO (no configurado o sin destinatarios): %s | para %s", asunto, dest)
        return False
    try:
        msg = EmailMessage()
        msg["Subject"] = asunto
        msg["From"] = user
        msg["To"] = ", ".join(dest)
        msg.set_content(cuerpo)
        for nombre, datos, mime in (adjuntos or []):
            maintype, _, subtype = (mime or "application/octet-stream").partition("/")
            msg.add_attachment(datos, maintype=maintype, subtype=subtype, filename=nombre)
        port = int(os.environ.get("SMTP_PORT", "587"))
        with smtplib.SMTP(host, port, timeout=15) as s:
            s.starttls()
            s.login(user, pw)
            s.send_message(msg)
        return True
    except Exception as e:
        app.logger.error("No se pudo enviar el correo: %s", e)
        return False

def send_admin_alert(asunto, cuerpo, extra=None):
    """Alerta al administrador (y opcionalmente a mas destinatarios)."""
    dest = [os.environ.get("ADMIN_EMAIL", "")]
    if extra:
        dest += extra
    return enviar_correo(asunto, cuerpo, dest)

def siguiente_cedula(db):
    """Cedula del checador: A001, A002 ... A999, B001 ...
    Unica por trabajador, independiente de la obra."""
    row = db.execute(
        "SELECT cedula FROM empleados WHERE cedula IS NOT NULL AND cedula<>'' "
        "ORDER BY cedula DESC LIMIT 1").fetchone()
    if not row or not row["cedula"]:
        return "A001"
    letra = row["cedula"][0]
    num = int(row["cedula"][1:])
    if num >= 999:
        letra = chr(ord(letra) + 1)
        num = 1
    else:
        num += 1
    return f"{letra}{num:03d}"

def siguiente_clave_otro_pago(db):
    """Clave B001, B002... para el roster de Otros pagos semanales.
    Es una serie propia, independiente de la cedula del checador (empleados.cedula):
    no se usa para asistencia ni calculo de nomina, solo para identificar el registro."""
    row = db.execute(
        "SELECT clave FROM otros_pagos WHERE clave IS NOT NULL AND clave<>'' "
        "ORDER BY clave DESC LIMIT 1").fetchone()
    if not row or not row["clave"]:
        return "B001"
    num = int(row["clave"][1:])
    return f"B{num + 1:03d}"


def resolver_puesto(db, obra_nom):
    """En la carga masiva, el texto de la columna Puesto ya no se usa para elegir
    el puesto (queda solo de referencia): siempre se asigna el puesto activo de
    menor sueldo semanal de la obra indicada, para revisar y ajustar el sueldo
    de cada quien despues. Devuelve el puesto (id, nombre, sueldo, estado de la
    obra) o None si la obra no existe o no tiene puestos activos."""
    return db.execute(
        "SELECT p.id, p.nombre AS puesto, p.sueldo_semanal, o.estado FROM puestos p "
        "JOIN obras o ON o.id=p.obra_id "
        "WHERE lower(o.nombre)=lower(?) AND p.activo=1 "
        "ORDER BY p.sueldo_semanal ASC LIMIT 1",
        ((obra_nom or "").strip(),)).fetchone()


ESTADOS = [
    "Aguascalientes", "Baja California", "Baja California Sur", "Campeche",
    "Chiapas", "Chihuahua", "Ciudad de Mexico", "Coahuila", "Colima",
    "Durango", "Estado de Mexico", "Guanajuato", "Guerrero", "Hidalgo",
    "Jalisco", "Michoacan", "Morelos", "Nayarit", "Nuevo Leon", "Oaxaca",
    "Puebla", "Queretaro", "Quintana Roo", "San Luis Potosi", "Sinaloa",
    "Sonora", "Tabasco", "Tamaulipas", "Tlaxcala", "Veracruz", "Yucatan",
    "Zacatecas",
]
TIPOS_CUENTA = ["Tarjeta", "CLABE interbancaria", "Numero de cuenta"]

# Clasificaciones del catalogo de sueldos (para totalizar la nomina por grupo)
CLASIFICACIONES = [
    "Gerencia de obra", "Residencia de obra", "Administracion", "Montadores",
    "Soldadores", "Laminadores", "Ayudantes generales", "Operadores",
    "Seguridad", "Almacen", "Topografia",
]

MOTIVOS_BAJA = [
    "Renuncia voluntaria", "Termino de la obra o contrato", "Rescision de contrato",
    "Abandono de empleo", "Ausentismo", "Defuncion", "Otro",
]

def titulo(s):
    """Primera letra de cada palabra en mayuscula y el resto en minuscula.
    Mantiene los acentos y no rompe apostrofes."""
    if not s:
        return s
    return " ".join(w[:1].upper() + w[1:].lower() for w in str(s).split())

SIGLAS_CONOCIDAS = {"ISSTE", "ISSSTE", "IMSS", "INFONAVIT", "SAT", "CFE", "PEMEX",
                    "CDMX", "SEP", "STPS", "UMF", "SA", "CV", "RFC", "CURP", "NSS",
                    "SAPI", "SC", "SRL"}

def titulo_obra(s):
    """Como titulo(), pero conserva las SIGLAS conocidas (ISSTE, IMSS, SAT, ...).
    Sirve para nombres de obra y puesto sin arruinar palabras normales como OBRA."""
    if not s:
        return s
    out = []
    for w in str(s).split():
        if w.upper().strip(".,") in SIGLAS_CONOCIDAS:
            out.append(w.upper())
        else:
            out.append(w[:1].upper() + w[1:].lower())
    return " ".join(out)
# NSS generico para personal pensionado que no requiere alta en el IMSS.
# Se exime de la validacion normal (digito verificador) y de la regla de
# NSS unico, para que varios pensionados puedan compartirlo.
NSS_GENERICO = "00000000000"

# Catalogo de bancos disponibles en Mexico (los mas usados para nomina).
BANCOS = [
    "BBVA Mexico", "Banorte", "Santander", "Citibanamex", "HSBC",
    "Scotiabank", "Inbursa", "Banco Azteca", "BanBajio", "Banregio",
    "Afirme", "Banca Mifel", "Multiva", "CIBanco", "Intercam Banco",
    "Banco Ve por Mas", "Compartamos Banco", "Banco del Bienestar",
    "BanCoppel", "STP", "Klar", "Nu Mexico", "Hey Banco", "Mercado Pago",
    "Otro",
]

# ---------------------------------------------------------------------------
# Base de datos
# ---------------------------------------------------------------------------
SCHEMA = """
CREATE TABLE IF NOT EXISTS users (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    username TEXT UNIQUE NOT NULL,
    password_hash TEXT NOT NULL,
    nombre TEXT NOT NULL,
    role TEXT NOT NULL DEFAULT 'capturista',
    obra_id INTEGER
);
CREATE TABLE IF NOT EXISTS obras (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    proyecto TEXT,
    contrato TEXT,
    nombre TEXT NOT NULL,
    estado TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS puestos (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    obra_id INTEGER NOT NULL,
    nombre TEXT NOT NULL,
    sueldo_semanal REAL NOT NULL DEFAULT 0,
    viaticos_semanales REAL NOT NULL DEFAULT 0,
    FOREIGN KEY (obra_id) REFERENCES obras(id)
);
CREATE TABLE IF NOT EXISTS empleados (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    ref TEXT,
    nombre TEXT NOT NULL,
    primer_apellido TEXT NOT NULL,
    segundo_apellido TEXT,
    curp TEXT,
    rfc TEXT,
    cp_fiscal TEXT,
    nss TEXT,
    sexo TEXT,
    fecha_nacimiento TEXT,
    puesto_id INTEGER,
    fecha_alta TEXT,
    importe_alta_imss REAL DEFAULT 0,
    infonavit_monto REAL DEFAULT 0,
    exime_docs INTEGER DEFAULT 0,
    autoriza_tercero TEXT,
    estatus TEXT DEFAULT 'activo',
    FOREIGN KEY (puesto_id) REFERENCES puestos(id)
);
CREATE TABLE IF NOT EXISTS cuentas_bancarias (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    empleado_id INTEGER NOT NULL,
    institucion TEXT NOT NULL,
    tipo_cuenta TEXT NOT NULL,
    numero TEXT UNIQUE NOT NULL,
    FOREIGN KEY (empleado_id) REFERENCES empleados(id)
);
CREATE TABLE IF NOT EXISTS documentos (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    empleado_id INTEGER NOT NULL,
    tipo TEXT NOT NULL,
    filename TEXT,
    contenido BLOB,
    subido_en TEXT,
    subido_por TEXT,
    UNIQUE(empleado_id, tipo),
    FOREIGN KEY (empleado_id) REFERENCES empleados(id)
);
CREATE TABLE IF NOT EXISTS user_obras (
    user_id INTEGER NOT NULL,
    obra_id INTEGER NOT NULL,
    UNIQUE(user_id, obra_id)
);
-- Fase 2: asistencia semanal (viernes a jueves, base de 6 dias)
CREATE TABLE IF NOT EXISTS asistencia_semanas (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    obra_id INTEGER NOT NULL,
    fecha_inicio TEXT NOT NULL,   -- viernes (ISO)
    fecha_fin TEXT NOT NULL,      -- jueves (ISO)
    semana_num INTEGER,
    anio INTEGER,
    estatus TEXT DEFAULT 'borrador',   -- borrador / cerrada
    creada_en TEXT,
    creada_por TEXT,
    UNIQUE(obra_id, fecha_inicio),
    FOREIGN KEY (obra_id) REFERENCES obras(id)
);
CREATE TABLE IF NOT EXISTS asistencia (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    semana_id INTEGER NOT NULL,
    empleado_id INTEGER NOT NULL,
    d1 TEXT, d2 TEXT, d3 TEXT, d4 TEXT, d5 TEXT, d6 TEXT, d7 TEXT,  -- VIE..JUE
    he_150 REAL DEFAULT 0,   -- horas extra al x1.5
    he_200 REAL DEFAULT 0,   -- horas extra al x2.0
    observaciones TEXT,
    UNIQUE(semana_id, empleado_id),
    FOREIGN KEY (semana_id) REFERENCES asistencia_semanas(id),
    FOREIGN KEY (empleado_id) REFERENCES empleados(id)
);
CREATE TABLE IF NOT EXISTS parametros (
    clave TEXT PRIMARY KEY,
    valor TEXT
);
CREATE TABLE IF NOT EXISTS nominas (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    obra_id INTEGER NOT NULL,
    fecha_inicio TEXT NOT NULL,
    fecha_fin TEXT NOT NULL,
    semana_num INTEGER,
    anio INTEGER,
    jornada REAL,
    aplico_retardos INTEGER DEFAULT 0,
    total_neto REAL DEFAULT 0,
    estatus TEXT DEFAULT 'pendiente',
    autorizada_por TEXT,
    autorizada_en TEXT,
    creada_en TEXT,
    creada_por TEXT,
    FOREIGN KEY (obra_id) REFERENCES obras(id)
);
CREATE TABLE IF NOT EXISTS nomina_detalle (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    nomina_id INTEGER NOT NULL,
    empleado_id INTEGER,
    cedula TEXT,
    nombre TEXT,
    dias REAL DEFAULT 0,
    faltas REAL DEFAULT 0,
    retardos REAL DEFAULT 0,
    vacaciones REAL DEFAULT 0,
    sueldo REAL DEFAULT 0,
    viaticos REAL DEFAULT 0,
    he_horas REAL DEFAULT 0,
    he_importe REAL DEFAULT 0,
    infonavit REAL DEFAULT 0,
    desc_nomina REAL DEFAULT 0,
    desc_otra REAL DEFAULT 0,
    desc_retardos REAL DEFAULT 0,
    neto REAL DEFAULT 0,
    baja_fecha TEXT,
    nota TEXT,
    FOREIGN KEY (nomina_id) REFERENCES nominas(id)
);
CREATE TABLE IF NOT EXISTS clasificaciones (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    nombre TEXT UNIQUE NOT NULL
);
CREATE TABLE IF NOT EXISTS otros_pagos (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    clave TEXT UNIQUE,
    nombre TEXT NOT NULL,
    banco TEXT,
    tipo_cuenta TEXT,
    numero_cuenta TEXT,
    monto_semanal REAL NOT NULL DEFAULT 0,
    activo INTEGER NOT NULL DEFAULT 1,
    creado_en TEXT
);
CREATE TABLE IF NOT EXISTS otros_pagos_lotes (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    anio INTEGER NOT NULL,
    semana_num INTEGER NOT NULL,
    nomina_id INTEGER,
    obra_id INTEGER,
    generado_por TEXT,
    generado_en TEXT,
    detalle_json TEXT,
    UNIQUE(anio, semana_num)
);
CREATE TABLE IF NOT EXISTS bitacora (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    fecha TEXT,
    usuario TEXT,
    rol TEXT,
    accion TEXT,
    detalle TEXT
);
CREATE TABLE IF NOT EXISTS contratos (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    empleado_id INTEGER NOT NULL,
    nombre_archivo TEXT,
    generado_en TEXT,
    generado_por TEXT,
    estatus TEXT DEFAULT 'generado',   -- generado / revisado
    revisado_por TEXT,
    revisado_en TEXT,
    contenido BLOB,
    FOREIGN KEY (empleado_id) REFERENCES empleados(id)
);
"""

def get_db():
    if "db" not in g:
        g.db = sqlite3.connect(DB_PATH)
        g.db.row_factory = sqlite3.Row
        g.db.execute("PRAGMA foreign_keys = ON")
    return g.db

@app.teardown_appcontext
def close_db(exc):
    db = g.pop("db", None)
    if db is not None:
        db.close()

def init_db():
    db = sqlite3.connect(DB_PATH)
    db.executescript(SCHEMA)
    # Migracion suave: agrega obra_id a users si la base venia de una version previa
    try:
        db.execute("ALTER TABLE users ADD COLUMN obra_id INTEGER")
    except sqlite3.OperationalError:
        pass
    # Migracion suave: cedula del checador y observaciones en empleados
    for col, tipo in (("cedula", "TEXT"), ("observaciones", "TEXT"),
                      ("estatus_docs", "TEXT DEFAULT 'Pendiente de carga'"),
                      ("fecha_registro", "TEXT"), ("fecha_solicitud", "TEXT"),
                      ("fecha_baja", "TEXT"), ("motivo_baja", "TEXT"),
                      ("viaticos_semanales", "REAL"), ("bono_semanal", "REAL DEFAULT 0"),
                      ("nss_generico", "INTEGER DEFAULT 0"), ("autoriza_nss_generico", "TEXT")):
        try:
            db.execute(f"ALTER TABLE empleados ADD COLUMN {col} {tipo}")
        except sqlite3.OperationalError:
            pass
    # Migracion suave: clasificacion en el catalogo de puestos y en el detalle de nomina
    try:
        db.execute("ALTER TABLE puestos ADD COLUMN clasificacion TEXT")
    except sqlite3.OperationalError:
        pass
    # Migracion suave: puesto activo/inactivo (se puede ocultar del catalogo sin borrarlo)
    try:
        db.execute("ALTER TABLE puestos ADD COLUMN activo INTEGER NOT NULL DEFAULT 1")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE nomina_detalle ADD COLUMN clasificacion TEXT")
    except sqlite3.OperationalError:
        pass
    # Migracion suave: autorizacion de nomina
    for col, tipo in (("estatus", "TEXT DEFAULT 'pendiente'"),
                      ("autorizada_por", "TEXT"), ("autorizada_en", "TEXT")):
        try:
            db.execute(f"ALTER TABLE nominas ADD COLUMN {col} {tipo}")
        except sqlite3.OperationalError:
            pass
    # Migracion suave: sueldo y viaticos contratados en el detalle de nomina
    for col in ("sueldo_contratado", "viaticos_contratado"):
        try:
            db.execute(f"ALTER TABLE nomina_detalle ADD COLUMN {col} REAL DEFAULT 0")
        except sqlite3.OperationalError:
            pass
    try:
        db.execute("ALTER TABLE nomina_detalle ADD COLUMN bono REAL DEFAULT 0")
    except sqlite3.OperationalError:
        pass
    # Migracion suave: sembrar la tabla de clasificaciones con la lista fija anterior
    if db.execute("SELECT COUNT(*) FROM clasificaciones").fetchone()[0] == 0:
        for nombre in CLASIFICACIONES:
            try:
                db.execute("INSERT INTO clasificaciones(nombre) VALUES(?)", (nombre,))
            except sqlite3.IntegrityError:
                pass
    # Crear usuario administrador la primera vez
    row = db.execute("SELECT COUNT(*) FROM users").fetchone()
    if row[0] == 0:
        pw = os.environ.get("ADMIN_PASSWORD", "cambiar123")
        db.execute(
            "INSERT INTO users(username, password_hash, nombre, role) VALUES(?,?,?,?)",
            ("admin", generate_password_hash(pw), "Administrador", "admin"),
        )
    # Migracion: mapear roles antiguos a los 3 nuevos
    db.execute("UPDATE users SET role='superintendente' "
               "WHERE role IN ('direccion','gerencia','gerente_obra','autorizador')")
    db.execute("UPDATE users SET role='residente' WHERE role='capturista'")
    # Migracion: pasar la obra unica (obra_id) a la tabla user_obras
    try:
        for r in db.execute("SELECT id, obra_id FROM users WHERE obra_id IS NOT NULL").fetchall():
            try:
                db.execute("INSERT INTO user_obras(user_id, obra_id) VALUES(?,?)", (r[0], r[1]))
            except sqlite3.IntegrityError:
                pass
    except sqlite3.OperationalError:
        pass
    # Parametros por defecto (salario minimo 2026 CONASAMI, editable por el admin)
    for clave, valor in (("sm_general", "315.04"),
                         ("sm_zlfn", "440.87"),
                         ("sm_vigencia", "2026-01-01"),
                         ("propuesta_extra", "15"),
                         ("jornada_horas", "8"),
                         ("porcentaje_despacho", "4"),
                         ("email_contador", "acerospoltech@gmail.com"),
                         ("empresa_razon_social", "POLTECH ACERO Y CONSTRUCCION, S.A. DE C.V."),
                         ("empresa_rfc", "PAC2408281N9"),
                         ("empresa_representante", "HUMBERTO FLORES PRADO"),
                         ("empresa_domicilio", "Calle Sur 27, Manzana 27 Lote 254, Col. Leyes de Reforma 1ra Seccion, Iztapalapa, Ciudad de Mexico, C.P. 09310"),
                         ("empresa_instrumento", "52,465"),
                         ("empresa_volumen", "1,086"),
                         ("empresa_fecha_escritura", "28/08/2024"),
                         ("empresa_notario_num", "10"),
                         ("empresa_notario_nombre", "ROBERTO MENDOZA NAVA"),
                         ("empresa_notario_ciudad", "Chalco, Estado de Mexico"),
                         ("apimarket_token", ""),
                         ("apimarket_url_nss", "https://apimarket.mx/api/imss/grupo/localizar-nss"),
                         ("apimarket_url_vigencia", "https://apimarket.mx/api/imss/grupo/consultar-vigencia")):
        db.execute("INSERT OR IGNORE INTO parametros(clave, valor) VALUES(?,?)", (clave, valor))
    # Rellenar datos de empresa en bases ya existentes (sin pisar lo que ya se edito)
    def _param_actual(clave):
        r = db.execute("SELECT valor FROM parametros WHERE clave=?", (clave,)).fetchone()
        return (r[0] if r else "") or ""
    if not _param_actual("empresa_rfc").strip():
        db.execute("INSERT INTO parametros(clave, valor) VALUES('empresa_rfc', ?) "
                   "ON CONFLICT(clave) DO UPDATE SET valor=excluded.valor", ("PAC2408281N9",))
    _dom_old = "Colonia Leyes de Reforma, Alcaldia de Iztapalapa, Ciudad de Mexico"
    _dom_new = ("Calle Sur 27, Manzana 27 Lote 254, Col. Leyes de Reforma 1ra Seccion, "
                "Iztapalapa, Ciudad de Mexico, C.P. 09310")
    if _param_actual("empresa_domicilio").strip() in ("", _dom_old):
        db.execute("INSERT INTO parametros(clave, valor) VALUES('empresa_domicilio', ?) "
                   "ON CONFLICT(clave) DO UPDATE SET valor=excluded.valor", (_dom_new,))
    if not _param_actual("apimarket_url_nss").strip():
        db.execute("INSERT INTO parametros(clave, valor) VALUES('apimarket_url_nss', ?) "
                   "ON CONFLICT(clave) DO UPDATE SET valor=excluded.valor",
                   ("https://apimarket.mx/api/imss/grupo/localizar-nss",))
    if not _param_actual("apimarket_url_vigencia").strip():
        db.execute("INSERT INTO parametros(clave, valor) VALUES('apimarket_url_vigencia', ?) "
                   "ON CONFLICT(clave) DO UPDATE SET valor=excluded.valor",
                   ("https://apimarket.mx/api/imss/grupo/consultar-vigencia",))
    # Normalizacion unica: nombre/apellidos de empleados, y nombres de puesto y obra
    if _param_actual("norm_nombres_hecho") != "1":
        for eid, nom, a1, a2 in db.execute(
                "SELECT id, nombre, primer_apellido, segundo_apellido FROM empleados").fetchall():
            db.execute("UPDATE empleados SET nombre=?, primer_apellido=?, segundo_apellido=? WHERE id=?",
                       (titulo(nom), titulo(a1), titulo(a2), eid))
        for pid, pnom in db.execute("SELECT id, nombre FROM puestos").fetchall():
            db.execute("UPDATE puestos SET nombre=? WHERE id=?", (titulo_obra(pnom), pid))
        for oid, onom in db.execute("SELECT id, nombre FROM obras").fetchall():
            db.execute("UPDATE obras SET nombre=? WHERE id=?", (titulo_obra(onom), oid))
        db.execute("INSERT INTO parametros(clave, valor) VALUES('norm_nombres_hecho','1') "
                   "ON CONFLICT(clave) DO UPDATE SET valor='1'")
    db.commit()
    db.close()

init_db()

# ---------------------------------------------------------------------------
# Autenticacion y permisos
# ---------------------------------------------------------------------------
def login_required(f):
    @wraps(f)
    def wrapper(*a, **k):
        if not session.get("user_id"):
            return redirect(url_for("login"))
        return f(*a, **k)
    return wrapper

def min_rank(rank):
    """Decorador: exige un rango minimo de rol."""
    def deco(f):
        @wraps(f)
        def wrapper(*a, **k):
            if not session.get("user_id"):
                return redirect(url_for("login"))
            if role_rank(session.get("role", "")) < rank:
                abort(403)
            return f(*a, **k)
        return wrapper
    return deco

@app.context_processor
def inject_user():
    return {
        "current_user": {
            "id": session.get("user_id"),
            "nombre": session.get("nombre"),
            "role": session.get("role"),
            "role_label": role_label(session.get("role", "")),
            "rank": role_rank(session.get("role", "")),
        },
        "GERENTE_RANK": GERENTE_RANK,
        "ADMIN_RANK": ADMIN_RANK,
        "APP_VERSION": APP_VERSION,
    }


@app.template_filter("money")
def money(v):
    """Formatea numeros como $1,234.56 (coma en miles, punto en centavos)."""
    try:
        return "${:,.2f}".format(float(v or 0))
    except (ValueError, TypeError):
        return v


# Estados del proceso de documentos de cada trabajador
ESTATUS_DOCS = ["Pendiente de carga", "Informacion incompleta",
                "Carga completa", "Validado"]
DIAS_RETENCION = 15  # dias para validar la documentacion antes de retener pago

# Documentos que se suben escaneados por trabajador
TIPOS_DOC = [
    ("INE", "INE / Identificacion oficial (PDF)"),
    ("Comprobante", "Comprobante de domicilio (PDF)"),
    ("Contrato", "Contrato firmado (PDF)"),
    ("Fiscal", "Constancia de Situacion Fiscal / RFC (PDF)"),
]
DOCS_REQUERIDOS = ["INE", "Comprobante", "Contrato"]  # necesarios para "Carga completa"


def recomputar_estatus_docs(db, emp_id):
    """Actualiza solo el estatus de documentos segun lo que ya se subio."""
    tipos = [r["tipo"] for r in db.execute(
        "SELECT tipo FROM documentos WHERE empleado_id=?", (emp_id,)).fetchall()]
    presentes = [t for t in DOCS_REQUERIDOS if t in tipos]
    fila = db.execute("SELECT estatus_docs FROM empleados WHERE id=?", (emp_id,)).fetchone()
    actual = fila["estatus_docs"] if fila else None
    if len(presentes) == 0:
        nuevo = "Pendiente de carga"
    elif len(presentes) < len(DOCS_REQUERIDOS):
        nuevo = "Informacion incompleta"
    else:
        nuevo = "Validado" if actual == "Validado" else "Carga completa"
    db.execute("UPDATE empleados SET estatus_docs=? WHERE id=?", (nuevo, emp_id))


def validar_solicitud_imss(fecha_solicitud, fecha_alta):
    """La solicitud de alta IMSS debe ser <= fecha de alta y como maximo 5 dias antes."""
    if not fecha_solicitud or not fecha_alta:
        return None
    try:
        fs = date.fromisoformat(fecha_solicitud[:10])
        fa = date.fromisoformat(fecha_alta[:10])
    except ValueError:
        return None
    if fs > fa:
        return "La fecha de solicitud de alta IMSS no puede ser posterior a la fecha de alta."
    if (fa - fs).days > 5:
        return "La fecha de solicitud de alta IMSS debe ser como maximo 5 dias antes de la fecha de alta."
    return None


def dias_para_retencion(emp):
    """Dias que faltan para completar la documentacion. Se ancla a la fecha de
    REGISTRO en el sistema (no a la fecha de alta real en el IMSS, que puede ser
    anterior si el trabajador ya llevaba tiempo laborando antes de capturarlo aqui).
    Devuelve (texto, clase_color)."""
    if (emp["estatus_docs"] or "") == "Validado":
        return ("Validado", "success")
    base = emp["fecha_registro"] or emp["fecha_alta"]
    if not base:
        return ("-", "secondary")
    try:
        d0 = date.fromisoformat(str(base)[:10])
    except ValueError:
        return ("-", "secondary")
    restan = DIAS_RETENCION - (date.today() - d0).days
    if restan <= 0:
        return ("Retenido", "danger")
    clase = "warning" if restan <= 5 else "secondary"
    return (f"{restan} dias", clase)

# ---------------------------------------------------------------------------
# Validaciones (NSS y CURP)  -- sin librerias externas
# ---------------------------------------------------------------------------
def solo_digitos(s):
    return re.sub(r"\D", "", s or "")

def luhn_ok(num):
    total, alt = 0, False
    for ch in reversed(num):
        d = int(ch)
        if alt:
            d *= 2
            if d > 9:
                d -= 9
        total += d
        alt = not alt
    return total % 10 == 0

def nss_valido(nss):
    n = solo_digitos(nss)
    return len(n) == 11 and luhn_ok(n)

def clabe_valida(num):
    """CLABE interbancaria: 18 digitos con digito verificador (pesos 3,7,1)."""
    n = solo_digitos(num)
    if len(n) != 18:
        return False
    pesos = [3, 7, 1]
    suma = sum((int(n[i]) * pesos[i % 3]) % 10 for i in range(17))
    control = (10 - (suma % 10)) % 10
    return control == int(n[17])

def validar_cuenta(tipo, numero):
    """Candado de formato para la cuenta bancaria. Devuelve un mensaje de error o None."""
    n = solo_digitos(numero)
    if not n:
        return None  # cuenta opcional; si no hay numero, no se valida aqui
    t = (tipo or "").lower()
    if "clabe" in t:
        if len(n) != 18:
            return f"La CLABE debe tener 18 digitos (tiene {len(n)}). Revisa si falta o sobra algun numero."
        if not clabe_valida(n):
            return "La CLABE no pasa la validacion del digito verificador. Revisa que no haya un numero mal."
    elif "tarjeta" in t:
        if len(n) not in (15, 16):
            return f"El numero de tarjeta debe tener 16 digitos (tiene {len(n)})."
        if not luhn_ok(n):
            return "El numero de tarjeta no pasa la validacion (digito verificador). Revisa los numeros."
    else:  # numero de cuenta
        if not (8 <= len(n) <= 20):
            return f"El numero de cuenta debe tener entre 8 y 20 digitos (tiene {len(n)})."
    return None

CURP_RE = re.compile(
    r"^[A-Z][AEIOUX][A-Z]{2}\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])"
    r"[HM][A-Z]{2}[B-DF-HJ-NP-TV-Z]{3}[A-Z0-9]\d$"
)

def curp_valida(curp):
    return bool(CURP_RE.match((curp or "").strip().upper()))

def curp_sexo(curp):
    return "H" if curp[10].upper() == "H" else "M"

def curp_fecha(curp):
    try:
        yy = int(curp[4:6]); mm = int(curp[6:8]); dd = int(curp[8:10])
        siglo = 2000 if not curp[16].isdigit() else 1900
        return date(siglo + yy, mm, dd).isoformat()
    except Exception:
        return None

def revisar_curp(curp, nombre, ap1, ap2, sexo, fecha_nac):
    """Devuelve una lista de advertencias (no bloquea)."""
    avisos = []
    curp = (curp or "").strip().upper()
    if not curp_valida(curp):
        return avisos  # el formato ya se valida aparte
    def inicial(x):
        x = (x or "").strip().upper()
        return x[0] if x else ""
    if ap1 and curp[0] != inicial(ap1):
        avisos.append("La CURP no coincide con la inicial del primer apellido.")
    if ap2 and curp[2] not in (inicial(ap2), "X"):
        avisos.append("La CURP no coincide con la inicial del segundo apellido.")
    if nombre and curp[3] not in (inicial(nombre), "X"):
        avisos.append("La CURP no coincide con la inicial del nombre.")
    if sexo and curp_sexo(curp) != sexo:
        avisos.append("El sexo no coincide con el que indica la CURP.")
    if fecha_nac and curp_fecha(curp) and curp_fecha(curp) != fecha_nac:
        avisos.append("La fecha de nacimiento no coincide con la que indica la CURP.")
    return avisos

# ---------------------------------------------------------------------------
# Fase 2: asistencia semanal (viernes a jueves, base de 6 dias)
# ---------------------------------------------------------------------------
# d1..d7 = de viernes a jueves. El domingo (d3) es visible pero NO cuenta
# para la base de 6 dias.
DIAS_SEMANA = [("d1", "VIE"), ("d2", "SAB"), ("d3", "DOM"), ("d4", "LUN"),
               ("d5", "MAR"), ("d6", "MIE"), ("d7", "JUE")]
DIAS_BASE = ["d1", "d2", "d4", "d5", "d6", "d7"]   # 6 dias, sin domingo
CODIGOS_ASIS = [
    ("A", "Asistio"),
    ("F", "Falta"),
    ("R", "Retardo"),
    ("D", "Descanso"),
]

def viernes_de(fecha):
    """Viernes que inicia la semana de nomina (viernes-jueves) que contiene 'fecha'."""
    offset = (fecha.weekday() - 4) % 7   # lunes=0 ... viernes=4
    return fecha - timedelta(days=offset)

def fechas_de_dias(viernes):
    """Etiqueta corta dd/mm para cada dia de la semana."""
    return {f"d{i+1}": (viernes + timedelta(days=i)).strftime("%d/%m") for i in range(7)}

def contar_asistencia(cods):
    """cods: dict d1..d7 con codigos. Cuenta solo los 6 dias base (domingo no).
    Devuelve (dias_trabajados, faltas, retardos). Un retardo cuenta como dia trabajado."""
    dias = faltas = retardos = 0
    for k in DIAS_BASE:
        c = (cods.get(k) or "").upper()
        if c in ("A", "R"):
            dias += 1
        if c == "F":
            faltas += 1
        if c == "R":
            retardos += 1
    return dias, faltas, retardos

def obras_visibles(db):
    """Lista de obras que el usuario puede ver (respeta rol/asignacion)."""
    vis = obras_del_usuario(db)
    if vis is None:
        return db.execute("SELECT * FROM obras ORDER BY nombre").fetchall()
    if vis:
        ph = ",".join("?" * len(vis))
        return db.execute(f"SELECT * FROM obras WHERE id IN ({ph}) ORDER BY nombre", vis).fetchall()
    return []

def puede_ver_obra(db, obra_id):
    vis = obras_del_usuario(db)
    return vis is None or obra_id in (vis or [])


# --- Parametros del sistema (salario minimo, etc.) ---
def get_param(db, clave, default=None):
    row = db.execute("SELECT valor FROM parametros WHERE clave=?", (clave,)).fetchone()
    return row["valor"] if row else default

def set_param(db, clave, valor):
    db.execute("INSERT INTO parametros(clave, valor) VALUES(?,?) "
               "ON CONFLICT(clave) DO UPDATE SET valor=excluded.valor", (clave, str(valor)))

def registrar_bitacora(db, accion, detalle=""):
    """Registra un movimiento en la bitacora (quien, cuando, que)."""
    db.execute(
        "INSERT INTO bitacora(fecha, usuario, rol, accion, detalle) VALUES(?,?,?,?,?)",
        (datetime.now().isoformat(timespec="seconds"),
         session.get("nombre", ""), session.get("role", ""), accion, detalle))


# ---------------------------------------------------------------------------
# Contratos: numero a letra y generacion del docx
# ---------------------------------------------------------------------------
MESES_ES = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio",
            "agosto", "septiembre", "octubre", "noviembre", "diciembre"]

def fecha_larga(fecha_iso):
    """Convierte 'AAAA-MM-DD' a '13 de mayo de 1995'. Si no se puede leer, regresa None."""
    try:
        d = date.fromisoformat(str(fecha_iso)[:10])
    except ValueError:
        return None
    return f"{d.day} de {MESES_ES[d.month-1]} de {d.year}"

def normalizar_fecha(valor):
    """Convierte una fecha a AAAA-MM-DD (formato interno) sin importar como haya
    llegado desde un Excel: fecha nativa de Excel, DD-MM-AAAA o DD/MM/AAAA (formato
    mexicano) o ya en AAAA-MM-DD. Si no se reconoce el formato, se regresa el texto
    tal cual para que quien valida la fila la rechace en vez de guardar una fecha
    incorrecta."""
    if isinstance(valor, (datetime, date)):
        return valor.strftime("%Y-%m-%d")
    s = "" if valor is None else str(valor).strip()
    if not s:
        return ""
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", s):
        return s
    m = re.fullmatch(r"(\d{1,2})[/-](\d{1,2})[/-](\d{4})", s)
    if m:
        d, mo, y = (int(g) for g in m.groups())
        try:
            return date(y, mo, d).isoformat()
        except ValueError:
            return s
    return s

def fecha_valida(s):
    """True si s es una fecha en formato interno AAAA-MM-DD reconocible."""
    return bool(re.fullmatch(r"\d{4}-\d{2}-\d{2}", s or ""))

def _centenas_letra(n):
    UNIDADES = ["", "UNO", "DOS", "TRES", "CUATRO", "CINCO", "SEIS", "SIETE", "OCHO", "NUEVE",
                "DIEZ", "ONCE", "DOCE", "TRECE", "CATORCE", "QUINCE", "DIECISEIS",
                "DIECISIETE", "DIECIOCHO", "DIECINUEVE", "VEINTE"]
    DECENAS = ["", "", "VEINTI", "TREINTA", "CUARENTA", "CINCUENTA", "SESENTA",
               "SETENTA", "OCHENTA", "NOVENTA"]
    CENTENAS = ["", "CIENTO", "DOSCIENTOS", "TRESCIENTOS", "CUATROCIENTOS", "QUINIENTOS",
                "SEISCIENTOS", "SETECIENTOS", "OCHOCIENTOS", "NOVECIENTOS"]
    if n == 0:
        return ""
    if n == 100:
        return "CIEN"
    c, resto = divmod(n, 100)
    out = CENTENAS[c]
    if resto:
        if resto <= 20:
            out = (out + " " + UNIDADES[resto]).strip()
        else:
            d, u = divmod(resto, 10)
            if d == 2:
                out = (out + " VEINTI" + UNIDADES[u].lower()).strip() if u else (out + " VEINTE").strip()
            else:
                seg = DECENAS[d] + (" Y " + UNIDADES[u] if u else "")
                out = (out + " " + seg).strip()
    return out.strip()

def numero_letra(n):
    """Entero a letras en mayusculas (hasta millones)."""
    n = int(n)
    if n == 0:
        return "CERO"
    partes = []
    millones, resto = divmod(n, 1_000_000)
    miles, cientos = divmod(resto, 1000)
    if millones:
        partes.append("UN MILLON" if millones == 1 else _centenas_letra(millones) + " MILLONES")
    if miles:
        partes.append("MIL" if miles == 1 else _centenas_letra(miles) + " MIL")
    if cientos:
        partes.append(_centenas_letra(cientos))
    return " ".join(p for p in partes if p).strip()

def pesos_letra(monto):
    """Ej. 330.04 -> 'TRESCIENTOS TREINTA PESOS 04/100 M.N.'"""
    try:
        monto = float(monto)
    except (TypeError, ValueError):
        monto = 0.0
    entero = int(monto)
    centavos = int(round((monto - entero) * 100))
    return f"{numero_letra(entero)} PESOS {centavos:02d}/100 M.N."

def edad_de(fecha_nac):
    try:
        f = date.fromisoformat(str(fecha_nac)[:10])
    except (ValueError, TypeError):
        return ""
    hoy = date.today()
    return hoy.year - f.year - ((hoy.month, hoy.day) < (f.month, f.day))

CLASIFICACION_FUNCIONES = {
    "Gerencia de obra": {
        "objetivo": "Planear, coordinar y supervisar la ejecuci\u00f3n de la obra, asegurando el "
                    "cumplimiento del programa, la calidad, la seguridad y el presupuesto autorizado.",
        "funciones": [
            "Coordinar a las cuadrillas y al personal de supervisi\u00f3n a su cargo.",
            "Supervisar avances de obra, calidad de los trabajos y cumplimiento del programa.",
            "Gestionar recursos, materiales y proveedores necesarios para la obra.",
            "Verificar el cumplimiento de las normas de seguridad aplicables.",
            "Reportar a la direcci\u00f3n de la empresa el estado y los resultados de la obra.",
        ],
    },
    "Residencia de obra": {
        "objetivo": "Coordinar la operaci\u00f3n diaria de la obra, dando seguimiento al personal, "
                    "los avances y las condiciones de seguridad en el sitio de trabajo.",
        "funciones": [
            "Supervisar diariamente el avance y la calidad de los trabajos en obra.",
            "Coordinar al personal, cuadrillas y proveedores presentes en el sitio.",
            "Verificar el cumplimiento de las medidas de seguridad y uso de EPP.",
            "Reportar incidencias, avances y necesidades a la gerencia de obra.",
        ],
    },
    "Administracion": {
        "objetivo": "Apoyar en las funciones administrativas, documentales y de control de la "
                    "obra o de las oficinas de la empresa.",
        "funciones": [
            "Capturar, organizar y resguardar documentaci\u00f3n administrativa.",
            "Apoyar en tr\u00e1mites, pagos y control de proveedores.",
            "Atender al personal y proveedores en asuntos administrativos.",
            "Elaborar reportes y mantener actualizados los archivos a su cargo.",
        ],
    },
    "Montadores": {
        "objetivo": "Ensamblar, alinear y montar estructuras met\u00e1licas, equipos o componentes "
                    "conforme a planos, especificaciones y procedimientos de seguridad.",
        "funciones": [
            "Interpretar planos, croquis y \u00f3rdenes de trabajo de montaje.",
            "Preparar, presentar y alinear piezas y estructuras conforme a tolerancias.",
            "Apoyar y ejecutar maniobras de montaje o izaje solo con capacitaci\u00f3n, permiso y equipo "
            "requeridos.",
            "Verificar dimensiones, plomos y niveles del trabajo realizado.",
            "Cuidar herramienta, equipo y materiales a su cargo.",
        ],
    },
    "Soldadores": {
        "objetivo": "Ejecutar procesos de soldadura en estructuras y componentes met\u00e1licos "
                    "conforme a especificaciones, procedimientos de calidad y normas de seguridad.",
        "funciones": [
            "Interpretar planos, especificaciones y procedimientos de soldadura aplicables.",
            "Preparar juntas y soldar \u00fanicamente conforme a procedimiento y calificaci\u00f3n vigente.",
            "Inspeccionar el trabajo propio e identificar defectos o desviaciones.",
            "Cuidar el equipo de soldadura, consumibles y materiales a su cargo.",
            "Reportar de inmediato riesgos, incidentes o condiciones inseguras.",
        ],
    },
    "Laminadores": {
        "objetivo": "Preparar, cortar, conformar y procesar l\u00e1mina y perfiles met\u00e1licos conforme "
                    "a planos, especificaciones y tolerancias.",
        "funciones": [
            "Trazar, medir y cortar l\u00e1mina y perfiles conforme a orden de trabajo.",
            "Operar la maquinaria y herramienta autorizada para su puesto.",
            "Verificar medidas, tolerancias y calidad del material procesado.",
            "Reportar desviaciones, fallas de equipo o necesidades de mantenimiento.",
        ],
    },
    "Ayudantes generales": {
        "objetivo": "Apoyar las actividades operativas de la obra o el taller bajo la instrucci\u00f3n "
                    "directa del personal calificado y del supervisor a cargo.",
        "funciones": [
            "Apoyar en la carga, traslado y acomodo de materiales y herramienta.",
            "Mantener el orden y la limpieza del \u00e1rea de trabajo asignada.",
            "Apoyar en maniobras y actividades bajo supervisi\u00f3n directa.",
            "Cumplir las instrucciones de seguridad y usar el equipo de protecci\u00f3n proporcionado.",
        ],
    },
    "Operadores": {
        "objetivo": "Operar maquinaria y equipo autorizado conforme a su capacitaci\u00f3n y a los "
                    "procedimientos de seguridad aplicables.",
        "funciones": [
            "Realizar la inspecci\u00f3n previa del equipo antes de operarlo.",
            "Operar la maquinaria conforme al manual, la capacitaci\u00f3n y las instrucciones recibidas.",
            "Reportar de inmediato fallas, necesidades de mantenimiento o condiciones inseguras.",
            "Respetar zonas de exclusi\u00f3n, se\u00f1alizaci\u00f3n y permisos de trabajo aplicables.",
        ],
    },
    "Seguridad": {
        "objetivo": "Vigilar el cumplimiento de las medidas de seguridad, orden y control de "
                    "acceso en la obra o instalaci\u00f3n asignada.",
        "funciones": [
            "Controlar el acceso de personal, visitantes y veh\u00edculos.",
            "Verificar el uso correcto del equipo de protecci\u00f3n personal.",
            "Reportar actos y condiciones inseguras al supervisor correspondiente.",
            "Apoyar en la respuesta ante incidentes o emergencias conforme a los protocolos vigentes.",
        ],
    },
    "Almacen": {
        "objetivo": "Recibir, resguardar, controlar y entregar materiales, herramienta y equipo "
                    "de la obra o el taller.",
        "funciones": [
            "Registrar las entradas y salidas de materiales, herramienta y equipo.",
            "Mantener actualizado el inventario y los resguardos correspondientes.",
            "Entregar y recibir bienes mediante vale o resguardo firmado.",
            "Reportar de inmediato faltantes, da\u00f1os o irregularidades detectadas.",
        ],
    },
    "Topografia": {
        "objetivo": "Realizar los levantamientos, trazos y verificaciones topogr\u00e1ficas "
                    "necesarias para la ejecuci\u00f3n de la obra.",
        "funciones": [
            "Operar el equipo topogr\u00e1fico conforme a su capacitaci\u00f3n.",
            "Realizar trazos, niveles y verificaciones conforme al proyecto.",
            "Registrar y reportar las mediciones realizadas.",
            "Verificar el cumplimiento de las referencias y ejes del proyecto.",
        ],
    },
}
CLASIFICACION_FUNCIONES_DEFAULT = {
    "objetivo": "Ejecutar las actividades propias de su puesto conforme a las instrucciones del "
                "supervisor, las especificaciones aplicables y las normas de seguridad vigentes.",
    "funciones": [
        "Realizar las tareas asignadas por su supervisor relacionadas con su categor\u00eda.",
        "Cuidar la herramienta, el equipo y los materiales a su cargo.",
        "Cumplir las normas de seguridad, calidad y conducta aplicables a su puesto.",
        "Reportar de inmediato cualquier incidente, riesgo o desviaci\u00f3n detectada.",
    ],
}


def generar_contrato_docx(db, emp):
    """Arma el contrato individual de trabajo (docx) con los datos del empleado
    y de la empresa, siguiendo el formato aprobado (28 cl\u00e1usulas + 3 anexos).
    Devuelve (bytes, nombre_archivo). Los datos que el sistema no captura quedan
    resaltados en amarillo para llenarse a mano antes de firmar."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

    P = lambda k, d="": (get_param(db, k, d) or d)

    razon = P("empresa_razon_social")
    rep = P("empresa_representante")
    dom = P("empresa_domicilio")
    rfc_emp = P("empresa_rfc")
    instr = P("empresa_instrumento")
    fesc = P("empresa_fecha_escritura")
    not_num = P("empresa_notario_num")
    not_nom = P("empresa_notario_nombre")
    not_ciu = P("empresa_notario_ciudad")

    nombre = titulo(" ".join(x for x in [emp["nombre"], emp["primer_apellido"], emp["segundo_apellido"]] if x))
    puesto = titulo_obra(emp["puesto"]) if emp["puesto"] else ""
    obra_nom = emp["obra"] or ""
    estado_obra = emp["estado"] or ""
    sueldo_semanal = _num(emp["sueldo_semanal"])
    salario_diario = round(sueldo_semanal / 7.0, 2) if sueldo_semanal else 0
    salario_cotizacion = _num(emp["importe_alta_imss"])
    sexo_txt = "MASCULINO" if (emp["sexo"] or "").upper().startswith("H") else ("FEMENINO" if (emp["sexo"] or "").upper().startswith("M") else None)
    edad = edad_de(emp["fecha_nacimiento"])
    fecha_nac_txt = fecha_larga(emp["fecha_nacimiento"]) if emp["fecha_nacimiento"] else None
    fecha_alta_txt = fecha_larga(emp["fecha_alta"]) if emp["fecha_alta"] else None
    hoy = date.today()
    fecha_txt = f"{hoy.day} de {MESES_ES[hoy.month-1]} de {hoy.year}"
    es_pensionado = bool(emp["nss_generico"])
    nss_txt = "No requiere (trabajador pensionado, no requiere nueva alta ante el IMSS)" if es_pensionado else emp["nss"]
    tiene_cuenta = db.execute(
        "SELECT 1 FROM cuentas_bancarias WHERE empleado_id=?", (emp["id"],)).fetchone() is not None
    medio_pago = "TRANSFERENCIA BANCARIA" if tiene_cuenta else None
    clasif = emp["clasificacion"] or ""
    func = CLASIFICACION_FUNCIONES.get(clasif, CLASIFICACION_FUNCIONES_DEFAULT)

    doc = Document()
    base = doc.styles["Normal"]; base.font.name = "Arial"; base.font.size = Pt(10)

    def par(texto="", *, bold=False, center=False, size=None, space=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space)
        if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(texto)
        r.bold = bold
        if size: r.font.size = Pt(size)
        return p

    LINEA = "  ____________________________  "   # espacio para escribir a mano

    def par_seg(segmentos, *, center=False, space=6, bold_base=False):
        """Arma un p\u00e1rrafo con segmentos; un segmento (texto, True) va resaltado en amarillo."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
        for seg in segmentos:
            if isinstance(seg, tuple):
                r = p.add_run(seg[0]); r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                r.bold = bold_base
            else:
                r = p.add_run(seg); r.bold = bold_base
        return p

    def blank():
        """Espacio resaltado para llenar a mano."""
        return (LINEA, True)

    def dato(v):
        """Dato normal; si falta, queda como espacio resaltado para llenar a mano."""
        return v if (v is not None and str(v).strip()) else (LINEA, True)

    def titulo_seccion(texto):
        par(texto, bold=True, center=True, size=12, space=10)

    def bullet(texto):
        par(f"\u2610 {texto}", space=4)

    # =================== CONTRATO ===================
    titulo_seccion("CONTRATO INDIVIDUAL DE TRABAJO")
    par("RELACI\u00d3N POR TIEMPO INDETERMINADO", bold=True, center=True, space=10)
    par_seg([("CONTROL PREVIO A FIRMA: sustituir o confirmar todos los datos resaltados en "
              "amarillo. No firmar, fechar ni poner huellas mientras exista un campo pendiente. "
              "El resaltado deber\u00e1 eliminarse en la versi\u00f3n final.", True)], space=10)
    par_seg([
        "Contrato individual de trabajo que celebran, por una parte, ", dato(razon),
        ", representada por ", dato(rep), ", en su car\u00e1cter de ", blank(),
        ", a quien se denominar\u00e1 EL PATR\u00d3N; y, por la otra, ", dato(nombre),
        ", a quien se denominar\u00e1 EL TRABAJADOR; conjuntamente LAS PARTES, conforme a las "
        "siguientes declaraciones y cl\u00e1usulas.",
    ])
    par("DECLARACIONES", bold=True, center=True)
    par("I. Declara EL PATR\u00d3N", bold=True)
    par_seg([
        "Que es una sociedad legalmente constituida conforme a las leyes mexicanas mediante "
        "instrumento notarial n\u00famero ", dato(instr), " de fecha ", dato(fesc),
        ", otorgado ante la fe del Lic. ", dato(not_nom), ", Notario P\u00fablico n\u00famero ", dato(not_num),
        " de ", dato(not_ciu), ", e inscrito bajo folio mercantil electr\u00f3nico ", blank(),
        " de fecha ", blank(), ".",
    ])
    par_seg([
        "Que su Registro Federal de Contribuyentes es ", dato(rfc_emp),
        " y su registro patronal ante el IMSS es ", blank(), ".",
    ])
    par_seg([
        "Que se\u00f1ala como domicilio para efectos de este contrato el ubicado en ", dato(dom),
        ", y como correo de recursos humanos ", blank(), ".",
    ])
    par_seg([
        "Que su representante cuenta con facultades suficientes y vigentes para suscribir este "
        "contrato, seg\u00fan ", blank(), ".",
    ])
    par_seg([
        "Que requiere servicios personales subordinados para el puesto de ", dato(puesto),
        ", conforme a la descripci\u00f3n objetiva de funciones incluida en el Anexo 1.",
    ])
    par("II. Declara EL TRABAJADOR", bold=True)
    par_seg(["Nombre completo: ", dato(nombre)])
    par_seg(["Nacionalidad: ", "MEXICANA"])
    par_seg(["Fecha de nacimiento y edad al firmar: ", dato(fecha_nac_txt), " / ",
             dato(str(edad) + " a\u00f1os" if edad != "" else None)])
    par_seg(["Sexo y estado civil: ", dato(sexo_txt), " / ", blank()])
    par_seg(["CURP: ", dato(emp["curp"])])
    par_seg(["RFC: ", dato(emp["rfc"])])
    par_seg(["N\u00famero de Seguridad Social: ", dato(nss_txt)])
    par_seg(["Domicilio particular completo: ", blank()])
    par_seg(["Tel\u00e9fono y correo personal: ", blank(), " / ", blank()])
    par_seg(["Identificaci\u00f3n oficial: ", blank()])
    par("Que la informaci\u00f3n y documentos entregados son aut\u00e9nticos y se obliga a comunicar por "
        "escrito cualquier cambio de domicilio o datos de contacto dentro de los diez d\u00edas "
        "naturales siguientes. Esta manifestaci\u00f3n no implica renuncia de derechos ni autoriza "
        "descuentos no permitidos por la ley.")
    par("Que cuenta con la capacidad, experiencia y aptitudes necesarias para prestar los "
        "servicios contratados, y que informar\u00e1 de inmediato cualquier restricci\u00f3n m\u00e9dica o "
        "condici\u00f3n de riesgo relevante para realizar de manera segura las labores, sin perjuicio "
        "de sus derechos de privacidad y no discriminaci\u00f3n.")
    par("III. Declaran LAS PARTES", bold=True)
    par_seg([
        "Que existe o ha existido prestaci\u00f3n de servicios previa. Para evitar contradicciones, "
        "reconocen como fecha real de ingreso y antig\u00fcedad la que resulte de la evidencia laboral "
        "y de seguridad social: ", dato(fecha_alta_txt), ".",
    ])
    par("Que este instrumento documenta hacia futuro las condiciones vigentes de una relaci\u00f3n por "
        "tiempo indeterminado y no reduce salarios, prestaciones ni derechos ya devengados. Por "
        "existir servicios previos, no se establece un nuevo periodo a prueba ni de capacitaci\u00f3n "
        "inicial.")
    par_seg([
        "Que a la fecha de firma existe el siguiente instrumento colectivo aplicable: ", blank(), ".",
    ])
    par("CL\u00c1USULAS", bold=True, center=True)
    par("PRIMERA. NATURALEZA Y ANTIG\u00dcEDAD", bold=True)
    par_seg([
        "La relaci\u00f3n de trabajo es por tiempo indeterminado. EL PATR\u00d3N reconoce la antig\u00fcedad de "
        "EL TRABAJADOR desde ", dato(fecha_alta_txt),
        ". Cualquier diferencia entre esta fecha y documentos anteriores deber\u00e1 aclararse con base "
        "en los registros del IMSS, n\u00f3mina, transferencias y controles de asistencia, sin simular "
        "ni alterar hechos pasados.",
    ])
    par("SEGUNDA. PUESTO, \u00c1REA Y DEPENDENCIA", bold=True)
    par_seg([
        "EL TRABAJADOR prestar\u00e1 servicios como ", dato(puesto), ", adscrito al \u00e1rea de ",
        dato(obra_nom), " y reportar\u00e1 a ", blank(), " o a quien legalmente lo sustituya. El puesto "
        "no es de confianza salvo que las funciones reales satisfagan los requisitos de la Ley "
        "Federal del Trabajo.",
    ])
    par("TERCERA. SERVICIOS Y ALCANCE DEL PUESTO", bold=True)
    par("Las funciones se describen con precisi\u00f3n en el Anexo 1, que forma parte integrante del "
        "contrato. EL TRABAJADOR ejecutar\u00e1 tambi\u00e9n actividades conexas o complementarias "
        "compatibles con su categor\u00eda, capacitaci\u00f3n, aptitudes y condiciones de seguridad, cuando "
        "sean razonablemente instruidas por EL PATR\u00d3N.")
    par("No se entender\u00e1 autorizada la operaci\u00f3n de maquinaria, trabajo en alturas, izaje, "
        "espacios confinados, soldadura especializada, maniobras el\u00e9ctricas o cualquier actividad "
        "de riesgo para la que EL TRABAJADOR no cuente con capacitaci\u00f3n, autorizaci\u00f3n, permiso de "
        "trabajo y equipo de protecci\u00f3n aplicables.")
    par("CUARTA. SUBORDINACI\u00d3N E INSTRUCCIONES", bold=True)
    par("EL TRABAJADOR prestar\u00e1 el servicio bajo la direcci\u00f3n de EL PATR\u00d3N, por conducto de sus "
        "supervisores autorizados, y cumplir\u00e1 instrucciones relacionadas con el trabajo, calidad, "
        "seguridad, uso de materiales y programaci\u00f3n. Las instrucciones no podr\u00e1n implicar "
        "renuncia de derechos, reducci\u00f3n salarial ni exposici\u00f3n a un riesgo grave e inminente.")
    par("Las \u00f3rdenes relevantes de producci\u00f3n, cambios de frente, autorizaciones de tiempo "
        "extraordinario, entrega de equipo y observaciones de desempe\u00f1o deber\u00e1n documentarse por "
        "medios verificables.")
    par("QUINTA. LUGAR DE TRABAJO Y ASIGNACIONES TEMPORALES", bold=True)
    par_seg([
        "El centro de trabajo ordinario ser\u00e1 ", blank(), ". Por la naturaleza de la construcci\u00f3n y "
        "fabricaci\u00f3n, EL TRABAJADOR podr\u00e1 ser asignado temporalmente a talleres, obras o "
        "instalaciones de clientes en ", dato(estado_obra or None),
        ", siempre que la asignaci\u00f3n sea l\u00edcita, razonable, compatible con el puesto y comunicada "
        "por escrito.",
    ])
    par("Cuando la asignaci\u00f3n requiera traslado fuera de la residencia habitual o genere gastos "
        "extraordinarios, EL PATR\u00d3N cubrir\u00e1 o reembolsar\u00e1, contra comprobaci\u00f3n y conforme a "
        "pol\u00edtica entregada, los conceptos legalmente procedentes de transporte, hospedaje y "
        "alimentaci\u00f3n. Ning\u00fan traslado reducir\u00e1 el salario o las prestaciones.")
    par("SEXTA. JORNADA Y HORARIO ORDINARIO", bold=True)
    par_seg([
        "La jornada ordinaria vigente al firmar ser\u00e1: lunes a viernes de 8:00 a 18:00 horas y "
        "s\u00e1bados de 8:00 a 12:00 horas, con un total de 48 horas efectivas semanales. El periodo "
        "para alimentos ser\u00e1 de una hora diaria y EL TRABAJADOR S\u00cd podr\u00e1 disponer libremente de ese "
        "tiempo y salir del \u00e1rea de trabajo. Si debe permanecer a disposici\u00f3n de EL PATR\u00d3N, el "
        "tiempo se computar\u00e1 como jornada.",
    ])
    par("La jornada efectiva nunca exceder\u00e1 el m\u00e1ximo legal aplicable. Conforme al r\u00e9gimen "
        "transitorio vigente: 48 horas semanales en 2026; 46 en 2027; 44 en 2028; 42 en 2029; y "
        "40 a partir de 2030. LAS PARTES ajustar\u00e1n por escrito la distribuci\u00f3n antes de cada "
        "reducci\u00f3n, sin disminuir salario ni prestaciones.")
    par("La distribuci\u00f3n del horario podr\u00e1 modificarse por acuerdo escrito para atender "
        "necesidades de producci\u00f3n u obra, respetando los l\u00edmites diarios, descansos, vida digna "
        "y dem\u00e1s disposiciones legales.")
    par("S\u00c9PTIMA. DESCANSOS", bold=True)
    par_seg([
        "El d\u00eda de descanso semanal ordinario ser\u00e1 ", blank(), ", con goce de salario \u00edntegro. El "
        "trabajo en domingo generar\u00e1 la prima dominical legal y el trabajo en d\u00eda de descanso "
        "semanal u obligatorio se pagar\u00e1 conforme a la ley, sin que se considere incluido en el "
        "salario ordinario.",
    ])
    par("OCTAVA. ASISTENCIA Y REGISTRO DE JORNADA", bold=True)
    par("EL TRABAJADOR registrar\u00e1 personalmente sus entradas, salidas, descansos y, en su caso, "
        "tiempo extraordinario mediante el sistema autorizado. Queda prohibido registrar por otra "
        "persona, alterar controles o trabajar fuera de registro.")
    par("EL PATR\u00d3N conservar\u00e1 los controles de jornada y proporcionar\u00e1, cuando legalmente proceda, "
        "acceso o constancia. El uso de biometr\u00eda estar\u00e1 sujeto a un aviso de privacidad "
        "espec\u00edfico, medidas de seguridad y finalidad laboral leg\u00edtima. La falla del sistema "
        "deber\u00e1 reportarse de inmediato al supervisor y a recursos humanos por un medio "
        "verificable.")
    par("NOVENA. TIEMPO EXTRAORDINARIO", bold=True)
    par("El tiempo extraordinario s\u00f3lo se laborar\u00e1 por necesidad excepcional y con autorizaci\u00f3n "
        "previa y verificable del supervisor designado. La falta de autorizaci\u00f3n no elimina el "
        "pago del tiempo efectivamente laborado cuando EL PATR\u00d3N lo orden\u00f3, conoci\u00f3 o toler\u00f3; "
        "podr\u00e1, sin embargo, dar lugar a medidas disciplinarias v\u00e1lidas si se incumpli\u00f3 el "
        "procedimiento.")
    par("Las horas extraordinarias se registrar\u00e1n y pagar\u00e1n con los recargos legales. Durante "
        "2026 y 2027, el l\u00edmite ordinario transitorio es de nueve horas extraordinarias por "
        "semana; despu\u00e9s se aplicar\u00e1n los m\u00e1ximos vigentes. La suma de jornada ordinaria y "
        "extraordinaria no podr\u00e1 superar el l\u00edmite diario legal.")
    par("D\u00c9CIMA. SALARIO Y FORMA DE PAGO", bold=True)
    par_seg([
        "EL TRABAJADOR percibir\u00e1 un salario bruto semanal de ", dato(f"${sueldo_semanal:,.2f}" if sueldo_semanal else None),
        f" ({pesos_letra(sueldo_semanal)})" if sueldo_semanal else blank(),
        ", equivalente para efectos ordinarios a un salario diario base de ",
        dato(f"${salario_diario:,.2f}" if salario_diario else None),
        ". Se pagar\u00e1 el d\u00eda ", blank(), " mediante ", dato(medio_pago),
        ", en la cuenta designada por EL TRABAJADOR, sin costo para \u00e9ste.",
    ])
    par("El salario semanal incluye el pago del descanso semanal en los t\u00e9rminos legales, pero no "
        "incluye horas extraordinarias efectivamente laboradas, primas, trabajo en d\u00edas de "
        "descanso u obligatorio, vi\u00e1ticos ni otras prestaciones que deban pagarse por separado.")
    par("EL PATR\u00d3N expedir\u00e1 los CFDI de n\u00f3mina y comprobantes correspondientes. EL TRABAJADOR "
        "revisar\u00e1 y reportar\u00e1 discrepancias; la firma o recepci\u00f3n de un recibo no implica renuncia "
        "a diferencias que legalmente procedan.")
    par("D\u00c9CIMA PRIMERA. DEDUCCIONES", bold=True)
    par("\u00danicamente se efectuar\u00e1n retenciones y descuentos permitidos por la ley y con los "
        "requisitos aplicables. No habr\u00e1 multas, descuentos autom\u00e1ticos por herramientas, "
        "materiales, errores, retardos o da\u00f1os, ni compensaciones contra salario fuera de los "
        "casos legalmente autorizados.")
    par("La falta de prestaci\u00f3n efectiva por una ausencia podr\u00e1 reflejarse proporcionalmente en "
        "n\u00f3mina, sin perjuicio de analizar justificantes, incapacidades y derechos de descanso. "
        "Cualquier convenio de descuento permitido deber\u00e1 identificar monto, causa y forma de "
        "amortizaci\u00f3n.")
    par("D\u00c9CIMA SEGUNDA. PRESTACIONES", bold=True)
    par("EL TRABAJADOR disfrutar\u00e1 de vacaciones conforme a su antig\u00fcedad, con al menos doce d\u00edas "
        "continuos cuando nazca el primer derecho, salvo que decida distribuirlos, y de una prima "
        "vacacional no inferior al veinticinco por ciento. Las vacaciones se conceder\u00e1n dentro de "
        "los seis meses siguientes al aniversario correspondiente y se documentar\u00e1n.")
    par("EL PATR\u00d3N pagar\u00e1 aguinaldo anual no inferior a quince d\u00edas de salario antes del veinte de "
        "diciembre, o la parte proporcional; participaci\u00f3n de utilidades cuando proceda; descansos "
        "obligatorios; prima dominical; permisos de paternidad; licencias e incapacidades, y las "
        "dem\u00e1s prestaciones previstas por la ley o condiciones m\u00e1s favorables vigentes.")
    par("Cualquier bono o incentivo se regir\u00e1 por un plan escrito que identifique periodo, "
        "m\u00e9tricas, condiciones de devengo y fecha de pago. Su denominaci\u00f3n no alterar\u00e1 la "
        "integraci\u00f3n salarial que legalmente corresponda.")
    par("D\u00c9CIMA TERCERA. SEGURIDAD SOCIAL", bold=True)
    par_seg([
        "EL PATR\u00d3N mantendr\u00e1 a EL TRABAJADOR inscrito ante el IMSS con el salario base de "
        "cotizaci\u00f3n legalmente integrado y efectuar\u00e1 las aportaciones al INFONAVIT y al sistema de "
        "ahorro para el retiro. Fecha de alta o reingreso: ", dato(fecha_alta_txt),
        "; salario base de cotizaci\u00f3n inicial o vigente: ",
        dato("No aplica (pensionado)" if es_pensionado else (f"${salario_cotizacion:,.2f}" if salario_cotizacion else None)),
        ".",
    ])
    par("D\u00c9CIMA CUARTA. CAPACITACI\u00d3N Y EVALUACI\u00d3N", bold=True)
    par("EL TRABAJADOR participar\u00e1 en los planes y programas de capacitaci\u00f3n, seguridad, calidad y "
        "productividad aplicables, asistir\u00e1 puntualmente y presentar\u00e1 las evaluaciones "
        "relacionadas con su puesto. EL PATR\u00d3N documentar\u00e1 temarios, asistencia, resultados, "
        "constancias y acciones de refuerzo.")
    par("Las evaluaciones ser\u00e1n objetivas, conocidas, relacionadas con las funciones del Anexo 1 y "
        "no discriminatorias. Una evaluaci\u00f3n desfavorable no sustituye por s\u00ed sola las causas ni el "
        "procedimiento legal de rescisi\u00f3n.")
    par("D\u00c9CIMA QUINTA. SEGURIDAD, SALUD Y EQUIPO DE PROTECCI\u00d3N", bold=True)
    par("EL PATR\u00d3N proporcionar\u00e1 capacitaci\u00f3n, herramientas y equipo de protecci\u00f3n personal "
        "adecuados. EL TRABAJADOR se obliga a usarlos correctamente, participar en pl\u00e1ticas de "
        "seguridad, respetar permisos de trabajo, bloqueo y se\u00f1alizaci\u00f3n, y reportar de inmediato "
        "actos inseguros, incidentes, casi accidentes, defectos de equipo y riesgos.")
    par("EL TRABAJADOR no retirar\u00e1 guardas ni anular\u00e1 dispositivos de seguridad y podr\u00e1 "
        "interrumpir la actividad y comunicarla al supervisor cuando perciba un riesgo grave e "
        "inminente, sin abandonar injustificadamente el centro de trabajo.")
    par("D\u00c9CIMA SEXTA. HERRAMIENTAS, MATERIALES Y RESGUARDOS", bold=True)
    par("Las entregas y devoluciones se documentar\u00e1n mediante resguardos con descripci\u00f3n, n\u00famero "
        "de serie, estado, fecha y firmas. EL TRABAJADOR conservar\u00e1 los bienes con diligencia, los "
        "usar\u00e1 s\u00f3lo para fines autorizados y reportar\u00e1 p\u00e9rdidas o da\u00f1os de inmediato.")
    par("La responsabilidad se determinar\u00e1 despu\u00e9s de escuchar a EL TRABAJADOR y revisar evidencia "
        "de entrega, uso, desgaste normal, capacitaci\u00f3n y causa del da\u00f1o. No se descontar\u00e1 "
        "cantidad alguna del salario salvo que la ley lo permita y se cumplan todos sus "
        "requisitos.")
    par("D\u00c9CIMA S\u00c9PTIMA. CALIDAD, CONDUCTA Y PROHIBICIONES", bold=True)
    par("EL TRABAJADOR cumplir\u00e1 planos, especificaciones, tolerancias, \u00f3rdenes de fabricaci\u00f3n, "
        "inspecciones y controles de calidad aplicables; mantendr\u00e1 orden y limpieza; tratar\u00e1 con "
        "respeto al personal, clientes y proveedores; y evitar\u00e1 violencia, hostigamiento, acoso, "
        "discriminaci\u00f3n, amenazas, ri\u00f1as, actos deshonestos y da\u00f1os intencionales.")
    par("Queda prohibido presentarse o laborar bajo efectos que comprometan la seguridad, consumir "
        "alcohol o drogas durante la jornada o introducir armas, salvo herramientas de trabajo "
        "expresamente autorizadas. Las medidas preventivas y pruebas m\u00e9dicas deber\u00e1n ser "
        "pertinentes al riesgo, respetuosas de la dignidad y confidenciales.")
    par("D\u00c9CIMA OCTAVA. CONFIDENCIALIDAD", bold=True)
    par("EL TRABAJADOR mantendr\u00e1 reservada la informaci\u00f3n t\u00e9cnica, comercial, financiera, de "
        "costos, cotizaciones, clientes, proveedores, planos, listas de materiales, procesos y "
        "datos personales a la que acceda por raz\u00f3n del trabajo, cuando no sea p\u00fablica y EL "
        "PATR\u00d3N adopte medidas razonables para conservarla confidencial.")
    par("La obligaci\u00f3n no impide revelar informaci\u00f3n a una autoridad competente, denunciar hechos "
        "il\u00edcitos, ejercer derechos laborales o usar conocimientos generales y experiencia "
        "profesional. Al terminar la relaci\u00f3n, EL TRABAJADOR devolver\u00e1 soportes, documentos y "
        "accesos. EL PATR\u00d3N podr\u00e1 ejercer las acciones legales procedentes por uso o revelaci\u00f3n "
        "il\u00edcitos, sin retener salarios ni prestaciones.")
    par("D\u00c9CIMA NOVENA. DOCUMENTOS Y RESULTADOS DE TRABAJO", bold=True)
    par("Los planos modificados, listas de corte, formatos, reportes, procedimientos y dem\u00e1s "
        "resultados elaborados dentro de las funciones, con recursos de EL PATR\u00d3N y para sus "
        "proyectos, pertenecer\u00e1n a EL PATR\u00d3N en la medida permitida por la legislaci\u00f3n aplicable, "
        "respetando los derechos morales irrenunciables.")
    par("EL TRABAJADOR no retirar\u00e1 originales ni copias sin autorizaci\u00f3n y entregar\u00e1 archivos, "
        "contrase\u00f1as institucionales, avances y documentaci\u00f3n al cambiar de puesto o terminar la "
        "relaci\u00f3n.")
    par("VIG\u00c9SIMA. DATOS PERSONALES", bold=True)
    par("EL PATR\u00d3N tratar\u00e1 los datos personales y, en su caso, sensibles de EL TRABAJADOR conforme "
        "al aviso de privacidad para personal que se entregar\u00e1 por separado. S\u00f3lo se recabar\u00e1n "
        "datos necesarios para la relaci\u00f3n laboral, seguridad social, seguridad y cumplimiento "
        "legal, con las medidas de protecci\u00f3n aplicables.")
    par("La firma de este contrato no sustituye el aviso de privacidad ni implica consentimiento "
        "ilimitado. Los sistemas biom\u00e9tricos, videovigilancia o evaluaciones m\u00e9dicas deber\u00e1n "
        "informarse espec\u00edficamente cuando se utilicen.")
    par("VIG\u00c9SIMA PRIMERA. AVISOS Y ACTUALIZACI\u00d3N DE DATOS", bold=True)
    par("Los avisos operativos podr\u00e1n realizarse personalmente, por correo institucional, "
        "plataforma o mensajer\u00eda al dato registrado, conservando evidencia. EL TRABAJADOR "
        "actualizar\u00e1 domicilio y contacto. Los avisos de rescisi\u00f3n o actos que exijan formalidad "
        "se practicar\u00e1n exactamente en la forma prevista por la ley.")
    par("VIG\u00c9SIMA SEGUNDA. REGLAMENTO INTERIOR Y DISCIPLINA", bold=True)
    par_seg([
        "Ser\u00e1 aplicable exclusivamente el Reglamento Interior de Trabajo de EL PATR\u00d3N que haya "
        "sido elaborado, depositado, entregado y publicado conforme a la ley. Datos de dep\u00f3sito: ",
        blank(), ".",
    ])
    par("Las medidas disciplinarias deber\u00e1n estar previstas en un reglamento v\u00e1lido, ser "
        "proporcionales, documentadas y aplicarse despu\u00e9s de escuchar a EL TRABAJADOR. No se "
        "impondr\u00e1n multas. La suspensi\u00f3n disciplinaria, cuando legalmente proceda, no exceder\u00e1 de "
        "ocho d\u00edas.")
    par("VIG\u00c9SIMA TERCERA. AUSENCIAS, RETARDOS Y PERMISOS", bold=True)
    par("EL TRABAJADOR notificar\u00e1 ausencias o retardos tan pronto como sea razonablemente posible "
        "y entregar\u00e1 los justificantes correspondientes. Las incapacidades se acreditar\u00e1n con la "
        "documentaci\u00f3n del IMSS o la legalmente procedente.")
    par("Las faltas se analizar\u00e1n individualmente; \u00fanicamente podr\u00e1n producir consecuencias "
        "conforme a la ley, al reglamento v\u00e1lidamente depositado y a la evidencia. No se "
        "considerar\u00e1 autom\u00e1ticamente toda ausencia como falta grave.")
    par("VIG\u00c9SIMA CUARTA. RESCISI\u00d3N Y TERMINACI\u00d3N", bold=True)
    par("La relaci\u00f3n s\u00f3lo podr\u00e1 rescindirse o terminarse por las causas y mediante los "
        "procedimientos previstos en la Ley Federal del Trabajo. Si EL PATR\u00d3N rescinde, entregar\u00e1 "
        "aviso escrito que identifique claramente las conductas y fechas, o lo comunicar\u00e1 al "
        "Tribunal competente dentro del plazo legal.")
    par("No se firmar\u00e1n renuncias, finiquitos, constancias de terminaci\u00f3n ni hojas en blanco "
        "anticipadamente. Todo convenio de terminaci\u00f3n o liquidaci\u00f3n deber\u00e1 contener hechos y "
        "conceptos desglosados y, para m\u00e1xima certeza, ratificarse ante el Centro de Conciliaci\u00f3n "
        "o Tribunal competente.")
    par("VIG\u00c9SIMA QUINTA. MODIFICACIONES Y CONDICIONES M\u00c1S FAVORABLES", bold=True)
    par("Las modificaciones a condiciones esenciales constar\u00e1n por escrito y respetar\u00e1n la ley y "
        "derechos adquiridos. La tolerancia u omisi\u00f3n aislada no modifica por s\u00ed sola el contrato; "
        "las condiciones m\u00e1s favorables que se acrediten continuar\u00e1n vigentes cuando legalmente "
        "corresponda.")
    par("La nulidad de una estipulaci\u00f3n no afectar\u00e1 las dem\u00e1s; ser\u00e1 sustituida por la disposici\u00f3n "
        "legal aplicable sin renuncia de derechos.")
    par("VIG\u00c9SIMA SEXTA. INSTRUMENTOS COLECTIVOS Y POL\u00cdTICAS", bold=True)
    par("Si existe contrato colectivo aplicable, sus condiciones prevalecer\u00e1n cuando sean m\u00e1s "
        "favorables o legalmente obligatorias. Las pol\u00edticas t\u00e9cnicas, de seguridad, calidad, "
        "vi\u00e1ticos, uso de sistemas y protecci\u00f3n de informaci\u00f3n complementar\u00e1n este contrato sin "
        "reducir derechos ni sustituir un Reglamento Interior de Trabajo cuando \u00e9ste sea exigible.")
    par("VIG\u00c9SIMA S\u00c9PTIMA. LEGISLACI\u00d3N Y AUTORIDAD COMPETENTE", bold=True)
    par("Para lo no previsto se aplicar\u00e1 la Ley Federal del Trabajo y dem\u00e1s normas vigentes. LAS "
        "PARTES acudir\u00e1n primero a la conciliaci\u00f3n prejudicial cuando sea obligatoria y, en su "
        "caso, al Tribunal laboral competente determinado por la ley. No existe renuncia "
        "anticipada a competencia, derechos o acciones.")
    par("VIG\u00c9SIMA OCTAVA. LECTURA, COPIAS Y FIRMA", bold=True)
    par_seg([
        "LAS PARTES declaran que leyeron el contrato y sus anexos, pudieron formular preguntas, "
        "comprenden su alcance y reciben un ejemplar \u00edntegro en la fecha de firma. Se firma en ",
        dato(estado_obra or None), " el ", fecha_txt, ".",
    ], space=24)

    par("FIRMAS DEL CONTRATO", bold=True, center=True, space=10)
    par("EL TRABAJADOR", space=2)
    par("____________________________________________", space=2)
    par(nombre, space=12)
    par("EL PATR\u00d3N", space=2)
    par("____________________________________________", space=2)
    par(f"C. {rep}", space=2)
    par_seg(["En representación legal de: ", dato(razon)], space=12)
    par("TESTIGO 1", space=2)
    par("____________________________________________", space=2)
    par("_______________________________ (nombre completo y firma)", space=12)
    par("TESTIGO 2", space=2)
    par("____________________________________________", space=2)
    par("_______________________________ (nombre completo y firma)", space=2)

    # =================== ANEXO 1 ===================
    doc.add_page_break()
    titulo_seccion("ANEXO 1. DESCRIPCI\u00d3N DEL PUESTO Y CRITERIOS DE DESEMPE\u00d1O")
    par_seg(["Puesto: ", dato(puesto)])
    par_seg(["\u00c1rea / proyecto: ", dato(obra_nom)])
    par_seg(["Jefe inmediato: ", blank()])
    par_seg(["Fecha de vigencia: ", fecha_txt])
    par("Objetivo del puesto", bold=True)
    par(func["objetivo"])
    par("Funciones esenciales", bold=True)
    for f in func["funciones"]:
        bullet(f)
    bullet("Cumplir instrucciones l\u00edcitas relacionadas con el puesto y actividades conexas "
           "compatibles con su capacidad y seguridad.")
    par("Criterios objetivos de desempe\u00f1o", bold=True)
    par_seg(["Seguridad: ", "Uso correcto de EPP, cumplimiento de permisos y procedimientos, "
             "reporte oportuno de riesgos e incidentes y ausencia de actos deliberadamente "
             "inseguros."])
    par_seg(["Calidad: ", "Cumplimiento de planos, especificaciones y tolerancias documentadas; "
             "identificaci\u00f3n de piezas y reducci\u00f3n de retrabajo atribuible a ejecuci\u00f3n."])
    par_seg(["Productividad: ", "Cumplimiento razonable de \u00f3rdenes y plazos considerando "
             "complejidad, recursos, liberaciones, capacitaci\u00f3n y condiciones de obra."])
    par_seg(["Asistencia: ", "Registros personales y completos, puntualidad y justificaci\u00f3n "
             "oportuna de ausencias conforme a la pol\u00edtica v\u00e1lida."])
    par_seg(["Cuidado de bienes: ", "Resguardos completos, uso autorizado, mantenimiento b\u00e1sico y "
             "reporte inmediato de fallas, p\u00e9rdidas o da\u00f1os."])
    par_seg(["Conducta: ", "Respeto, cooperaci\u00f3n, comunicaci\u00f3n con supervisi\u00f3n y cumplimiento de "
             "reglas contra violencia, hostigamiento, acoso y discriminaci\u00f3n."])
    par("Registro de evaluaci\u00f3n", bold=True)
    par_seg(["Periodo evaluado: ", blank()])
    par_seg(["Evidencias revisadas: ", blank()])
    par_seg(["Resultado y hechos observables: ", blank()])
    par_seg(["Apoyo o capacitaci\u00f3n acordada: ", blank()])
    par_seg(["Seguimiento: ", blank()])
    par("La firma acredita recepci\u00f3n y conocimiento del contenido, no conformidad obligatoria con "
        "una calificaci\u00f3n ni renuncia de derechos.", space=10)
    par("EL TRABAJADOR: ____________________________    FECHA: ______________")
    par("SUPERVISOR: _______________________________    FECHA: ______________")

    # =================== ANEXO 2 ===================
    doc.add_page_break()
    titulo_seccion("ANEXO 2. DESIGNACI\u00d3N DE BENEFICIARIOS")
    par("Para los efectos previstos en la Ley Federal del Trabajo respecto de salarios y "
        "prestaciones devengadas y no cobradas por fallecimiento o desaparici\u00f3n derivada de un "
        "acto delincuencial, EL TRABAJADOR designa a las siguientes personas. Los porcentajes "
        "deber\u00e1n sumar 100% y la designaci\u00f3n podr\u00e1 actualizarse por escrito.")
    for i in (1, 2, 3):
        par(f"Beneficiario {i}", bold=True)
        par_seg(["Nombre completo: ", blank()])
        par_seg(["Parentesco o relaci\u00f3n: ", blank()])
        par_seg(["CURP: ", blank()])
        par_seg(["Domicilio y tel\u00e9fono: ", blank()])
        par_seg(["Porcentaje: ", blank()])
    par("TOTAL: 100%", bold=True, space=10)
    par("Firma de EL TRABAJADOR: ____________________________________________")
    par("Lugar y fecha: ______________________________________________________")

    # =================== ANEXO 3 ===================
    doc.add_page_break()
    titulo_seccion("ANEXO 3. CONSTANCIA DE ENTREGA Y EXPEDIENTE PROBATORIO")
    par("Este anexo se completa al momento de la firma. Cada elemento entregado o recibido debe "
        "tener fecha, nombre y firma; cuando corresponda, anexar copia del comprobante.")
    for txt in [
        "Ejemplar completo del contrato y anexos entregado a EL TRABAJADOR.",
        "Copia de identificaci\u00f3n, CURP, RFC, constancia de situaci\u00f3n fiscal y NSS recibida y "
        "verificada.",
        "Aviso de alta o modificaci\u00f3n ante IMSS y salario base de cotizaci\u00f3n incorporados al "
        "expediente.",
        "Aviso de privacidad para personal entregado; aviso espec\u00edfico de "
        "biometr\u00eda/videovigilancia, si aplica.",
        "Reglamento Interior de Trabajo de la raz\u00f3n social correcta entregado, con datos de "
        "dep\u00f3sito ante CFCRL.",
        "Descripci\u00f3n de puesto, reglas de seguridad, pol\u00edtica de tiempo extraordinario y canal de "
        "reporte entregados.",
        "Equipo de protecci\u00f3n personal entregado mediante resguardo con estado, fecha y firmas.",
        "Herramientas, equipos, dispositivos o accesos entregados mediante resguardo individual.",
        "Capacitaci\u00f3n inicial y espec\u00edfica registrada con temario, instructor, fecha, evaluaci\u00f3n y "
        "constancia.",
        "Cuenta bancaria autorizada por EL TRABAJADOR para pago de n\u00f3mina y medio de entrega de "
        "CFDI registrados.",
        "Contacto de emergencia y designaci\u00f3n de beneficiarios actualizados.",
    ]:
        bullet(txt)
    par_seg(["Observaciones y documentos pendientes: ", blank()], space=10)
    par("EL TRABAJADOR confirma haber recibido \u00fanicamente los documentos marcados como "
        "entregados, sin que esta constancia sustituya su contenido ni implique renuncia de "
        "derechos.")
    par("EL TRABAJADOR: ____________________________    FECHA: ______________")
    par("POR EL PATR\u00d3N: ____________________________    FECHA: ______________")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    slug = re.sub(r"[^A-Za-z0-9]+", "_", nombre).strip("_")
    return buf.getvalue(), f"Contrato_{slug or 'trabajador'}.docx"

def generar_baja_docx(db, emp):
    """Escrito/aviso de baja de trabajador (docx) para tramite ante el IMSS."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

    P = lambda k, d="": (get_param(db, k, d) or d)
    razon = P("empresa_razon_social"); rfc_emp = P("empresa_rfc")
    dom = P("empresa_domicilio"); rep = P("empresa_representante")

    nombre = titulo(" ".join(x for x in [emp["nombre"], emp["primer_apellido"], emp["segundo_apellido"]] if x))
    puesto = titulo_obra(emp["puesto"]) if emp["puesto"] else ""
    obra = emp["obra"] or ""
    fbaja = emp["fecha_baja"] or ""
    motivo = emp["motivo_baja"] or ""
    hoy = date.today()
    fecha_txt = f"{hoy.day} de {MESES_ES[hoy.month-1]} de {hoy.year}"

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)

    def par(txt="", *, bold=False, center=False, size=None, color=None, space=6, just=False):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
        p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if center else
                       WD_ALIGN_PARAGRAPH.JUSTIFY if just else WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(txt); r.bold = bold
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor(*color)
        return p

    def par_hl(pre, valor, post=".", **kw):
        """Parrafo con un dato; si el dato falta, queda resaltado para llenar a mano."""
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(kw.get("space", 6))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(pre)
        if valor and str(valor).strip():
            r = p.add_run(str(valor)); r.bold = True
        else:
            r = p.add_run("  ____________________  "); r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        p.add_run(post)
        return p

    if razon:
        par(razon, bold=True, center=True, size=13, color=(0x16, 0x23, 0x3C), space=2)
    if rfc_emp or dom:
        par(f"RFC: {rfc_emp}   {dom}", center=True, size=9, color=(0x55, 0x55, 0x55), space=12)
    par("AVISO DE BAJA DE TRABAJADOR", bold=True, center=True, size=13, color=(0xE1, 0x28, 0x1A), space=14)
    par(f"Ciudad de México, a {fecha_txt}.", space=14)

    par("A QUIEN CORRESPONDA / INSTITUTO MEXICANO DEL SEGURO SOCIAL:", bold=True, space=10)
    par(f"Por medio del presente, la empresa {razon} hace constar la BAJA del trabajador "
        f"que a continuación se detalla, para los efectos administrativos y de seguridad "
        f"social correspondientes:", just=True, space=12)

    par_hl("Nombre del trabajador: ", nombre)
    par_hl("Número de Seguridad Social (NSS): ", emp["nss"])
    par_hl("CURP: ", emp["curp"])
    par_hl("RFC: ", emp["rfc"])
    par_hl("Puesto: ", puesto)
    par_hl("Obra / centro de trabajo: ", obra)
    par_hl("Fecha de ingreso: ", emp["fecha_alta"])
    par_hl("Fecha de baja (último día laborado): ", fbaja)
    par_hl("Motivo de la baja: ", motivo, space=14)

    par("Se solicita realizar los trámites correspondientes ante el IMSS y demás "
        "instancias aplicables. Sin otro particular, quedo a sus órdenes.", just=True, space=28)

    par("_________________________________________", center=True, space=2)
    par_hl2 = par(f"{rep}", center=True, bold=True, space=0) if rep else par("____________________", center=True, space=0)
    par(f"En representación de {razon}", center=True, size=10, space=0)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    slug = re.sub(r"[^A-Za-z0-9]+", "_", nombre).strip("_")
    return buf.getvalue(), f"Aviso_baja_{slug or emp['cedula']}.docx"


def crear_contrato_para(db, empleado_id, avisar_a=None):
    """Genera y guarda el contrato de un empleado. Si 'avisar_a' es un correo,
    lo envia como adjunto. Devuelve el id del contrato o None."""
    emp = db.execute(
        "SELECT e.*, p.nombre AS puesto, p.clasificacion, p.sueldo_semanal, p.viaticos_semanales, "
        "o.estado, o.nombre AS obra "
        "FROM empleados e JOIN puestos p ON p.id=e.puesto_id "
        "JOIN obras o ON o.id=p.obra_id WHERE e.id=?", (empleado_id,)).fetchone()
    if not emp:
        return None
    datos, archivo = generar_contrato_docx(db, emp)
    cur = db.execute(
        "INSERT INTO contratos(empleado_id, nombre_archivo, generado_en, generado_por, "
        "estatus, contenido) VALUES(?,?,?,?, 'generado', ?)",
        (empleado_id, archivo, datetime.now().isoformat(timespec="seconds"),
         session.get("nombre", ""), datos))
    registrar_bitacora(db, "Contrato generado",
                       f"{emp['nombre']} {emp['primer_apellido']} (cedula {emp['cedula']})")
    if avisar_a:
        dest = list(dict.fromkeys(
            [avisar_a, os.environ.get("ADMIN_EMAIL", "")]))   # el administrador siempre recibe copia
        dest = [d for d in dest if d]
        enviar_correo(
            f"POLTECH - Contrato generado: {emp['primer_apellido']} {emp['nombre']}",
            (f"Se genero el contrato de {emp['nombre']} {emp['primer_apellido']} "
             f"{emp['segundo_apellido'] or ''} (cedula {emp['cedula']}, puesto {emp['puesto']}).\n"
             f"Se adjunta para su revision. Los datos que el sistema no captura quedan como linea "
             f"para completar a mano."),
            dest,
            adjuntos=[(archivo, datos,
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")])
    return cur.lastrowid

# Estados dentro de la Zona Libre de la Frontera Norte (para proponer el SM correcto)
ESTADOS_ZLFN = {"Baja California", "Sonora", "Chihuahua", "Coahuila",
                "Nuevo Leon", "Tamaulipas"}

def propuesta_salario_alta(db, estado=None):
    """Propone el salario diario de alta IMSS: salario minimo del dia + extra.
    Usa el SM de la Zona Libre de la Frontera Norte si el estado pertenece a ella."""
    try:
        extra = float(get_param(db, "propuesta_extra", "15") or 15)
    except (ValueError, TypeError):
        extra = 15.0
    clave = "sm_zlfn" if (estado in ESTADOS_ZLFN) else "sm_general"
    try:
        sm = float(get_param(db, clave, "0") or 0)
    except (ValueError, TypeError):
        sm = 0.0
    return round(sm + extra, 2)


# ---------------------------------------------------------------------------
# Rutas: sesion
# ---------------------------------------------------------------------------
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        u = request.form.get("username", "").strip()
        p = request.form.get("password", "")
        row = get_db().execute("SELECT * FROM users WHERE username=?", (u,)).fetchone()
        if row and check_password_hash(row["password_hash"], p):
            session["user_id"] = row["id"]
            session["nombre"] = row["nombre"]
            session["role"] = row["role"]
            session["username"] = row["username"]
            return redirect(url_for("dashboard"))
        flash("Usuario o contrasena incorrectos.", "danger")
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# ---------------------------------------------------------------------------
# Dashboard
# ---------------------------------------------------------------------------
@app.route("/")
@login_required
def dashboard():
    db = get_db()
    vis = obras_del_usuario(db)
    if vis is None:
        stats = {
            "empleados": db.execute("SELECT COUNT(*) FROM empleados WHERE estatus='activo'").fetchone()[0],
            "obras": db.execute("SELECT COUNT(*) FROM obras").fetchone()[0],
            "puestos": db.execute("SELECT COUNT(*) FROM puestos").fetchone()[0],
            "cuentas": db.execute("SELECT COUNT(*) FROM cuentas_bancarias").fetchone()[0],
        }
    elif vis:
        ph = ",".join("?" * len(vis))
        stats = {
            "empleados": db.execute(f"SELECT COUNT(*) FROM empleados e JOIN puestos p ON p.id=e.puesto_id "
                                    f"WHERE e.estatus='activo' AND p.obra_id IN ({ph})", vis).fetchone()[0],
            "obras": len(vis),
            "puestos": db.execute(f"SELECT COUNT(*) FROM puestos WHERE obra_id IN ({ph})", vis).fetchone()[0],
            "cuentas": db.execute(f"SELECT COUNT(*) FROM cuentas_bancarias c JOIN empleados e ON e.id=c.empleado_id "
                                  f"JOIN puestos p ON p.id=e.puesto_id WHERE p.obra_id IN ({ph})", vis).fetchone()[0],
        }
    else:
        stats = {"empleados": 0, "obras": 0, "puestos": 0, "cuentas": 0}
    usa_default = os.environ.get("ADMIN_PASSWORD") is None
    return render_template("dashboard.html", stats=stats, usa_default=usa_default)

# ---------------------------------------------------------------------------
# Obras
# ---------------------------------------------------------------------------
@app.route("/obras", methods=["GET", "POST"])
@login_required
def obras():
    db = get_db()
    if request.method == "POST":
        if role_rank(session["role"]) < GERENTE_RANK:
            abort(403)
        db.execute(
            "INSERT INTO obras(proyecto, contrato, nombre, estado) VALUES(?,?,?,?)",
            (request.form.get("proyecto", "").strip(),
             request.form.get("contrato", "").strip(),
             request.form.get("nombre", "").strip(),
             request.form.get("estado", "")),
        )
        db.commit()
        flash("Obra registrada.", "success")
        return redirect(url_for("obras"))
    vis = obras_del_usuario(db)
    if vis is None:
        filas = db.execute("SELECT * FROM obras ORDER BY nombre").fetchall()
    elif vis:
        ph = ",".join("?" * len(vis))
        filas = db.execute(f"SELECT * FROM obras WHERE id IN ({ph}) ORDER BY nombre", vis).fetchall()
    else:
        filas = []
    return render_template("obras_list.html", obras=filas, estados=ESTADOS)


@app.route("/obra/<int:obra_id>/editar", methods=["GET", "POST"])
@min_rank(GERENTE_RANK)
def obra_editar(obra_id):
    db = get_db()
    obra = db.execute("SELECT * FROM obras WHERE id=?", (obra_id,)).fetchone()
    if not obra:
        abort(404)
    if not puede_ver_obra(db, obra_id):
        abort(403)
    if request.method == "POST":
        db.execute(
            "UPDATE obras SET proyecto=?, contrato=?, nombre=?, estado=? WHERE id=?",
            (request.form.get("proyecto", "").strip(),
             request.form.get("contrato", "").strip(),
             request.form.get("nombre", "").strip() or obra["nombre"],
             request.form.get("estado", ""), obra_id))
        db.commit()
        registrar_bitacora(db, "Edicion de obra", f"Obra {obra_id}")
        db.commit()
        flash("Obra actualizada.", "success")
        return redirect(url_for("obras"))
    return render_template("obra_form.html", obra=obra, estados=ESTADOS)

# ---------------------------------------------------------------------------
# Catalogo de sueldos por obra (solo gerente de obra hacia arriba)
# ---------------------------------------------------------------------------
@app.route("/catalogo", methods=["GET", "POST"])
@login_required
def catalogo():
    db = get_db()
    if request.method == "POST":
        if role_rank(session["role"]) < GERENTE_RANK:
            abort(403)
        db.execute(
            "INSERT INTO puestos(obra_id, nombre, sueldo_semanal, viaticos_semanales, clasificacion) "
            "VALUES(?,?,?,?,?)",
            (request.form.get("obra_id"),
             titulo(request.form.get("nombre", "").strip()),
             float(request.form.get("sueldo_semanal") or 0),
             float(request.form.get("viaticos_semanales") or 0),
             request.form.get("clasificacion", "").strip() or None),
        )
        db.commit()
        flash("Puesto agregado al catalogo.", "success")
        return redirect(url_for("catalogo"))
    vis = obras_del_usuario(db)
    filtro_obra = (request.args.get("obra_id") or "").strip()
    filtro_clasif = (request.args.get("clasificacion") or "").strip()
    where, args = [], []
    if vis is None:
        obras_l = db.execute("SELECT * FROM obras ORDER BY nombre").fetchall()
    elif vis:
        ph = ",".join("?" * len(vis))
        obras_l = db.execute(f"SELECT * FROM obras WHERE id IN ({ph}) ORDER BY nombre", vis).fetchall()
        where.append(f"o.id IN ({ph})"); args += vis
    else:
        obras_l = []
        where.append("0")
    if filtro_obra:
        where.append("o.id = ?"); args.append(filtro_obra)
    if filtro_clasif:
        where.append("p.clasificacion = ?"); args.append(filtro_clasif)
    sql = ("SELECT p.*, p.nombre AS puesto, o.nombre AS obra, o.estado FROM puestos p "
           "JOIN obras o ON o.id=p.obra_id")
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY o.nombre, p.nombre"
    puestos = db.execute(sql, args).fetchall()
    clasif_rows = db.execute(
        "SELECT c.id, c.nombre, "
        "(SELECT COUNT(*) FROM puestos p WHERE p.clasificacion = c.nombre) AS usos "
        "FROM clasificaciones c ORDER BY c.nombre"
    ).fetchall()
    return render_template("catalogo_list.html", obras=obras_l, puestos=puestos,
                           filtro_obra=filtro_obra, filtro_clasif=filtro_clasif,
                           clasificaciones=[c["nombre"] for c in clasif_rows],
                           clasif_rows=clasif_rows)


@app.route("/clasificacion/agregar", methods=["POST"])
@min_rank(GERENTE_RANK)
def clasificacion_agregar():
    db = get_db()
    nombre = titulo(request.form.get("nombre", "").strip())
    if not nombre:
        flash("Escribe un nombre para la clasificacion.", "warning")
    else:
        try:
            db.execute("INSERT INTO clasificaciones(nombre) VALUES(?)", (nombre,))
            db.commit()
            flash(f"Clasificacion '{nombre}' agregada.", "success")
        except sqlite3.IntegrityError:
            flash(f"La clasificacion '{nombre}' ya existe.", "warning")
    return redirect(url_for("catalogo"))


@app.route("/clasificacion/<int:clasificacion_id>/eliminar", methods=["POST"])
@min_rank(GERENTE_RANK)
def clasificacion_eliminar(clasificacion_id):
    db = get_db()
    c = db.execute("SELECT * FROM clasificaciones WHERE id=?", (clasificacion_id,)).fetchone()
    if not c:
        abort(404)
    usos = db.execute("SELECT COUNT(*) FROM puestos WHERE clasificacion=?", (c["nombre"],)).fetchone()[0]
    if usos > 0:
        flash(f"No se puede eliminar '{c['nombre']}': esta asignada a {usos} puesto(s). "
              "Cambia esos puestos a otra clasificacion primero.", "danger")
    else:
        db.execute("DELETE FROM clasificaciones WHERE id=?", (clasificacion_id,))
        db.commit()
        flash(f"Clasificacion '{c['nombre']}' eliminada.", "success")
    return redirect(url_for("catalogo"))


@app.route("/puesto/<int:puesto_id>/editar", methods=["GET", "POST"])
@min_rank(GERENTE_RANK)
def puesto_editar(puesto_id):
    db = get_db()
    p = db.execute("SELECT p.*, o.nombre AS obra FROM puestos p JOIN obras o ON o.id=p.obra_id "
                   "WHERE p.id=?", (puesto_id,)).fetchone()
    if not p:
        abort(404)
    if not puede_ver_obra(db, p["obra_id"]):
        abort(403)
    if request.method == "POST":
        db.execute(
            "UPDATE puestos SET nombre=?, sueldo_semanal=?, viaticos_semanales=?, clasificacion=? "
            "WHERE id=?",
            (titulo(request.form.get("nombre", "").strip()),
             float(request.form.get("sueldo_semanal") or 0),
             float(request.form.get("viaticos_semanales") or 0),
             request.form.get("clasificacion", "").strip() or None, puesto_id))
        db.commit()
        flash("Puesto actualizado.", "success")
        return redirect(url_for("catalogo"))
    clasificaciones = [r["nombre"] for r in
                       db.execute("SELECT nombre FROM clasificaciones ORDER BY nombre").fetchall()]
    return render_template("puesto_form.html", p=p, clasificaciones=clasificaciones)


@app.route("/puesto/<int:puesto_id>/desactivar", methods=["POST"])
@min_rank(GERENTE_RANK)
def puesto_desactivar(puesto_id):
    db = get_db()
    p = db.execute("SELECT * FROM puestos WHERE id=?", (puesto_id,)).fetchone()
    if not p:
        abort(404)
    if not puede_ver_obra(db, p["obra_id"]):
        abort(403)
    db.execute("UPDATE puestos SET activo=0 WHERE id=?", (puesto_id,))
    db.commit()
    flash(f"Puesto '{p['nombre']}' desactivado. Ya no aparecera para dar de alta gente nueva, "
          "pero los trabajadores que ya lo tienen siguen funcionando normal.", "success")
    return redirect(url_for("catalogo"))


@app.route("/puesto/<int:puesto_id>/reactivar", methods=["POST"])
@min_rank(GERENTE_RANK)
def puesto_reactivar(puesto_id):
    db = get_db()
    p = db.execute("SELECT * FROM puestos WHERE id=?", (puesto_id,)).fetchone()
    if not p:
        abort(404)
    if not puede_ver_obra(db, p["obra_id"]):
        abort(403)
    db.execute("UPDATE puestos SET activo=1 WHERE id=?", (puesto_id,))
    db.commit()
    flash(f"Puesto '{p['nombre']}' reactivado.", "success")
    return redirect(url_for("catalogo"))

# ---------------------------------------------------------------------------
# Personal
# ---------------------------------------------------------------------------
@app.route("/personal")
@login_required
def personal():
    db = get_db()
    q = request.args.get("q", "").strip()
    f_obra = request.args.get("obra", "").strip()
    f_puesto = request.args.get("puesto", "").strip()
    desde = request.args.get("desde", "").strip()
    hasta = request.args.get("hasta", "").strip()

    sql = ("SELECT e.*, p.nombre AS puesto, o.nombre AS obra, o.estado AS plaza "
           "FROM empleados e "
           "LEFT JOIN puestos p ON p.id=e.puesto_id "
           "LEFT JOIN obras o ON o.id=p.obra_id WHERE 1=1")
    args = []
    if q:
        sql += (" AND (e.nombre LIKE ? OR e.primer_apellido LIKE ? OR "
                "e.segundo_apellido LIKE ? OR e.nss LIKE ? OR e.cedula LIKE ?)")
        like = f"%{q}%"
        args += [like, like, like, like, like]
    if f_obra:
        sql += " AND o.nombre = ?"; args.append(f_obra)
    if f_puesto:
        sql += " AND p.nombre = ?"; args.append(f_puesto)
    if desde:
        sql += " AND e.fecha_alta >= ?"; args.append(desde)
    if hasta:
        sql += " AND e.fecha_alta <= ?"; args.append(hasta)
    vis = obras_del_usuario(db)
    if vis is not None:
        if vis:
            ph = ",".join("?" * len(vis))
            sql += f" AND o.id IN ({ph})"; args += vis
        else:
            sql += " AND 1=0"
    sql += " ORDER BY e.primer_apellido, e.nombre"

    filas = []
    for r in db.execute(sql, args).fetchall():
        d = dict(r)
        texto, clase = dias_para_retencion(r)
        d["dias_texto"] = texto
        d["dias_clase"] = clase
        filas.append(d)

    obras = db.execute("SELECT DISTINCT nombre FROM obras ORDER BY nombre").fetchall()
    puestos = db.execute("SELECT DISTINCT nombre FROM puestos ORDER BY nombre").fetchall()
    return render_template("personal_list.html", empleados=filas,
                           obras=obras, puestos=puestos,
                           filtros={"q": q, "obra": f_obra, "puesto": f_puesto,
                                    "desde": desde, "hasta": hasta})


@app.route("/personal/rellenar-fecha-solicitud", methods=["POST"])
@min_rank(ADMIN_RANK)
def personal_rellenar_fecha_solicitud():
    """Herramienta de una sola vez: a los trabajadores que ya tienen fecha de alta
    pero les falta la fecha de solicitud de alta IMSS (comun en los subidos por
    carga masiva antes de este arreglo), se las calcula: alta - 5 dias.
    No pisa ninguna fecha de solicitud que ya este capturada."""
    db = get_db()
    filas = db.execute(
        "SELECT id, fecha_alta FROM empleados "
        "WHERE (fecha_solicitud IS NULL OR fecha_solicitud='') AND fecha_alta IS NOT NULL "
        "AND fecha_alta<>''").fetchall()
    actualizados = 0
    for r in filas:
        try:
            fs = (date.fromisoformat(str(r["fecha_alta"])[:10]) - timedelta(days=5)).isoformat()
        except ValueError:
            continue
        db.execute("UPDATE empleados SET fecha_solicitud=? WHERE id=?", (fs, r["id"]))
        actualizados += 1
    db.commit()
    registrar_bitacora(db, "Rellenar fecha de solicitud IMSS",
                        f"{actualizados} trabajador(es) actualizados (alta - 5 dias)")
    db.commit()
    flash(f"Se calculo la fecha de solicitud IMSS para {actualizados} trabajador(es) "
          "que no la tenian.", "success")
    return redirect(url_for("personal"))


@app.route("/personal/<int:emp_id>/editar", methods=["GET", "POST"])
@min_rank(ADMIN_RANK)
def personal_editar(emp_id):
    db = get_db()
    emp = db.execute("SELECT * FROM empleados WHERE id=?", (emp_id,)).fetchone()
    if not emp:
        abort(404)
    puestos = db.execute(
        "SELECT p.id, p.nombre AS puesto, p.sueldo_semanal, p.viaticos_semanales, "
        "p.obra_id, o.nombre AS obra, o.estado FROM puestos p JOIN obras o ON o.id=p.obra_id "
        "WHERE p.activo=1 OR p.id=? ORDER BY o.nombre, p.nombre", (emp["puesto_id"],)).fetchall()
    obras_form = db.execute("SELECT * FROM obras ORDER BY nombre").fetchall()
    propuesta_por_obra = {o["id"]: propuesta_salario_alta(db, o["estado"]) for o in obras_form}

    if request.method == "POST":
        f = request.form
        nombre = titulo(f.get("nombre", "").strip())
        ap1 = titulo(f.get("primer_apellido", "").strip())
        curp = f.get("curp", "").strip().upper()
        pensionado = 1 if f.get("nss_generico") else 0
        autoriza_pension = f.get("autoriza_nss_generico", "").strip()
        errores = []
        if not nombre: errores.append("El nombre es obligatorio.")
        if not ap1: errores.append("El primer apellido es obligatorio.")
        if not f.get("puesto_id"): errores.append("Selecciona un puesto.")
        if not curp_valida(curp):
            errores.append("La CURP no tiene un formato valido.")
        if pensionado and role_rank(session.get("role", "")) < GERENTE_RANK:
            errores.append("Solo el superintendente o el administrador pueden marcar NSS generico de pensionados.")

        if pensionado:
            if not autoriza_pension:
                errores.append("Para el NSS generico de pensionados, indica que administrador autorizo.")
            nss = NSS_GENERICO
        else:
            nss = solo_digitos(f.get("nss", ""))
            if nss and not nss_valido(nss):
                errores.append("El NSS no es valido.")

        err_sol = validar_solicitud_imss(f.get("fecha_solicitud", ""), f.get("fecha_alta", ""))
        if err_sol: errores.append(err_sol)
        # NSS unico excluyendo al propio empleado (el NSS generico si se puede compartir)
        if nss and nss != NSS_GENERICO:
            otro = db.execute(
                "SELECT cedula FROM empleados WHERE nss=? AND id<>?",
                (nss, emp_id)).fetchone()
            if otro:
                errores.append(f"Otro trabajador ya tiene ese NSS (cedula {otro['cedula']}).")

        if errores:
            for e in errores: flash(e, "danger")
            datos = dict(emp); datos.update(f)
            return render_template("personal_form.html", puestos=puestos,
                                   obras=obras_form, propuesta_por_obra=propuesta_por_obra,
                                   estados=ESTADOS, tipos=TIPOS_CUENTA, bancos=BANCOS,
                                   estatus_docs=ESTATUS_DOCS, datos=datos,
                                   editar=True, emp_id=emp_id)

        # Si no viene la fecha de solicitud (el JS no corrio), se calcula: alta - 5 dias.
        fecha_sol = f.get("fecha_solicitud", "").strip()
        if not fecha_sol and f.get("fecha_alta"):
            try:
                fecha_sol = (date.fromisoformat(f.get("fecha_alta")) - timedelta(days=5)).isoformat()
            except ValueError:
                fecha_sol = ""

        db.execute(
            """UPDATE empleados SET nombre=?, primer_apellido=?, segundo_apellido=?,
               curp=?, rfc=?, cp_fiscal=?, nss=?, sexo=?, fecha_nacimiento=?,
               puesto_id=?, fecha_alta=?, fecha_solicitud=?, importe_alta_imss=?,
               infonavit_monto=?, bono_semanal=?,
               nss_generico=?, autoriza_nss_generico=?,
               estatus_docs=?, observaciones=? WHERE id=?""",
            (nombre, ap1, titulo(f.get("segundo_apellido", "").strip()), curp,
             f.get("rfc", "").strip().upper(), f.get("cp_fiscal", "").strip(),
             nss, f.get("sexo", ""), f.get("fecha_nacimiento", ""),
             f.get("puesto_id"), f.get("fecha_alta"), fecha_sol,
             float(f.get("importe_alta_imss") or 0), float(f.get("infonavit_monto") or 0),
             float(f.get("bono_semanal") or 0),
             pensionado, autoriza_pension,
             f.get("estatus_docs", "Pendiente de carga"),
             f.get("observaciones", "").strip(), emp_id))
        if pensionado and not emp["nss_generico"]:
            registrar_bitacora(db, "NSS generico de pensionado",
                                f"Empleado {emp['cedula']} ({nombre} {ap1}), autorizo: {autoriza_pension}")
        db.commit()
        flash(f"Empleado actualizado (cedula {emp['cedula']}).", "success")
        return redirect(url_for("personal"))

    return render_template("personal_form.html", puestos=puestos,
                           obras=obras_form, propuesta_por_obra=propuesta_por_obra,
                           estados=ESTADOS, tipos=TIPOS_CUENTA, bancos=BANCOS,
                           estatus_docs=ESTATUS_DOCS, datos=dict(emp),
                           editar=True, emp_id=emp_id)


# ---------------------------------------------------------------------------
# Documentos del empleado (subida de PDFs)  -- superintendente / residente
# ---------------------------------------------------------------------------
@app.route("/personal/<int:emp_id>/documentos")
@login_required
def documentos_empleado(emp_id):
    db = get_db()
    emp = db.execute(
        "SELECT e.*, p.nombre AS puesto, o.nombre AS obra FROM empleados e "
        "LEFT JOIN puestos p ON p.id=e.puesto_id LEFT JOIN obras o ON o.id=p.obra_id "
        "WHERE e.id=?", (emp_id,)).fetchone()
    if not emp:
        abort(404)
    subidos = {r["tipo"]: r for r in db.execute(
        "SELECT id, tipo, filename, subido_en, subido_por FROM documentos "
        "WHERE empleado_id=?", (emp_id,)).fetchall()}
    return render_template("documentos_empleado.html", emp=emp,
                           tipos_doc=TIPOS_DOC, requeridos=DOCS_REQUERIDOS,
                           subidos=subidos)


@app.route("/personal/<int:emp_id>/documentos/subir", methods=["POST"])
@login_required
def documento_subir(emp_id):
    db = get_db()
    tipo = request.form.get("tipo", "")
    archivo = request.files.get("archivo")
    if tipo not in [t[0] for t in TIPOS_DOC]:
        flash("Tipo de documento no valido.", "danger")
        return redirect(url_for("documentos_empleado", emp_id=emp_id))
    if not archivo or not archivo.filename.lower().endswith(".pdf"):
        flash("Sube un archivo en formato PDF.", "danger")
        return redirect(url_for("documentos_empleado", emp_id=emp_id))
    contenido = archivo.read()
    db.execute("DELETE FROM documentos WHERE empleado_id=? AND tipo=?", (emp_id, tipo))
    db.execute(
        "INSERT INTO documentos(empleado_id, tipo, filename, contenido, subido_en, subido_por) "
        "VALUES(?,?,?,?,?,?)",
        (emp_id, tipo, archivo.filename, contenido, date.today().isoformat(),
         session.get("nombre", "")))
    recomputar_estatus_docs(db, emp_id)
    db.commit()
    flash("Documento subido.", "success")
    return redirect(url_for("documentos_empleado", emp_id=emp_id))


@app.route("/documento/<int:doc_id>")
@login_required
def documento_ver(doc_id):
    db = get_db()
    d = db.execute("SELECT filename, contenido FROM documentos WHERE id=?", (doc_id,)).fetchone()
    if not d:
        abort(404)
    return send_file(io.BytesIO(d["contenido"]), mimetype="application/pdf",
                     download_name=d["filename"] or "documento.pdf")


@app.route("/personal/<int:emp_id>/documentos/<tipo>/eliminar", methods=["POST"])
@login_required
def documento_eliminar(emp_id, tipo):
    db = get_db()
    db.execute("DELETE FROM documentos WHERE empleado_id=? AND tipo=?", (emp_id, tipo))
    recomputar_estatus_docs(db, emp_id)
    db.commit()
    flash("Documento eliminado.", "success")
    return redirect(url_for("documentos_empleado", emp_id=emp_id))


@app.route("/personal/<int:emp_id>/documentos/validar", methods=["POST"])
@min_rank(GERENTE_RANK)
def documentos_validar(emp_id):
    db = get_db()
    tipos = [r["tipo"] for r in db.execute(
        "SELECT tipo FROM documentos WHERE empleado_id=?", (emp_id,)).fetchall()]
    if not all(t in tipos for t in DOCS_REQUERIDOS):
        flash("Faltan documentos requeridos para poder validar.", "danger")
    else:
        db.execute("UPDATE empleados SET estatus_docs='Validado' WHERE id=?", (emp_id,))
        db.commit()
        flash("Documentacion validada.", "success")
    return redirect(url_for("documentos_empleado", emp_id=emp_id))

@app.route("/personal/nuevo", methods=["GET", "POST"])
@login_required
def personal_nuevo():
    db = get_db()
    vis = obras_del_usuario(db)
    if vis is None:
        puestos = db.execute(
            "SELECT p.id, p.nombre AS puesto, p.sueldo_semanal, p.viaticos_semanales, "
            "p.obra_id, o.nombre AS obra, o.estado FROM puestos p JOIN obras o ON o.id=p.obra_id "
            "WHERE p.activo=1 ORDER BY o.nombre, p.nombre").fetchall()
    elif vis:
        ph = ",".join("?" * len(vis))
        puestos = db.execute(
            "SELECT p.id, p.nombre AS puesto, p.sueldo_semanal, p.viaticos_semanales, "
            "p.obra_id, o.nombre AS obra, o.estado FROM puestos p JOIN obras o ON o.id=p.obra_id "
            f"WHERE p.activo=1 AND p.obra_id IN ({ph}) ORDER BY o.nombre, p.nombre", vis).fetchall()
    else:
        puestos = []
    obras_form = obras_visibles(db)
    # propuesta de salario de alta por obra (SM de su zona + extra)
    propuesta_por_obra = {o["id"]: propuesta_salario_alta(db, o["estado"]) for o in obras_form}
    propuesta_salario = propuesta_salario_alta(db)   # SM general + extra (default)

    if request.method == "POST":
        f = request.form
        nombre = titulo(f.get("nombre", "").strip())
        ap1 = titulo(f.get("primer_apellido", "").strip())
        ap2 = titulo(f.get("segundo_apellido", "").strip())
        curp = f.get("curp", "").strip().upper()
        nss = f.get("nss", "").strip()
        sexo = f.get("sexo", "")
        fecha_nac = f.get("fecha_nacimiento", "")
        exime = 1 if f.get("exime_docs") else 0
        autoriza = f.get("autoriza_tercero", "").strip()
        pensionado = 1 if f.get("nss_generico") else 0
        autoriza_pension = f.get("autoriza_nss_generico", "").strip()

        errores = []
        if not nombre: errores.append("El nombre es obligatorio.")
        if not ap1: errores.append("El primer apellido es obligatorio.")
        if not f.get("puesto_id"): errores.append("Selecciona un puesto (obra + categoria).")
        elif f.get("puesto_id") not in {str(p["id"]) for p in puestos}:
            errores.append("El puesto seleccionado no pertenece a una obra que puedas gestionar.")
        if not f.get("fecha_alta"): errores.append("La fecha de alta es obligatoria.")
        err_sol = validar_solicitud_imss(f.get("fecha_solicitud", ""), f.get("fecha_alta", ""))
        if err_sol: errores.append(err_sol)
        if exime and role_rank(session.get("role", "")) < GERENTE_RANK:
            errores.append("Solo el superintendente o el administrador pueden eximir documentos.")
        if pensionado and role_rank(session.get("role", "")) < GERENTE_RANK:
            errores.append("Solo el superintendente o el administrador pueden marcar NSS generico de pensionados.")

        # CURP: formato obligatorio
        if not curp_valida(curp):
            errores.append("La CURP no tiene un formato valido (18 caracteres).")

        # NSS: obligatorio, salvo que se exima con autorizacion de un tercero, o que sea
        # personal pensionado con NSS generico (autorizado por un administrador).
        if pensionado:
            if not autoriza_pension:
                errores.append("Para el NSS generico de pensionados, indica que administrador autorizo.")
            nss_norm = NSS_GENERICO
        elif exime:
            if not autoriza:
                errores.append("Para eximir documentos, indica quien autoriza (tercero).")
            nss_norm = ""
        else:
            if not nss:
                errores.append("El NSS es obligatorio (o marca 'eximir' con autorizacion).")
            elif not nss_valido(nss):
                errores.append("El NSS no es valido (deben ser 11 digitos con digito verificador correcto).")
            nss_norm = solo_digitos(nss)

        # NSS unico: no puede repetirse (salvo el NSS generico de pensionados, que si se comparte)
        if nss_norm and nss_norm != NSS_GENERICO:
            ya = db.execute(
                "SELECT cedula, nombre, primer_apellido FROM empleados WHERE nss=?",
                (nss_norm,)).fetchone()
            if ya:
                errores.append(
                    f"Ya existe un trabajador con ese NSS (cedula {ya['cedula']}: "
                    f"{ya['nombre']} {ya['primer_apellido']}).")

        # Cuenta bancaria opcional
        cta_num = solo_digitos(f.get("cta_numero", ""))
        cta_inst = f.get("cta_institucion", "").strip()
        cta_tipo = f.get("cta_tipo", "")
        if cta_num:
            dup = db.execute(
                "SELECT e.nombre, e.primer_apellido, e.segundo_apellido, e.estatus, "
                "o.nombre AS obra FROM cuentas_bancarias c "
                "JOIN empleados e ON e.id=c.empleado_id "
                "LEFT JOIN puestos p ON p.id=e.puesto_id "
                "LEFT JOIN obras o ON o.id=p.obra_id "
                "WHERE c.numero=?", (cta_num,)).fetchone()
            if dup:
                titular = " ".join(x for x in [dup["nombre"], dup["primer_apellido"],
                                               dup["segundo_apellido"]] if x)
                obra_dup = dup["obra"] or "sin obra"
                estatus_dup = dup["estatus"] or "activo"
                errores.append(
                    f"Esa cuenta ya esta registrada a nombre de {titular} "
                    f"(obra: {obra_dup}, estatus: {estatus_dup}). No se permiten duplicados.")
                # Avisar al administrador por correo (no bloquea si el correo no esta configurado)
                nuevo = " ".join(x for x in [nombre, ap1, ap2] if x)
                send_admin_alert(
                    "POLTECH - Intento de cuenta bancaria duplicada",
                    (f"Se intento dar de alta a '{nuevo}' con la cuenta {cta_num}, "
                     f"que ya pertenece a '{titular}' (obra: {obra_dup}, "
                     f"estatus: {estatus_dup}).\n\nRevisar antes de continuar."))
            if not cta_inst or not cta_tipo:
                errores.append("Para la cuenta bancaria, captura institucion y tipo de cuenta.")
            err_cta = validar_cuenta(cta_tipo, cta_num)
            if err_cta:
                errores.append(err_cta)

        if errores:
            for e in errores:
                flash(e, "danger")
            return render_template("personal_form.html", puestos=puestos,
                                   obras=obras_form, propuesta_por_obra=propuesta_por_obra,
                                   propuesta_salario=propuesta_salario,
                                   estados=ESTADOS, tipos=TIPOS_CUENTA, bancos=BANCOS, datos=f)

        # Si no viene la fecha de solicitud (el JS no corrio), se calcula: alta - 5 dias.
        fecha_sol = f.get("fecha_solicitud", "").strip()
        if not fecha_sol and f.get("fecha_alta"):
            try:
                fecha_sol = (date.fromisoformat(f.get("fecha_alta")) - timedelta(days=5)).isoformat()
            except ValueError:
                fecha_sol = ""

        cedula = siguiente_cedula(db)
        observaciones = f.get("observaciones", "").strip()
        cur = db.execute(
            """INSERT INTO empleados
               (cedula, nombre, primer_apellido, segundo_apellido, curp, rfc, cp_fiscal,
                nss, sexo, fecha_nacimiento, puesto_id, fecha_alta, fecha_solicitud,
                fecha_registro, importe_alta_imss, infonavit_monto,
                bono_semanal, exime_docs, autoriza_tercero, nss_generico,
                autoriza_nss_generico, observaciones, estatus_docs)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (cedula, nombre, ap1, ap2, curp,
             f.get("rfc", "").strip().upper(), f.get("cp_fiscal", "").strip(),
             nss_norm, sexo, fecha_nac, f.get("puesto_id"),
             f.get("fecha_alta"), fecha_sol,
             date.today().isoformat(), float(f.get("importe_alta_imss") or 0),
             float(f.get("infonavit_monto") or 0),
             float(f.get("bono_semanal") or 0),
             exime, autoriza, pensionado, autoriza_pension,
             observaciones, "Pendiente de carga"),
        )
        emp_id = cur.lastrowid
        if cta_num:
            db.execute(
                "INSERT INTO cuentas_bancarias(empleado_id, institucion, tipo_cuenta, numero) VALUES(?,?,?,?)",
                (emp_id, cta_inst, cta_tipo, cta_num),
            )
        if pensionado:
            registrar_bitacora(db, "NSS generico de pensionado",
                                f"Empleado {cedula} ({nombre} {ap1}), autorizo: {autoriza_pension}")
        db.commit()

        # Generar el contrato automaticamente y avisar por correo a quien dio el alta
        try:
            crear_contrato_para(db, emp_id, avisar_a=session.get("username"))
            db.commit()
        except Exception as e:
            app.logger.error("No se pudo generar el contrato: %s", e)

        for a in revisar_curp(curp, nombre, ap1, ap2, sexo, fecha_nac):
            flash("Aviso: " + a, "warning")
        flash(f"Empleado dado de alta con cedula {cedula}. Se genero su contrato para revision.", "success")
        return redirect(url_for("personal"))

    return render_template("personal_form.html", puestos=puestos,
                           obras=obras_form, propuesta_por_obra=propuesta_por_obra,
                           propuesta_salario=propuesta_salario,
                           estados=ESTADOS, tipos=TIPOS_CUENTA, bancos=BANCOS, datos={})

# ---------------------------------------------------------------------------
# Carga masiva de trabajadores (por Excel)
# ---------------------------------------------------------------------------
COLS_CARGA = ["NOMBRE(S)", "PRIMER APELLIDO", "SEGUNDO APELLIDO", "CURP", "RFC",
              "CP FISCAL", "NSS", "SEXO (H/M)", "FECHA NACIMIENTO (DD-MM-AAAA)",
              "OBRA", "PUESTO (referencia, no se usa)", "FECHA ALTA (DD-MM-AAAA)",
              "SALARIO ALTA IMSS", "INFONAVIT SEMANAL", "BANCO", "TIPO DE CUENTA",
              "NUMERO DE CUENTA", "OBSERVACIONES"]

@app.route("/personal/plantilla")
@login_required
def personal_plantilla():
    db = get_db()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Trabajadores"
    ws.append(COLS_CARGA)
    ws.append(["JUAN", "PEREZ", "LOPEZ", "PELJ800101HDFRPN09", "PELJ800101AB1",
               "54000", "12345678903", "H", "01-01-1980", "Torre Reforma",
               "(no se usa)", "28-07-2026", "3000", "0", "BBVA Mexico",
               "CLABE interbancaria", "012180001234567890", ""])
    from openpyxl.styles import Font, PatternFill
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="16233C")
    for i in range(1, len(COLS_CARGA) + 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = 20

    # -----------------------------------------------------------------
    # Hoja oculta "Listas": aqui viven los valores de los menus desplegables
    # (Sexo, Obra, Banco, Tipo de cuenta). El Puesto ya no tiene desplegable:
    # queda de referencia y no se usa para elegir el puesto (ver personal_carga).
    # -----------------------------------------------------------------
    from openpyxl.workbook.defined_name import DefinedName
    from openpyxl.worksheet.datavalidation import DataValidation

    obras_l = obras_visibles(db)
    lst = wb.create_sheet("Listas")
    lst.sheet_state = "hidden"

    def nombrar_rango(nombre, ref):
        wb.defined_names[nombre] = DefinedName(nombre, attr_text=ref)

    # Columna A: nombre de la obra.
    for idx, o in enumerate(obras_l, start=1):
        lst.cell(row=idx, column=1, value=o["nombre"])
    if obras_l:
        nombrar_rango("ListaObras", f"Listas!$A$1:$A${len(obras_l)}")

    # Columna D: bancos. Columna E: tipo de cuenta.
    for idx, b in enumerate(BANCOS, start=1):
        lst.cell(row=idx, column=4, value=b)
    nombrar_rango("ListaBancos", f"Listas!$D$1:$D${len(BANCOS)}")
    for idx, t in enumerate(TIPOS_CUENTA, start=1):
        lst.cell(row=idx, column=5, value=t)
    nombrar_rango("ListaTipoCuenta", f"Listas!$E$1:$E${len(TIPOS_CUENTA)}")

    # -----------------------------------------------------------------
    # Menus desplegables en la hoja "Trabajadores"
    # -----------------------------------------------------------------
    ultima_fila = 500  # deja espacio para pegar/escribir muchos trabajadores

    dv_sexo = DataValidation(type="list", formula1='"H,M"', allow_blank=True,
                             showErrorMessage=True, errorTitle="Valor no valido",
                             error="Escribe H (hombre) o M (mujer).")
    ws.add_data_validation(dv_sexo)
    dv_sexo.add(f"H2:H{ultima_fila}")

    if obras_l:
        dv_obra = DataValidation(type="list", formula1="ListaObras", allow_blank=True,
                                  showErrorMessage=True, errorTitle="Obra no valida",
                                  error="Elige una obra de la lista (ya dada de alta en Obras).")
        ws.add_data_validation(dv_obra)
        dv_obra.add(f"J2:J{ultima_fila}")

    dv_banco = DataValidation(type="list", formula1="ListaBancos", allow_blank=True,
                              showErrorMessage=False)
    ws.add_data_validation(dv_banco)
    dv_banco.add(f"O2:O{ultima_fila}")

    dv_tipo = DataValidation(type="list", formula1="ListaTipoCuenta", allow_blank=True,
                             showErrorMessage=False)
    ws.add_data_validation(dv_tipo)
    dv_tipo.add(f"P2:P{ultima_fila}")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    nombre = f"Plantilla de alta de trabajadores_{datetime.now():%Y-%m-%d_%H%M}.xlsx"
    return send_file(buf, as_attachment=True, download_name=nombre,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.route("/personal/carga", methods=["GET", "POST"])
@login_required
def personal_carga():
    if request.method == "POST":
        archivo = request.files.get("archivo")
        if not archivo or not archivo.filename.lower().endswith(".xlsx"):
            flash("Sube un archivo de Excel (.xlsx) usando la plantilla.", "danger")
            return render_template("personal_carga.html", resumen=None)

        db = get_db()
        wb = openpyxl.load_workbook(archivo, data_only=True)
        ws = wb.active
        creados = 0
        contratos_gen = 0
        errores = []
        duplicadas = []   # cuentas bancarias repetidas
        creados_detalle = []   # para que revisen el sueldo de cada quien despues

        for i, fila in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if fila is None or all(c is None or str(c).strip() == "" for c in fila):
                continue
            vals = (list(fila) + [None] * len(COLS_CARGA))[:len(COLS_CARGA)]
            (nombre, ap1, ap2, curp, rfc, cp, nss, sexo, fnac, obra_nom, puesto_nom,
             falta, salario, infonavit, banco, tipo_cta, num_cta, obs) = vals

            def txt(x): return "" if x is None else str(x).strip()
            fecha_txt = normalizar_fecha
            nombre, ap1, ap2 = titulo(txt(nombre)), titulo(txt(ap1)), titulo(txt(ap2))
            curp = txt(curp).upper()
            nss = solo_digitos(txt(nss))
            # el NSS a veces llega como numero con ".0"
            if nss.endswith(".0"):
                nss = nss[:-2]
            fnac = fecha_txt(fnac)
            falta = fecha_txt(falta)
            nombre_completo = " ".join(x for x in [nombre, ap1, ap2] if x) or "sin nombre"

            errs = []
            if not nombre: errs.append("nombre vacio")
            if not ap1: errs.append("primer apellido vacio")
            if not curp_valida(curp): errs.append("CURP invalida")
            if not nss: errs.append("NSS vacio")
            elif not nss_valido(nss): errs.append("NSS invalido")
            if fnac and not fecha_valida(fnac):
                errs.append(f"fecha de nacimiento '{fnac}' no reconocida (usa DD-MM-AAAA)")
            if not falta:
                errs.append("fecha de alta vacia")
            elif not fecha_valida(falta):
                errs.append(f"fecha de alta '{falta}' no reconocida (usa DD-MM-AAAA)")
            puesto = resolver_puesto(db, obra_nom)
            if not puesto:
                errs.append(f"la obra '{txt(obra_nom)}' no existe en el catalogo o no tiene puestos activos")
            # NSS duplicado
            if nss and db.execute("SELECT 1 FROM empleados WHERE nss=?", (nss,)).fetchone():
                errs.append("ya existe un trabajador con ese NSS")

            if errs:
                errores.append(f"Fila {i} ({nombre_completo}): " + "; ".join(errs))
                continue

            # Si no viene el salario de alta, se sugiere el minimo del estado de la obra + extra
            # (igual que en el alta individual), en vez de dejarlo en $0.
            salario_txt = txt(salario)
            if salario_txt:
                salario_final = float(salario_txt)
            else:
                salario_final = propuesta_salario_alta(db, puesto["estado"])

            # Fecha de solicitud de alta IMSS: 5 dias antes de la fecha de alta (igual
            # que en el alta individual; en la carga masiva no habia forma de calcularla).
            fecha_solicitud_calc = (date.fromisoformat(falta) - timedelta(days=5)).isoformat()

            cedula = siguiente_cedula(db)
            cur = db.execute(
                """INSERT INTO empleados
                   (cedula, nombre, primer_apellido, segundo_apellido, curp, rfc, cp_fiscal,
                    nss, sexo, fecha_nacimiento, puesto_id, fecha_alta, fecha_solicitud,
                    fecha_registro, importe_alta_imss, infonavit_monto, exime_docs,
                    autoriza_tercero, observaciones, estatus_docs)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (cedula, nombre, ap1, ap2, curp, txt(rfc).upper(), txt(cp), nss,
                 txt(sexo).upper()[:1], fnac, puesto["id"], falta, fecha_solicitud_calc,
                 date.today().isoformat(), salario_final, float(infonavit or 0),
                 0, "", txt(obs), "Pendiente de carga"))
            emp_id = cur.lastrowid
            creados += 1
            creados_detalle.append({
                "cedula": cedula, "nombre": nombre_completo, "obra": txt(obra_nom),
                "puesto": puesto["puesto"], "sueldo": puesto["sueldo_semanal"],
            })
            try:
                crear_contrato_para(db, emp_id, avisar_a=None)
                contratos_gen += 1
            except Exception as e:
                app.logger.error("Contrato en carga (fila %s): %s", i, e)

            # cuenta bancaria (opcional): revisar duplicado antes de insertar
            num_cta = solo_digitos(txt(num_cta))
            if num_cta:
                dup = db.execute(
                    "SELECT e.cedula, e.nombre, e.primer_apellido, e.estatus, o.nombre AS obra "
                    "FROM cuentas_bancarias c JOIN empleados e ON e.id=c.empleado_id "
                    "LEFT JOIN puestos p ON p.id=e.puesto_id "
                    "LEFT JOIN obras o ON o.id=p.obra_id WHERE c.numero=?",
                    (num_cta,)).fetchone()
                if dup:
                    titular = f"{dup['nombre']} {dup['primer_apellido']}".strip()
                    duplicadas.append(
                        f"Fila {i}: la cuenta {num_cta} de {nombre} {ap1} ya pertenece a "
                        f"{titular} (cedula {dup['cedula']}, obra {dup['obra'] or '-'}, "
                        f"estatus {dup['estatus'] or 'activo'}). No se guardo la cuenta.")
                else:
                    db.execute(
                        "INSERT INTO cuentas_bancarias(empleado_id, institucion, tipo_cuenta, numero) "
                        "VALUES(?,?,?,?)", (emp_id, txt(banco), txt(tipo_cta), num_cta))

        db.commit()

        if duplicadas:
            send_admin_alert(
                "POLTECH - Cuentas bancarias duplicadas en carga masiva",
                "Durante una carga masiva se detectaron cuentas repetidas:\n\n" +
                "\n".join(duplicadas))

        resumen = {"creados": creados, "contratos": contratos_gen,
                   "errores": errores, "duplicadas": duplicadas,
                   "creados_detalle": creados_detalle}
        return render_template("personal_carga.html", resumen=resumen)

    return render_template("personal_carga.html", resumen=None)

# ---------------------------------------------------------------------------
# Contratos generados
# ---------------------------------------------------------------------------
@app.route("/contratos")
@login_required
def contratos():
    db = get_db()
    vis = obras_del_usuario(db)
    sql = ("SELECT c.id, c.nombre_archivo, c.generado_en, c.generado_por, c.estatus, "
           "c.revisado_por, c.revisado_en, e.cedula, e.nombre, e.primer_apellido, "
           "e.segundo_apellido, p.nombre AS puesto, o.nombre AS obra, o.id AS obra_id "
           "FROM contratos c JOIN empleados e ON e.id=c.empleado_id "
           "JOIN puestos p ON p.id=e.puesto_id JOIN obras o ON o.id=p.obra_id WHERE 1=1")
    args = []
    if vis is not None:
        if vis:
            sql += " AND o.id IN (%s)" % ",".join("?" * len(vis)); args += vis
        else:
            sql += " AND 0"
    sql += " ORDER BY c.generado_en DESC"
    filas = db.execute(sql, args).fetchall()
    return render_template("contratos_list.html", contratos=filas)


@app.route("/contratos/<int:cid>/descargar")
@login_required
def contrato_descargar(cid):
    db = get_db()
    c = db.execute("SELECT c.*, p.obra_id FROM contratos c "
                   "JOIN empleados e ON e.id=c.empleado_id "
                   "JOIN puestos p ON p.id=e.puesto_id WHERE c.id=?", (cid,)).fetchone()
    if not c:
        abort(404)
    if not puede_ver_obra(db, c["obra_id"]):
        abort(403)
    return send_file(io.BytesIO(c["contenido"]), as_attachment=True,
                     download_name=c["nombre_archivo"],
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/contratos/<int:cid>/revisado", methods=["POST"])
@min_rank(GERENTE_RANK)
def contrato_revisado(cid):
    db = get_db()
    c = db.execute("SELECT c.*, p.obra_id FROM contratos c "
                   "JOIN empleados e ON e.id=c.empleado_id "
                   "JOIN puestos p ON p.id=e.puesto_id WHERE c.id=?", (cid,)).fetchone()
    if not c:
        abort(404)
    if not puede_ver_obra(db, c["obra_id"]):
        abort(403)
    db.execute("UPDATE contratos SET estatus='revisado', revisado_por=?, revisado_en=? WHERE id=?",
               (session.get("nombre"), datetime.now().isoformat(timespec="seconds"), cid))
    registrar_bitacora(db, "Contrato revisado", f"Contrato {cid}")
    db.commit()
    flash("Contrato marcado como revisado.", "success")
    return redirect(url_for("contratos"))


@app.route("/personal/<int:emp_id>/contrato", methods=["POST"])
@login_required
def contrato_regenerar(emp_id):
    db = get_db()
    emp = db.execute("SELECT e.id, p.obra_id FROM empleados e JOIN puestos p ON p.id=e.puesto_id "
                     "WHERE e.id=?", (emp_id,)).fetchone()
    if not emp:
        abort(404)
    if not puede_ver_obra(db, emp["obra_id"]):
        abort(403)
    try:
        crear_contrato_para(db, emp_id, avisar_a=session.get("username"))
        db.commit()
        flash("Contrato regenerado y enviado a tu correo.", "success")
    except Exception as e:
        app.logger.error("No se pudo regenerar el contrato: %s", e)
        flash("No se pudo generar el contrato.", "danger")
    return redirect(url_for("contratos"))


# ---------------------------------------------------------------------------
# Integracion API market (consulta de NSS por CURP y vigencia)
# ---------------------------------------------------------------------------
def _apimarket_call(db, url, params):
    """POST a un endpoint de API market. Los parametros (curp, nss) viajan en la
    URL como query string (?curp=...), con Authorization: Bearer <token>.
    Devuelve (ok, datos) donde datos es dict/lista (JSON) o texto/mensaje de error."""
    token = (get_param(db, "apimarket_token", "") or os.environ.get("APIMARKET_TOKEN", "")).strip()
    if not url:
        return False, "Falta configurar la URL del servicio en Parametros (API market)."
    if not token:
        return False, "Falta configurar el token de API market en Parametros."
    try:
        full = url
        limpios = {k: v for k, v in (params or {}).items() if v}
        if limpios:
            sep = "&" if "?" in url else "?"
            full = url + sep + urllib.parse.urlencode(limpios)
        req = urllib.request.Request(full, method="POST")
        req.add_header("Authorization", "Bearer " + token)
        req.add_header("Accept", "application/json")
        with urllib.request.urlopen(req, timeout=25) as resp:
            body = resp.read().decode("utf-8", "ignore")
        try:
            return True, json.loads(body)
        except ValueError:
            return True, body
    except urllib.error.HTTPError as e:
        detalle = e.read().decode("utf-8", "ignore")[:400] if hasattr(e, "read") else ""
        return False, f"HTTP {e.code}: {detalle}"
    except Exception as e:
        return False, str(e)


def _buscar_nss(obj):
    """Busca recursivamente un NSS (11 digitos) en la respuesta del API."""
    if isinstance(obj, dict):
        for k, v in obj.items():
            if isinstance(v, (str, int)) and re.fullmatch(r"\d{11}", str(v)):
                if any(t in k.lower() for t in ("nss", "seguro", "social")):
                    return str(v)
        for v in obj.values():
            r = _buscar_nss(v)
            if r:
                return r
    elif isinstance(obj, list):
        for v in obj:
            r = _buscar_nss(v)
            if r:
                return r
    elif isinstance(obj, (str, int)) and re.fullmatch(r"\d{11}", str(obj)):
        return str(obj)
    return None


def _vigencia_datos(resp):
    """Extrae los campos utiles de la respuesta de vigencia (dentro de 'data')."""
    d = resp.get("data") if isinstance(resp, dict) else None
    if not isinstance(d, dict):
        d = resp if isinstance(resp, dict) else {}
    def si_no(v):
        s = str(v).strip().upper()
        if s in ("SI", "SÍ", "TRUE", "1"): return "SI"
        if s in ("NO", "FALSE", "0"): return "NO"
        return str(v)
    return {
        "nss": d.get("nss") or d.get("NSS") or "",
        "servicio_medico": si_no(d.get("conDerechoSm", "")),
        "incapacidad": si_no(d.get("conDerechoInc", "")),
        "vigente_hasta": d.get("vigenteHasta") or "",
        "codigo": (resp.get("codigoValidacion") if isinstance(resp, dict) else "") or "",
        "mensaje": (resp.get("message") if isinstance(resp, dict) else "") or "",
    }


def generar_vigencia_docx(db, vig, curp):
    """Comprobante de vigencia de derechos (IMSS) en Word."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    razon = get_param(db, "empresa_razon_social", "") or ""
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)

    def par(txt, *, bold=False, center=False, size=None, color=None, space=6):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(txt); r.bold = bold
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor(*color)
        return p

    par("COMPROBANTE DE VIGENCIA DE DERECHOS (IMSS)", bold=True, center=True, size=14, color=(0x0A, 0x56, 0xC0))
    if razon:
        par(razon, center=True, size=10, color=(0x16, 0x23, 0x3C), space=12)
    par(f"Fecha de consulta: {date.today().strftime('%d/%m/%Y')}", space=12)
    par(f"Número de Seguridad Social (NSS): {vig['nss']}", bold=True, space=4)
    par(f"CURP: {curp}", bold=True, space=12)

    par(f"Con derecho a servicio médico: {vig['servicio_medico']}", space=4)
    par(f"Con derecho a incapacidades: {vig['incapacidad']}", space=4)
    if vig["vigente_hasta"]:
        par(f"Vigente hasta: {vig['vigente_hasta']}", space=12)
    else:
        par("", space=8)

    if vig["codigo"]:
        par(f"Código de validación: {vig['codigo']}", size=9, color=(0x55, 0x55, 0x55), space=4)
    par("Consulta realizada a través de API Market (fuente: IMSS). Documento informativo.",
        size=8, color=(0x88, 0x88, 0x88), space=0)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue(), f"Vigencia_{vig['nss'] or curp}.docx"


@app.route("/api/consulta-nss", methods=["POST"])
@login_required
def api_consulta_nss():
    db = get_db()
    curp = ((request.form.get("curp") or (request.json or {}).get("curp") or "")).strip().upper()
    if not re.fullmatch(r"[A-Z0-9]{18}", curp):
        return jsonify({"ok": False, "error": "CURP invalido (deben ser 18 caracteres)."})
    ok, data = _apimarket_call(db, get_param(db, "apimarket_url_nss", ""), {"curp": curp})
    if not ok:
        return jsonify({"ok": False, "error": data})
    return jsonify({"ok": True, "nss": _buscar_nss(data), "raw": data})


@app.route("/apimarket", methods=["GET", "POST"])
@login_required
def apimarket_tool():
    db = get_db()
    resultado = None
    if request.method == "POST":
        tipo = request.form.get("tipo")
        curp = (request.form.get("curp") or "").strip().upper()
        nss = solo_digitos(request.form.get("nss") or "")
        if tipo == "nss":
            ok, data = _apimarket_call(db, get_param(db, "apimarket_url_nss", ""), {"curp": curp})
            resultado = {"ok": ok, "tipo": "nss", "titulo": "Consulta de NSS por CURP",
                         "nss": _buscar_nss(data) if ok else None,
                         "data": data if ok else None, "error": None if ok else data}
        else:
            ok, data = _apimarket_call(db, get_param(db, "apimarket_url_vigencia", ""),
                                       {"nss": nss, "curp": curp})
            vig = _vigencia_datos(data) if ok else None
            resultado = {"ok": ok, "tipo": "vigencia", "titulo": "Consulta de vigencia (NSS y CURP)",
                         "vig": vig, "curp": curp, "data": data if ok else None,
                         "error": None if ok else data,
                         "vig_json": json.dumps({"vig": vig, "curp": curp}) if ok and vig else ""}
        resultado["raw"] = json.dumps(resultado.get("data"), indent=2, ensure_ascii=False) \
            if resultado.get("data") is not None else ""
    configurado = bool((get_param(db, "apimarket_token", "") or "").strip())
    return render_template("apimarket.html", resultado=resultado, configurado=configurado)


@app.route("/apimarket/vigencia/comprobante", methods=["POST"])
@login_required
def apimarket_vigencia_comprobante():
    db = get_db()
    try:
        payload = json.loads(request.form.get("vig_json") or "{}")
        vig = payload.get("vig") or {}
        curp = payload.get("curp") or ""
    except (ValueError, TypeError):
        flash("No se pudo generar el comprobante.", "danger")
        return redirect(url_for("apimarket_tool"))
    datos, archivo = generar_vigencia_docx(db, vig, curp)
    return send_file(io.BytesIO(datos), as_attachment=True, download_name=archivo,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/contratos/generar-todos", methods=["POST"])
@login_required
def contratos_generar_todos():
    db = get_db()
    vis = obras_del_usuario(db)
    sql = ("SELECT e.id FROM empleados e JOIN puestos p ON p.id=e.puesto_id "
           "WHERE (e.estatus IS NULL OR e.estatus != 'baja') "
           "AND e.id NOT IN (SELECT empleado_id FROM contratos)")
    args = []
    if vis is not None:
        if vis:
            sql += " AND p.obra_id IN (%s)" % ",".join("?" * len(vis)); args += vis
        else:
            sql += " AND 0"
    faltantes = [r["id"] for r in db.execute(sql, args).fetchall()]
    generados = 0
    for eid in faltantes:
        try:
            crear_contrato_para(db, eid, avisar_a=None)
            generados += 1
        except Exception as e:
            app.logger.error("Contrato masivo (emp %s): %s", eid, e)
    db.commit()
    if generados:
        flash(f"Se generaron {generados} contratos del personal dado de alta.", "success")
    else:
        flash("No habia contratos por generar (todo el personal activo ya tiene contrato).", "info")
    return redirect(url_for("contratos"))


@app.route("/contratos/regenerar-todos", methods=["POST"])
@min_rank(ADMIN_RANK)
def contratos_regenerar_todos():
    """Regenera el contrato de TODO el personal activo con el formato mas reciente,
    tenga o no un contrato ya generado. Cada regeneracion crea un registro nuevo en
    Contratos (no borra los anteriores, para conservar el historial)."""
    db = get_db()
    vis = obras_del_usuario(db)
    sql = ("SELECT e.id FROM empleados e JOIN puestos p ON p.id=e.puesto_id "
           "WHERE (e.estatus IS NULL OR e.estatus != 'baja')")
    args = []
    if vis is not None:
        if vis:
            sql += " AND p.obra_id IN (%s)" % ",".join("?" * len(vis)); args += vis
        else:
            sql += " AND 0"
    ids = [r["id"] for r in db.execute(sql, args).fetchall()]
    generados = 0
    for eid in ids:
        try:
            crear_contrato_para(db, eid, avisar_a=None)
            generados += 1
        except Exception as e:
            app.logger.error("Contrato regenerar-todos (emp %s): %s", eid, e)
    db.commit()
    flash(f"Se regeneraron {generados} contratos con el formato mas reciente.", "success")
    return redirect(url_for("contratos"))


# ---------------------------------------------------------------------------
# Bajas de personal (aviso/escrito de baja)
# ---------------------------------------------------------------------------
def _emp_para_baja(db, emp_id):
    return db.execute(
        "SELECT e.*, p.nombre AS puesto, o.nombre AS obra, o.id AS obra_id "
        "FROM empleados e JOIN puestos p ON p.id=e.puesto_id "
        "JOIN obras o ON o.id=p.obra_id WHERE e.id=?", (emp_id,)).fetchone()


@app.route("/personal/<int:emp_id>/baja", methods=["GET", "POST"])
@login_required
def personal_baja(emp_id):
    db = get_db()
    emp = _emp_para_baja(db, emp_id)
    if not emp:
        abort(404)
    if not puede_ver_obra(db, emp["obra_id"]):
        abort(403)
    if request.method == "POST":
        fecha = (request.form.get("fecha_baja") or "").strip()
        motivo = (request.form.get("motivo_baja") or "").strip()
        if not fecha:
            flash("Indica la fecha de baja.", "danger")
            return render_template("personal_baja.html", emp=emp, motivos=MOTIVOS_BAJA)
        db.execute("UPDATE empleados SET estatus='baja', fecha_baja=?, motivo_baja=? WHERE id=?",
                   (fecha, motivo, emp_id))
        registrar_bitacora(db, "Baja de trabajador",
                           f"{emp['nombre']} {emp['primer_apellido']} (cedula {emp['cedula']}) "
                           f"baja {fecha} - {motivo}")
        db.commit()
        send_admin_alert(
            "POLTECH - Aviso de baja de trabajador",
            (f"El trabajador {emp['nombre']} {emp['primer_apellido']} {emp['segundo_apellido'] or ''} "
             f"causa baja desde {fecha}.\nMotivo: {motivo or 'no especificado'}.\n"
             f"CURP: {emp['curp']}  RFC: {emp['rfc']}  NSS: {emp['nss']}  Puesto: {emp['puesto']}.\n\n"
             f"Favor de tramitar la baja ante el IMSS."),
            extra=[get_param(db, "email_contador", "")])
        flash("Trabajador dado de baja. Ya puedes descargar el aviso de baja.", "success")
        return redirect(url_for("bajas"))
    return render_template("personal_baja.html", emp=emp, motivos=MOTIVOS_BAJA)


@app.route("/personal/<int:emp_id>/escrito-baja")
@login_required
def escrito_baja(emp_id):
    db = get_db()
    emp = _emp_para_baja(db, emp_id)
    if not emp:
        abort(404)
    if not puede_ver_obra(db, emp["obra_id"]):
        abort(403)
    datos, archivo = generar_baja_docx(db, emp)
    return send_file(io.BytesIO(datos), as_attachment=True, download_name=archivo,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/personal/<int:emp_id>/reactivar", methods=["POST"])
@min_rank(GERENTE_RANK)
def personal_reactivar(emp_id):
    db = get_db()
    emp = _emp_para_baja(db, emp_id)
    if not emp:
        abort(404)
    if not puede_ver_obra(db, emp["obra_id"]):
        abort(403)
    db.execute("UPDATE empleados SET estatus='activo', fecha_baja=NULL, motivo_baja=NULL WHERE id=?", (emp_id,))
    registrar_bitacora(db, "Reingreso de trabajador",
                       f"{emp['nombre']} {emp['primer_apellido']} (cedula {emp['cedula']}) reactivado")
    db.commit()
    flash("Trabajador reactivado.", "success")
    return redirect(url_for("bajas"))


@app.route("/bajas")
@login_required
def bajas():
    db = get_db()
    vis = obras_del_usuario(db)
    f_obra = (request.args.get("obra_id") or "").strip()
    f_desde = (request.args.get("desde") or "").strip()
    f_hasta = (request.args.get("hasta") or "").strip()
    sql = ("SELECT e.id, e.cedula, e.nombre, e.primer_apellido, e.segundo_apellido, e.nss, "
           "e.curp, e.rfc, e.fecha_alta, e.fecha_baja, e.motivo_baja, p.nombre AS puesto, "
           "o.nombre AS obra, o.id AS obra_id FROM empleados e "
           "JOIN puestos p ON p.id=e.puesto_id JOIN obras o ON o.id=p.obra_id "
           "WHERE e.estatus='baja'")
    args = []
    if vis is not None:
        if vis:
            sql += " AND o.id IN (%s)" % ",".join("?" * len(vis)); args += vis
        else:
            sql += " AND 0"
    if f_obra:
        sql += " AND o.id=?"; args.append(f_obra)
    if f_desde:
        sql += " AND e.fecha_baja>=?"; args.append(f_desde)
    if f_hasta:
        sql += " AND e.fecha_baja<=?"; args.append(f_hasta)
    sql += " ORDER BY e.fecha_baja DESC, e.primer_apellido"
    filas = db.execute(sql, args).fetchall()
    obras_l = obras_visibles(db)
    return render_template("bajas.html", bajas=filas, obras=obras_l,
                           f_obra=f_obra, f_desde=f_desde, f_hasta=f_hasta)


@app.route("/bajas/excel")
@login_required
def bajas_excel():
    db = get_db()
    vis = obras_del_usuario(db)
    sql = ("SELECT e.cedula, e.nombre, e.primer_apellido, e.segundo_apellido, e.nss, e.curp, "
           "e.rfc, e.fecha_alta, e.fecha_baja, e.motivo_baja, p.nombre AS puesto, o.nombre AS obra "
           "FROM empleados e JOIN puestos p ON p.id=e.puesto_id JOIN obras o ON o.id=p.obra_id "
           "WHERE e.estatus='baja'")
    args = []
    if vis is not None:
        if vis:
            sql += " AND o.id IN (%s)" % ",".join("?" * len(vis)); args += vis
        else:
            sql += " AND 0"
    sql += " ORDER BY e.fecha_baja DESC"
    filas = db.execute(sql, args).fetchall()
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = Workbook(); ws = wb.active; ws.title = "Bajas"
    cab = ["CEDULA", "NOMBRE", "NSS", "CURP", "RFC", "PUESTO", "OBRA",
           "FECHA INGRESO", "FECHA BAJA", "MOTIVO"]
    ws.append(cab)
    for c in ws[1]:
        c.font = Font(name="Arial", bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="16233C")
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for e in filas:
        nom = titulo(" ".join(x for x in [e["primer_apellido"], e["segundo_apellido"], e["nombre"]] if x))
        ws.append([e["cedula"], nom, str(e["nss"] or ""), e["curp"], e["rfc"],
                   titulo_obra(e["puesto"] or ""), e["obra"], e["fecha_alta"],
                   e["fecha_baja"], e["motivo_baja"]])
    ws.column_dimensions["B"].width = 30
    for L2 in ("A", "C", "D", "E", "F", "G", "H", "I", "J"):
        ws.column_dimensions[L2].width = 16
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="Reporte_bajas.xlsx",
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ---------------------------------------------------------------------------
# Cuentas bancarias
# ---------------------------------------------------------------------------
@app.route("/cuentas")
@login_required
def cuentas():
    db = get_db()
    q = request.args.get("q", "").strip()
    sql = ("SELECT c.*, e.nombre, e.primer_apellido, e.segundo_apellido "
           "FROM cuentas_bancarias c JOIN empleados e ON e.id=c.empleado_id "
           "LEFT JOIN puestos p ON p.id=e.puesto_id "
           "LEFT JOIN obras o ON o.id=p.obra_id WHERE 1=1")
    args = []
    if q:
        sql += (" AND (e.nombre LIKE ? OR e.primer_apellido LIKE ? OR "
                "c.numero LIKE ? OR c.institucion LIKE ?)")
        like = f"%{q}%"; args += [like, like, like, like]
    vis = obras_del_usuario(db)
    if vis is not None:
        if vis:
            ph = ",".join("?" * len(vis))
            sql += f" AND o.id IN ({ph})"; args += vis
        else:
            sql += " AND 1=0"
    sql += " ORDER BY e.primer_apellido"
    filas = db.execute(sql, args).fetchall()
    return render_template("cuentas_list.html", cuentas=filas, q=q)


@app.route("/cuenta/<int:cid>/editar", methods=["GET", "POST"])
@min_rank(ADMIN_RANK)
def cuenta_editar(cid):
    db = get_db()
    c = db.execute(
        "SELECT c.*, e.nombre, e.primer_apellido, e.segundo_apellido "
        "FROM cuentas_bancarias c JOIN empleados e ON e.id=c.empleado_id WHERE c.id=?",
        (cid,)).fetchone()
    if not c:
        abort(404)
    if request.method == "POST":
        f = request.form
        inst = f.get("institucion", "").strip()
        tipo = f.get("tipo_cuenta", "").strip()
        numero = solo_digitos(f.get("numero", ""))
        errores = []
        if not numero:
            errores.append("El numero de cuenta es obligatorio.")
        err_cta = validar_cuenta(tipo, numero)
        if err_cta:
            errores.append(err_cta)
        dup = db.execute("SELECT id FROM cuentas_bancarias WHERE numero=? AND id<>?",
                         (numero, cid)).fetchone()
        if dup:
            errores.append("Ese numero de cuenta ya esta registrado en otra cuenta.")
        if errores:
            for e in errores:
                flash(e, "danger")
            datos = dict(c); datos.update({"institucion": inst, "tipo_cuenta": tipo, "numero": numero})
            return render_template("cuenta_form.html", c=datos, bancos=BANCOS, tipos=TIPOS_CUENTA)
        db.execute("UPDATE cuentas_bancarias SET institucion=?, tipo_cuenta=?, numero=? WHERE id=?",
                   (inst, tipo, numero, cid))
        db.commit()
        flash("Cuenta bancaria actualizada.", "success")
        return redirect(url_for("cuentas"))
    return render_template("cuenta_form.html", c=c, bancos=BANCOS, tipos=TIPOS_CUENTA)

# ---------------------------------------------------------------------------
# Asistencia semanal (Fase 2)
# ---------------------------------------------------------------------------
@app.route("/asistencia")
@login_required
def asistencia_inicio():
    db = get_db()
    obras = obras_visibles(db)
    vdef = viernes_de(date.today()).isoformat()
    # semanas capturadas recientes (solo obras visibles)
    filas = db.execute(
        "SELECT s.*, o.nombre AS obra, "
        "  (SELECT COUNT(*) FROM asistencia a WHERE a.semana_id=s.id) AS capturados "
        "FROM asistencia_semanas s JOIN obras o ON o.id=s.obra_id "
        "ORDER BY s.fecha_inicio DESC, o.nombre LIMIT 30").fetchall()
    if obras_del_usuario(db) is not None:
        visibles = {o["id"] for o in obras}
        filas = [f for f in filas if f["obra_id"] in visibles]
    return render_template("asistencia.html", obras=obras,
                           viernes_default=vdef, semanas=filas)


@app.route("/asistencia/captura", methods=["GET", "POST"])
@login_required
def asistencia_captura():
    db = get_db()
    if request.method == "POST":
        obra_id = int(request.form.get("obra_id") or 0)
        finicio = request.form.get("fecha_inicio", "")
    else:
        obra_id = int(request.args.get("obra_id") or 0)
        finicio = request.args.get("fecha_inicio", "")

    if not obra_id or not finicio:
        flash("Selecciona una obra y una semana.", "warning")
        return redirect(url_for("asistencia_inicio"))
    if not puede_ver_obra(db, obra_id):
        abort(403)
    try:
        f = date.fromisoformat(finicio[:10])
    except ValueError:
        flash("Fecha invalida.", "danger")
        return redirect(url_for("asistencia_inicio"))

    viernes = viernes_de(f)
    jueves = viernes + timedelta(days=6)
    obra = db.execute("SELECT * FROM obras WHERE id=?", (obra_id,)).fetchone()
    if not obra:
        flash("La obra no existe.", "danger")
        return redirect(url_for("asistencia_inicio"))

    sem = db.execute(
        "SELECT * FROM asistencia_semanas WHERE obra_id=? AND fecha_inicio=?",
        (obra_id, viernes.isoformat())).fetchone()

    empleados = db.execute(
        "SELECT e.id, e.cedula, e.nombre, e.primer_apellido, e.segundo_apellido, "
        "       p.nombre AS puesto "
        "FROM empleados e JOIN puestos p ON p.id=e.puesto_id "
        "WHERE p.obra_id=? AND e.estatus='activo' "
        "ORDER BY e.primer_apellido, e.nombre", (obra_id,)).fetchall()

    if request.method == "POST":
        if not sem:
            cur = db.execute(
                "INSERT INTO asistencia_semanas(obra_id, fecha_inicio, fecha_fin, "
                "semana_num, anio, estatus, creada_en, creada_por) "
                "VALUES(?,?,?,?,?, 'borrador', ?, ?)",
                (obra_id, viernes.isoformat(), jueves.isoformat(),
                 viernes.isocalendar()[1], viernes.year,
                 datetime.now().isoformat(timespec="seconds"), session.get("nombre")))
            semana_id = cur.lastrowid
        else:
            semana_id = sem["id"]

        for e in empleados:
            eid = e["id"]
            cods = {k: (request.form.get(f"{k}_{eid}") or "").upper()[:1] for k, _ in DIAS_SEMANA}
            def _num(name):
                try:
                    return max(0.0, float(request.form.get(name) or 0))
                except (ValueError, TypeError):
                    return 0.0
            he150 = _num(f"he150_{eid}")
            he200 = _num(f"he200_{eid}")
            obs = (request.form.get(f"obs_{eid}") or "").strip()
            db.execute(
                "INSERT INTO asistencia(semana_id, empleado_id, d1,d2,d3,d4,d5,d6,d7, "
                "he_150, he_200, observaciones) VALUES(?,?,?,?,?,?,?,?,?,?,?,?) "
                "ON CONFLICT(semana_id, empleado_id) DO UPDATE SET "
                "d1=excluded.d1, d2=excluded.d2, d3=excluded.d3, d4=excluded.d4, "
                "d5=excluded.d5, d6=excluded.d6, d7=excluded.d7, "
                "he_150=excluded.he_150, he_200=excluded.he_200, "
                "observaciones=excluded.observaciones",
                (semana_id, eid, cods["d1"], cods["d2"], cods["d3"], cods["d4"],
                 cods["d5"], cods["d6"], cods["d7"], he150, he200, obs))
        db.commit()
        flash("Asistencia guardada.", "success")
        return redirect(url_for("asistencia_captura", obra_id=obra_id,
                                fecha_inicio=viernes.isoformat()))

    # GET: cargar lo ya guardado (si existe) y armar la cuadricula
    guardado = {}
    if sem:
        for r in db.execute("SELECT * FROM asistencia WHERE semana_id=?", (sem["id"],)).fetchall():
            guardado[r["empleado_id"]] = dict(r)

    filas = []
    for e in empleados:
        g = guardado.get(e["id"])
        cods = {}
        for k, _ in DIAS_SEMANA:
            if g:
                cods[k] = g.get(k) or ""
            else:
                cods[k] = "D" if k == "d3" else "A"   # domingo descanso, resto asistio
        dias, faltas, retardos = contar_asistencia(cods)
        filas.append({
            "emp": e, "cods": cods,
            "he150": (g["he_150"] if g else 0) or 0,
            "he200": (g["he_200"] if g else 0) or 0,
            "obs": (g["observaciones"] if g else "") or "",
            "dias": dias, "faltas": faltas, "retardos": retardos,
        })

    return render_template(
        "asistencia_captura.html", obra=obra,
        viernes=viernes.isoformat(), jueves=jueves.isoformat(),
        semana_num=viernes.isocalendar()[1], anio=viernes.year,
        dias_semana=DIAS_SEMANA, dias_base=DIAS_BASE, codigos=CODIGOS_ASIS,
        fechas=fechas_de_dias(viernes), filas=filas,
        ya_existe=bool(sem))


@app.route("/asistencia/plantilla")
@login_required
def asistencia_plantilla():
    """Descarga el Excel de asistencia de una obra con el formato de 3 bloques:
    ASISTENCIA (codigos A/F/R/V/D), HORAS EXTRAS (horas por dia + TIPO + VALOR) y
    RETARDOS (por dia), mas columnas de descuentos. Hoja BAJAS de referencia."""
    db = get_db()
    obra_id = int(request.args.get("obra_id") or 0)
    finicio = request.args.get("fecha_inicio", "")
    if not obra_id or not puede_ver_obra(db, obra_id):
        flash("Selecciona una obra valida.", "warning")
        return redirect(url_for("asistencia_inicio"))
    obra = db.execute("SELECT * FROM obras WHERE id=?", (obra_id,)).fetchone()
    if not obra:
        abort(404)
    try:
        viernes = viernes_de(date.fromisoformat(finicio[:10]))
    except ValueError:
        viernes = viernes_de(date.today())
    jueves = viernes + timedelta(days=6)

    activos = db.execute(
        "SELECT e.cedula, e.nombre, e.primer_apellido, e.segundo_apellido, p.nombre AS puesto "
        "FROM empleados e JOIN puestos p ON p.id=e.puesto_id "
        "WHERE p.obra_id=? AND e.estatus='activo' "
        "ORDER BY e.primer_apellido, e.nombre", (obra_id,)).fetchall()
    bajas = db.execute(
        "SELECT e.cedula, e.nombre, e.primer_apellido, e.segundo_apellido, p.nombre AS puesto, "
        "e.fecha_alta FROM empleados e JOIN puestos p ON p.id=e.puesto_id "
        "WHERE p.obra_id=? AND e.estatus<>'activo' "
        "ORDER BY e.primer_apellido, e.nombre", (obra_id,)).fetchall()

    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.worksheet.datavalidation import DataValidation
    from openpyxl.utils import get_column_letter as L
    NAVY, ROJO, AMAR, AZUL = "16233C", "E1281A", "FFF3CD", "0A56C0"
    thin = Side(style="thin", color="BBBBBB")
    borde = Border(left=thin, right=thin, top=thin, bottom=thin)
    def nom(r):
        return " ".join(x for x in [r["primer_apellido"], r["segundo_apellido"], r["nombre"]] if x)
    def F(**k): return Font(name="Arial", **k)
    def hdr(cell, fill=NAVY):
        cell.font = F(bold=True, size=9, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=fill)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = borde

    fechas = [f"{(viernes+timedelta(days=i)):%d/%m}" for i in range(7)]
    etqs = ["VIE", "SAB", "DOM", "LUN", "MAR", "MIE", "JUE"]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Asistencia"
    ws["A1"] = "POLTECH ACERO Y CONSTRUCCION, S.A. DE C.V."
    ws["A1"].font = F(bold=True, size=13, color=NAVY)
    ws["A2"] = "CONTROL SEMANAL DE ASISTENCIA  -  Semana laboral: Viernes a Jueves"
    ws["A2"].font = F(size=10, italic=True, color="555555")
    for lc, lbl, vc, val in [
        ("A4", "OBRA:", "C4", obra["nombre"]),
        ("A5", "ESTADO:", "C5", obra["estado"]),
        ("H4", "SEMANA No.:", "J4", viernes.isocalendar()[1]),
        ("H5", "PERIODO:", "J5", f"{viernes:%d/%m/%Y} al {jueves:%d/%m/%Y}"),
        ("L4", "GERENTE DE OBRA:", "N4", ""),
        ("L5", "FECHA DE CAPTURA:", "N5", ""),
    ]:
        ws[lc] = lbl; ws[lc].font = F(bold=True, size=9, color=NAVY)
        ws[vc] = val; ws[vc].font = F(size=9)
        if val == "":
            ws[vc].fill = PatternFill("solid", fgColor=AMAR)

    # Fila 6: titulos de bloque (combinados)
    ws.merge_cells("D6:O6"); ws["D6"] = "ASISTENCIA"
    ws.merge_cells("P6:Y6"); ws["P6"] = "HORAS EXTRAS"
    ws.merge_cells("Z6:AF6"); ws["Z6"] = "RETARDOS"
    for cc, fill in (("D6", NAVY), ("P6", AZUL), ("Z6", ROJO)):
        ws[cc].font = F(bold=True, size=10, color="FFFFFF")
        ws[cc].fill = PatternFill("solid", fgColor=fill)
        ws[cc].alignment = Alignment(horizontal="center", vertical="center")

    # Fila 7: encabezados
    HROW = 7
    cab = ["No.", "CEDULA", "NOMBRE DEL TRABAJADOR"]
    cab += [f"{etqs[i]}\n{fechas[i]}" for i in range(7)]                 # D-J asistencia
    cab += ["DIAS", "FALTAS", "RET.", "VAC.", "BAJA\n(Ultimo dia trabajado)\nDD-MM-AAAA"]  # K-O
    cab += [f"{etqs[i]}\n{fechas[i]}" for i in range(7)]                 # P-V horas extra
    cab += ["TOTAL\nHORAS EXTRAS", "TIPO", "VALOR"]                      # W-Y
    cab += [f"{etqs[i]}\n{fechas[i]}" for i in range(7)]                 # Z-AF retardos
    cab += ["DESCUENTOS\nNOMINA", "DESCUENTOS A\nOTRA CUENTA"]           # AG-AH
    cab += ["OBSERVACIONES\n(a quien se deposita el desc. a otra cuenta)"]  # AI
    for i, t in enumerate(cab, start=1):
        hdr(ws.cell(HROW, i, t))

    r = HROW + 1
    for idx, emp in enumerate(activos, start=1):
        ws.cell(r, 1, idx)
        ws.cell(r, 2, emp["cedula"] or "")
        ws.cell(r, 3, nom(emp))
        # Asistencia (D-J): base=A, domingo(F col6)=D
        for col in (4, 5, 7, 8, 9, 10):
            ws.cell(r, col, "A")
        ws.cell(r, 6, "D")
        # Formulas de conteo
        ws.cell(r, 11).value = (f'=COUNTIF(D{r}:E{r},"A")+COUNTIF(G{r}:J{r},"A")'
                                f'+COUNTIF(D{r}:E{r},"R")+COUNTIF(G{r}:J{r},"R")'
                                f'+COUNTIF(D{r}:E{r},"V")+COUNTIF(G{r}:J{r},"V")')   # DIAS
        ws.cell(r, 12).value = f'=COUNTIF(D{r}:E{r},"F")+COUNTIF(G{r}:J{r},"F")'     # FALTAS
        ws.cell(r, 13).value = f'=SUM(Z{r}:AF{r})'                                   # RET.
        ws.cell(r, 14).value = f'=COUNTIF(D{r}:E{r},"V")+COUNTIF(G{r}:J{r},"V")'     # VAC
        ws.cell(r, 15, "")                                                          # BAJA
        # Horas extra (P-V) = 0, TOTAL, TIPO, VALOR
        for col in range(16, 23):
            ws.cell(r, col, 0)
        ws.cell(r, 23).value = f"=SUM(P{r}:V{r})"                                    # TOTAL HE
        ws.cell(r, 24, "")                                                          # TIPO
        ws.cell(r, 25, "")                                                          # VALOR
        # Retardos (Z-AF) = 0
        for col in range(26, 33):
            ws.cell(r, col, 0)
        ws.cell(r, 33, 0)                                                           # DESC NOMINA
        ws.cell(r, 34, 0)                                                           # DESC OTRA CUENTA
        ws.cell(r, 35, "")                                                          # OBSERVACIONES
        for col in range(1, 36):
            cell = ws.cell(r, col)
            cell.font = F(size=9); cell.border = borde
            if 4 <= col <= 34: cell.alignment = Alignment(horizontal="center")
        ws.cell(r, 35).fill = PatternFill("solid", fgColor=AMAR)
        ws.cell(r, 15).fill = PatternFill("solid", fgColor=AMAR)
        r += 1
    ult = r - 1

    if activos:
        def dv(kind, f1, rng, prompt=None, title=None):
            d = DataValidation(type=kind, formula1=f1, allow_blank=True)
            if kind == "decimal": d.operator = "greaterThanOrEqual"
            if prompt: d.prompt = prompt
            if title: d.promptTitle = title
            ws.add_data_validation(d); d.add(rng)
        dv("list", '"A,F,R,V,D"', f"D{HROW+1}:J{ult}",
           "A=Asistio  F=Falta  R=Retardo  V=Vacaciones  D=Descanso", "Asistencia")
        dv("decimal", "0", f"P{HROW+1}:V{ult}")                 # horas extra
        dv("list", '"Factor,Precio"', f"X{HROW+1}:X{ult}",
           "Factor = 1.5 o 2 veces el sueldo por hora.  Precio = precio por hora pactado.", "Tipo H.E.")
        dv("decimal", "0", f"Y{HROW+1}:Y{ult}")                 # valor
        dv("list", '"0,1"', f"Z{HROW+1}:AF{ult}",
           "1 = ese dia tuvo retardo.  0 = sin retardo.", "Retardos")
        dv("custom", "TRUE", f"O{HROW+1}:O{ult}",
           "Si hay baja, escribe la fecha como DD-MM-AAAA (ej. 15-08-2026).", "Fecha de baja")
        dv("decimal", "0", f"AG{HROW+1}:AH{ult}")               # descuentos

    anchos = ([5, 10, 30] + [6]*7 + [7, 8, 6, 6, 12] + [6]*7 + [9, 9, 8] + [6]*7 + [12, 13, 32])
    for i, w in enumerate(anchos, start=1):
        ws.column_dimensions[L(i)].width = w
    ws.freeze_panes = "D8"

    lr = ult + 2
    ws.cell(lr, 3, "Codigos de ASISTENCIA:  A = Asistio    F = Falta    R = Retardo    V = Vacaciones    D = Descanso").font = F(size=9, bold=True, color=NAVY)
    ws.cell(lr+1, 3, "RETARDOS: escribe 1 en el dia que llego tarde.  HORAS EXTRAS: horas por dia + TIPO (Factor/Precio) y VALOR.").font = F(size=8, italic=True, color="555555")
    ws.cell(lr+2, 3, "BAJA: escribe la fecha del ultimo dia trabajado si el trabajador causa baja.  El DOMINGO no cuenta para la base de 6 dias.").font = F(size=8, italic=True, color="555555")

    # ---- Hoja BAJAS (referencia) ----
    wb2 = wb.create_sheet("Bajas")
    wb2.append(["No.", "CEDULA", "NOMBRE DEL TRABAJADOR", "PUESTO", "FECHA DE ALTA"])
    for c in wb2[1]:
        c.font = F(bold=True, color="FFFFFF"); c.fill = PatternFill("solid", fgColor=ROJO)
    for idx, emp in enumerate(bajas, start=1):
        wb2.append([idx, emp["cedula"] or "", nom(emp), emp["puesto"], emp["fecha_alta"] or ""])
    for i, w in enumerate([5, 10, 30, 18, 14], start=1):
        wb2.column_dimensions[L(i)].width = w

    # ---- Hoja INSTRUCCIONES ----
    ins = wb.create_sheet("Instrucciones")
    lineas = [
        ("INSTRUCCIONES DE USO", True),
        ("", False),
        ("1. Solo se captura la hoja ASISTENCIA. La hoja BAJAS es solo de referencia.", False),
        ("2. Completa el encabezado (celdas amarillas): Gerente de obra y Fecha de captura.", False),
        ("3. BLOQUE ASISTENCIA: en cada dia escribe el codigo", False),
        ("     A = Asistio     F = Falta     R = Retardo     V = Vacaciones     D = Descanso", False),
        ("   Cada dia trae su fecha (VIE 24/07, etc.). El DOMINGO no cuenta para la base de 6 dias.", False),
        ("4. DIAS, FALTAS, RET. y VAC. se calculan solos. No los edites.", False),
        ("   DIAS pagables = dias con A, R o V (el retardo y la vacacion cuentan como dia).", False),
        ("5. BLOQUE HORAS EXTRAS: escribe las horas por dia. TOTAL se suma solo.", False),
        ("   TIPO = 'Factor' (1.5 o 2 veces el sueldo por hora) o 'Precio' (precio por hora pactado).", False),
        ("   VALOR = el factor (1.5 / 2) o el precio por hora en pesos.", False),
        ("6. BLOQUE RETARDOS: escribe 1 en el dia que el trabajador llego tarde (RET. los suma).", False),
        ("7. BAJA (ultimo dia trabajado): si ya no regreso, escribe la fecha de su ultimo dia.", False),
        ("   Con eso el sistema lo marca de baja y prepara el aviso para el IMSS.", False),
        ("8. DESCUENTOS NOMINA: importe a descontar (prestamo, herramienta, etc.).", False),
        ("   DESCUENTOS A OTRA CUENTA: importe que se le quita para depositar a otra persona.", False),
        ("9. Guarda el archivo y subelo en el sistema (Asistencia semanal > Calcular nomina).", False),
    ]
    for i, (txt, boldl) in enumerate(lineas, start=1):
        c = ins.cell(i, 1, txt)
        c.font = F(bold=boldl, size=(12 if boldl else 10), color=(NAVY if boldl else "333333"))
    ins.column_dimensions["A"].width = 108

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    obra_limpia = re.sub(r"[^0-9A-Za-zÁÉÍÓÚÑáéíóúñ ._-]+", "", obra["nombre"]).strip()
    nombre = f"Lista de asistencia semana {viernes.isocalendar()[1]} - {obra_limpia}.xlsx"
    return send_file(buf, as_attachment=True, download_name=nombre,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ---------------------------------------------------------------------------
# Nomina: subir el Excel de asistencia lleno y calcular
# ---------------------------------------------------------------------------
DIAS_BASE_COLS = [4, 5, 7, 8, 9, 10]   # D,E,G,H,I,J (sin domingo=F/col6)

def _num(v):
    try:
        return float(v)
    except (TypeError, ValueError):
        return 0.0

def _cod(v):
    return (str(v).strip().upper()[:1]) if v is not None else ""

def calcular_nomina(db, obra_id, ws, jornada, aplicar_retardos):
    """Lee la hoja 'Asistencia' (formato de 3 bloques) y calcula la nomina.
    Devuelve (detalles, avisos_baja, errores)."""
    detalles, avisos_baja, errores = [], [], []
    HROW = 7
    r = HROW + 1
    while True:
        ced = ws.cell(r, 2).value
        nombre = ws.cell(r, 3).value
        if (ced is None or str(ced).strip() == "") and (nombre is None or str(nombre).strip() == ""):
            break
        cedula = str(ced).strip() if ced else ""
        # buscar empleado por cedula en la obra
        emp = db.execute(
            "SELECT e.id, e.nombre, e.primer_apellido, e.segundo_apellido, e.curp, e.rfc, "
            "e.nss, e.infonavit_monto, e.bono_semanal AS bono_emp, "
            "p.nombre AS puesto, p.clasificacion, "
            "p.sueldo_semanal, p.viaticos_semanales "
            "FROM empleados e JOIN puestos p ON p.id=e.puesto_id "
            "WHERE e.cedula=? AND p.obra_id=?", (cedula, obra_id)).fetchone()
        if not emp:
            errores.append(f"Cedula '{cedula}' ({nombre}) no se encontro en esta obra; se omitio.")
            r += 1
            continue

        codigos = [_cod(ws.cell(r, c).value) for c in DIAS_BASE_COLS]
        dias_A = codigos.count("A")
        dias_R = codigos.count("R")
        dias_V = codigos.count("V")
        faltas = codigos.count("F")
        dias_pagables = dias_A + dias_R + dias_V
        dias_presente = dias_A + dias_R

        he_horas = sum(_num(ws.cell(r, c).value) for c in range(16, 23))   # P-V
        tipo = (str(ws.cell(r, 24).value or "").strip().lower())
        valor = _num(ws.cell(r, 25).value)
        retardos = sum(_num(ws.cell(r, c).value) for c in range(26, 33))   # Z-AF
        desc_nomina = _num(ws.cell(r, 33).value)
        desc_otra = _num(ws.cell(r, 34).value)
        obs = str(ws.cell(r, 35).value or "").strip()   # a quien se deposita el desc. a otra cuenta
        baja_fecha = normalizar_fecha(ws.cell(r, 15).value)
        if baja_fecha and not fecha_valida(baja_fecha):
            errores.append(
                f"Cedula '{cedula}' ({nombre}): la fecha de baja '{baja_fecha}' no se "
                f"reconocio (usa DD-MM-AAAA); no se registro la baja de esta fila.")
            baja_fecha = ""

        sueldo_semanal = _num(emp["sueldo_semanal"])
        # Sueldo y viaticos siempre vienen del puesto/categoria (catalogo de sueldos);
        # ya no hay un viatico individual por trabajador que lo sobre-escriba.
        viaticos_semanal = _num(emp["viaticos_semanales"])
        bono_semanal = _num(emp["bono_emp"])
        sueldo_diario = sueldo_semanal / 6.0
        viatico_diario = viaticos_semanal / 6.0
        bono_diario = bono_semanal / 6.0

        sueldo = round(sueldo_diario * dias_pagables, 2)
        viaticos = round(viatico_diario * dias_presente, 2)
        bono = round(bono_diario * dias_presente, 2)
        if tipo == "factor":
            he_importe = round(he_horas * (sueldo_diario / jornada) * valor, 2)
        elif tipo == "precio":
            he_importe = round(he_horas * valor, 2)
        else:
            he_importe = 0.0
        infonavit = _num(emp["infonavit_monto"])
        desc_retardos = 0.0
        if aplicar_retardos and retardos >= 3:
            desc_retardos = round((int(retardos) // 3) * sueldo_diario, 2)
        neto = round(sueldo + viaticos + bono + he_importe - infonavit - desc_nomina - desc_otra - desc_retardos, 2)

        nom_full = " ".join(x for x in [emp["primer_apellido"], emp["segundo_apellido"], emp["nombre"]] if x)
        detalles.append({
            "empleado_id": emp["id"], "cedula": cedula, "nombre": nom_full,
            "puesto": emp["puesto"], "clasificacion": emp["clasificacion"] or "Sin clasificar",
            "sueldo_contratado": round(sueldo_semanal, 2), "viaticos_contratado": round(viaticos_semanal, 2),
            "dias": dias_pagables, "faltas": faltas,
            "retardos": retardos, "vacaciones": dias_V, "sueldo": sueldo,
            "viaticos": viaticos, "bono": bono, "he_horas": he_horas, "he_importe": he_importe,
            "infonavit": infonavit, "desc_nomina": desc_nomina, "desc_otra": desc_otra,
            "desc_retardos": desc_retardos, "neto": neto, "baja_fecha": baja_fecha, "nota": obs,
        })
        if baja_fecha:
            avisos_baja.append({
                "empleado_id": emp["id"], "cedula": cedula, "nombre": nom_full,
                "curp": emp["curp"], "rfc": emp["rfc"], "nss": emp["nss"],
                "puesto": emp["puesto"], "baja_fecha": baja_fecha,
            })
        r += 1
    return detalles, avisos_baja, errores


@app.route("/nomina", methods=["GET", "POST"])
@login_required
def nomina():
    db = get_db()
    obras = obras_visibles(db)
    if request.method == "POST":
        obra_id = int(request.form.get("obra_id") or 0)
        finicio = request.form.get("fecha_inicio", "")
        aplicar_retardos = 1 if request.form.get("aplicar_retardos") else 0
        archivo = request.files.get("archivo")
        if not obra_id or not puede_ver_obra(db, obra_id):
            flash("Selecciona una obra valida.", "warning")
            return redirect(url_for("nomina"))
        if not archivo or not archivo.filename.lower().endswith(".xlsx"):
            flash("Sube el Excel de asistencia (.xlsx) que descargaste del sistema.", "danger")
            return redirect(url_for("nomina"))
        try:
            viernes = viernes_de(date.fromisoformat(finicio[:10]))
        except ValueError:
            viernes = viernes_de(date.today())
        jueves = viernes + timedelta(days=6)
        try:
            jornada = float(get_param(db, "jornada_horas", "8") or 8) or 8
        except (ValueError, TypeError):
            jornada = 8.0

        wb = openpyxl.load_workbook(archivo, data_only=True)
        ws = wb["Asistencia"] if "Asistencia" in wb.sheetnames else wb.active
        detalles, avisos_baja, errores = calcular_nomina(db, obra_id, ws, jornada, aplicar_retardos)
        if not detalles:
            for e in errores:
                flash(e, "warning")
            flash("No se pudo calcular: revisa que el archivo tenga trabajadores y el formato correcto.", "danger")
            return redirect(url_for("nomina"))

        total = round(sum(d["neto"] for d in detalles), 2)
        cur = db.execute(
            "INSERT INTO nominas(obra_id, fecha_inicio, fecha_fin, semana_num, anio, jornada, "
            "aplico_retardos, total_neto, creada_en, creada_por) VALUES(?,?,?,?,?,?,?,?,?,?)",
            (obra_id, viernes.isoformat(), jueves.isoformat(), viernes.isocalendar()[1],
             viernes.year, jornada, aplicar_retardos, total,
             datetime.now().isoformat(timespec="seconds"), session.get("nombre")))
        nomina_id = cur.lastrowid
        for d in detalles:
            db.execute(
                "INSERT INTO nomina_detalle(nomina_id, empleado_id, cedula, nombre, clasificacion, "
                "sueldo_contratado, viaticos_contratado, dias, faltas, "
                "retardos, vacaciones, sueldo, viaticos, bono, he_horas, he_importe, infonavit, "
                "desc_nomina, desc_otra, desc_retardos, neto, baja_fecha, nota) "
                "VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (nomina_id, d["empleado_id"], d["cedula"], d["nombre"], d["clasificacion"],
                 d["sueldo_contratado"], d["viaticos_contratado"], d["dias"], d["faltas"],
                 d["retardos"], d["vacaciones"], d["sueldo"], d["viaticos"], d["bono"], d["he_horas"],
                 d["he_importe"], d["infonavit"], d["desc_nomina"], d["desc_otra"],
                 d["desc_retardos"], d["neto"], d["baja_fecha"], d["nota"]))
        # marcar bajas reportadas
        for b in avisos_baja:
            db.execute("UPDATE empleados SET estatus='baja', fecha_baja=COALESCE(fecha_baja, ?) WHERE id=?",
                       (b["baja_fecha"], b["empleado_id"]))
            registrar_bitacora(db, "Baja de trabajador",
                               f"{b['nombre']} (cedula {b['cedula']}, NSS {b['nss']}) baja desde {b['baja_fecha']}")
            send_admin_alert(
                "POLTECH - Aviso de baja de trabajador",
                (f"El trabajador {b['nombre']} causa baja desde {b['baja_fecha']}.\n"
                 f"CURP: {b['curp']}  RFC: {b['rfc']}  NSS: {b['nss']}  Puesto: {b['puesto']}.\n\n"
                 f"Favor de tramitar la baja ante el IMSS."),
                extra=[get_param(db, "email_contador", "")])
        registrar_bitacora(db, "Calculo de nomina",
                           f"Obra {obra_id}, semana {viernes.isoformat()}, {len(detalles)} trabajadores, total {total:.2f}")
        db.commit()
        for e in errores:
            flash(e, "warning")
        return redirect(url_for("nomina_resultado", nomina_id=nomina_id))

    vdef = viernes_de(date.today()).isoformat()
    recientes = db.execute(
        "SELECT n.*, o.nombre AS obra FROM nominas n JOIN obras o ON o.id=n.obra_id "
        "ORDER BY n.creada_en DESC LIMIT 15").fetchall()
    if obras_del_usuario(db) is not None:
        vis = {o["id"] for o in obras}
        recientes = [n for n in recientes if n["obra_id"] in vis]
    return render_template("nomina.html", obras=obras, viernes_default=vdef, recientes=recientes)


def _con_contratado(db, det):
    """Convierte el detalle a dicts y, si el sueldo/viaticos contratado viene en 0
    (nominas calculadas antes de esta funcion), lo toma del catalogo actual del puesto."""
    filas = [dict(d) for d in det]
    faltan = [d for d in filas if not _num(d.get("sueldo_contratado"))
              and not _num(d.get("viaticos_contratado")) and d.get("empleado_id")]
    if faltan:
        ids = list({d["empleado_id"] for d in faltan})
        ph = ",".join("?" * len(ids))
        cat = {r["id"]: r for r in db.execute(
            f"SELECT e.id, p.sueldo_semanal, p.viaticos_semanales FROM empleados e "
            f"JOIN puestos p ON p.id=e.puesto_id WHERE e.id IN ({ph})", ids).fetchall()}
        for d in faltan:
            r = cat.get(d["empleado_id"])
            if r:
                d["sueldo_contratado"] = _num(r["sueldo_semanal"])
                d["viaticos_contratado"] = _num(r["viaticos_semanales"])
    return filas


@app.route("/nomina/<int:nomina_id>")
@login_required
def nomina_resultado(nomina_id):
    db = get_db()
    n = db.execute("SELECT n.*, o.nombre AS obra, o.estado FROM nominas n "
                   "JOIN obras o ON o.id=n.obra_id WHERE n.id=?", (nomina_id,)).fetchone()
    if not n:
        abort(404)
    if not puede_ver_obra(db, n["obra_id"]):
        abort(403)
    det = _con_contratado(db, db.execute(
        "SELECT * FROM nomina_detalle WHERE nomina_id=? ORDER BY nombre", (nomina_id,)).fetchall())
    tot = {k: sum(_num(d[k]) for d in det) for k in
           ("sueldo", "viaticos", "bono", "he_importe", "infonavit", "desc_nomina",
            "desc_otra", "desc_retardos", "neto")}
    bajas = [d for d in det if d["baja_fecha"]]
    # totales por clasificacion (cuanto se paga de cada grupo)
    por_clasif = {}
    for d in det:
        k = d["clasificacion"] or "Sin clasificar"
        g = por_clasif.setdefault(k, {"num": 0, "neto": 0.0})
        g["num"] += 1
        g["neto"] += _num(d["neto"])
    por_clasif = dict(sorted(por_clasif.items(), key=lambda kv: kv[1]["neto"], reverse=True))
    # Nomina real vs. lo que sale de caja
    caja_rows = []
    caja = {"devengado": 0.0, "infonavit": 0.0, "desc_nomina": 0.0, "desc_retardos": 0.0,
            "se_queda": 0.0, "traspaso": 0.0, "neto": 0.0, "sale": 0.0}
    for d in det:
        devengado = _num(d["sueldo"]) + _num(d["viaticos"]) + _num(d.get("bono")) + _num(d["he_importe"])
        se_queda = _num(d["infonavit"]) + _num(d["desc_nomina"]) + _num(d["desc_retardos"])
        traspaso = _num(d["desc_otra"])
        neto = _num(d["neto"])
        sale = round(neto + traspaso, 2)
        caja_rows.append({"nombre": d["nombre"], "cedula": d["cedula"],
                          "devengado": round(devengado, 2), "se_queda": round(se_queda, 2),
                          "traspaso": traspaso, "neto": neto, "sale": sale, "nota": d["nota"]})
        caja["devengado"] += devengado
        caja["infonavit"] += _num(d["infonavit"])
        caja["desc_nomina"] += _num(d["desc_nomina"])
        caja["desc_retardos"] += _num(d["desc_retardos"])
        caja["se_queda"] += se_queda
        caja["traspaso"] += traspaso
        caja["neto"] += neto
        caja["sale"] += sale
    caja = {k: round(v, 2) for k, v in caja.items()}
    pct_despacho = _num(get_param(db, "porcentaje_despacho", "0"))
    despacho = round(caja["neto"] * pct_despacho / 100.0, 2)
    total_a_pagar = round(caja["neto"] + despacho, 2)
    lote_otros_pagos = None
    if role_rank(session.get("role", "")) >= ADMIN_RANK:
        lote_otros_pagos = db.execute(
            "SELECT l.*, o.nombre AS obra FROM otros_pagos_lotes l LEFT JOIN obras o ON o.id=l.obra_id "
            "WHERE l.anio=? AND l.semana_num=?", (n["anio"], n["semana_num"])).fetchone()
    return render_template("nomina_resultado.html", n=n, det=det, tot=tot,
                           bajas=bajas, por_clasif=por_clasif, caja=caja, caja_rows=caja_rows,
                           pct_despacho=pct_despacho, despacho=despacho, total_a_pagar=total_a_pagar,
                           lote_otros_pagos=lote_otros_pagos)


@app.route("/nomina/<int:nomina_id>/eliminar", methods=["POST"])
@min_rank(GERENTE_RANK)
def nomina_eliminar(nomina_id):
    db = get_db()
    n = db.execute("SELECT * FROM nominas WHERE id=?", (nomina_id,)).fetchone()
    if not n:
        abort(404)
    if not puede_ver_obra(db, n["obra_id"]):
        abort(403)
    db.execute("DELETE FROM nomina_detalle WHERE nomina_id=?", (nomina_id,))
    db.execute("DELETE FROM nominas WHERE id=?", (nomina_id,))
    registrar_bitacora(db, "Eliminacion de nomina",
                       f"Nomina {nomina_id} (obra {n['obra_id']}, semana {n['fecha_inicio']})")
    db.commit()
    flash("Nomina eliminada.", "success")
    return redirect(url_for("nomina"))


@app.route("/nomina/<int:nomina_id>/autorizar", methods=["POST"])
@min_rank(GERENTE_RANK)
def nomina_autorizar(nomina_id):
    db = get_db()
    n = db.execute("SELECT * FROM nominas WHERE id=?", (nomina_id,)).fetchone()
    if not n:
        abort(404)
    if not puede_ver_obra(db, n["obra_id"]):
        abort(403)
    if (n["estatus"] or "pendiente") == "autorizada":
        flash("Esta nomina ya estaba autorizada.", "info")
        return redirect(url_for("nomina_resultado", nomina_id=nomina_id))
    db.execute("UPDATE nominas SET estatus='autorizada', autorizada_por=?, autorizada_en=? WHERE id=?",
               (session.get("nombre"), datetime.now().isoformat(timespec="seconds"), nomina_id))
    registrar_bitacora(db, "Autorizacion de nomina",
                       f"Nomina {nomina_id} (obra {n['obra_id']}, semana {n['fecha_inicio']}) autorizada y liberada")
    db.commit()
    # Avisar por correo a residentes, superintendentes y administradores de la obra
    obra = db.execute("SELECT nombre FROM obras WHERE id=?", (n["obra_id"],)).fetchone()
    correos = [r["username"] for r in db.execute(
        "SELECT DISTINCT u.username FROM users u "
        "LEFT JOIN user_obras uo ON uo.user_id=u.id "
        "WHERE u.role='admin' OR (u.role IN ('residente','superintendente') AND uo.obra_id=?)",
        (n["obra_id"],)).fetchall()]
    correos.append(os.environ.get("ADMIN_EMAIL", ""))
    enviar_correo(
        f"POLTECH - Nomina liberada: {obra['nombre'] if obra else ''} semana {n['semana_num']}",
        (f"La nomina de la obra {obra['nombre'] if obra else n['obra_id']} "
         f"(semana {n['semana_num']}/{n['anio']}, del {n['fecha_inicio']} al {n['fecha_fin']}) "
         f"fue autorizada y liberada por {session.get('nombre')}.\n"
         f"Total neto: {n['total_neto']:.2f}."),
        correos)
    flash("Nomina autorizada y liberada.", "success")
    return redirect(url_for("nomina_resultado", nomina_id=nomina_id))


@app.route("/nomina/<int:nomina_id>/reabrir", methods=["POST"])
@min_rank(ADMIN_RANK)
def nomina_reabrir(nomina_id):
    db = get_db()
    n = db.execute("SELECT * FROM nominas WHERE id=?", (nomina_id,)).fetchone()
    if not n:
        abort(404)
    db.execute("UPDATE nominas SET estatus='pendiente', autorizada_por=NULL, autorizada_en=NULL WHERE id=?",
               (nomina_id,))
    registrar_bitacora(db, "Reapertura de nomina",
                       f"Nomina {nomina_id} (obra {n['obra_id']}, semana {n['fecha_inicio']}) regresada a pendiente")
    db.commit()
    flash("Se quito la autorizacion; la nomina quedo pendiente de nuevo.", "success")
    return redirect(url_for("nomina_resultado", nomina_id=nomina_id))


@app.route("/nomina/<int:nomina_id>/excel")
@login_required
def nomina_excel(nomina_id):
    db = get_db()
    n = db.execute("SELECT n.*, o.nombre AS obra FROM nominas n JOIN obras o ON o.id=n.obra_id "
                   "WHERE n.id=?", (nomina_id,)).fetchone()
    if not n:
        abort(404)
    if not puede_ver_obra(db, n["obra_id"]):
        abort(403)
    det = _con_contratado(db, db.execute(
        "SELECT * FROM nomina_detalle WHERE nomina_id=? ORDER BY nombre", (nomina_id,)).fetchall())
    # datos bancarios por empleado (para el pago)
    cuentas = {}
    ids = [d["empleado_id"] for d in det if d["empleado_id"]]
    if ids:
        ph = ",".join("?" * len(ids))
        for c in db.execute(
                f"SELECT empleado_id, institucion, tipo_cuenta, numero FROM cuentas_bancarias "
                f"WHERE empleado_id IN ({ph})", ids).fetchall():
            cuentas.setdefault(c["empleado_id"], c)   # la primera cuenta del trabajador

    from openpyxl.styles import Font, PatternFill, Alignment
    NAVY = "16233C"
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Nomina"
    ws["A1"] = "POLTECH ACERO Y CONSTRUCCION, S.A. DE C.V."
    ws["A1"].font = Font(name="Arial", bold=True, size=13, color=NAVY)
    ws["A2"] = f"NOMINA SEMANAL  -  {n['obra']}  -  {n['fecha_inicio']} al {n['fecha_fin']}"
    ws["A2"].font = Font(name="Arial", size=10, italic=True, color="555555")
    if (n["estatus"] or "pendiente") == "autorizada":
        ws["A3"] = f"AUTORIZADA Y LIBERADA por {n['autorizada_por']} el {n['autorizada_en']}"
        ws["A3"].font = Font(name="Arial", size=9, bold=True, color="1E7E34")
    else:
        ws["A3"] = "PENDIENTE DE AUTORIZAR"
        ws["A3"].font = Font(name="Arial", size=9, bold=True, color="B8860B")
    cab = ["No.", "CEDULA", "NOMBRE", "DIAS", "FALTAS", "RET.", "VAC.", "SUELDO",
           "VIATICOS", "TOTAL HORAS EXTRAS", "IMPORTE HORAS EXTRAS", "INFONAVIT", "DESC. NOMINA",
           "DESC. OTRA", "DESC. RET.", "NETO A PAGAR", "BAJA",
           "TIPO CUENTA", "BANCO", "No. CUENTA", "CLASIFICACION",
           "SUELDO CONTRATADO", "VIATICOS CONTRATADO", "DEPOSITO A (obs)", "BONO"]
    HROW = 4
    for i, t in enumerate(cab, start=1):
        c = ws.cell(HROW, i, t)
        c.font = Font(name="Arial", bold=True, size=9, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor=NAVY)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    r = HROW + 1
    for idx, d in enumerate(det, start=1):
        cta = cuentas.get(d["empleado_id"])
        tipo_cta = cta["tipo_cuenta"] if cta else ""
        banco_cta = cta["institucion"] if cta else ""
        num_cta = str(cta["numero"]) if cta and cta["numero"] is not None else ""
        ws.append([idx, d["cedula"], d["nombre"], d["dias"], d["faltas"], d["retardos"],
                   d["vacaciones"], d["sueldo"], d["viaticos"], d["he_horas"], d["he_importe"],
                   d["infonavit"], d["desc_nomina"], d["desc_otra"], d["desc_retardos"],
                   d["neto"], d["baja_fecha"] or "", tipo_cta, banco_cta, num_cta,
                   d["clasificacion"] or "Sin clasificar",
                   _num(d["sueldo_contratado"]), _num(d["viaticos_contratado"]), d["nota"] or "",
                   _num(d["bono"])])
        ws.cell(r, 20).number_format = "@"   # No. de cuenta como texto
        r += 1
    # totales generales
    ws.append([])
    fila_tot = ["", "", "TOTALES", "", "", "", ""]
    for col in ("sueldo", "viaticos", "he_importe", "infonavit", "desc_nomina", "desc_otra", "desc_retardos", "neto"):
        fila_tot.append(round(sum(_num(d[col]) for d in det), 2))
    fila_tot.append("")
    ws.append(fila_tot)
    for col in list(range(8, 17)) + [22, 23, 25]:
        for row in range(HROW + 1, r + 2):
            cell = ws.cell(row, col)
            if isinstance(cell.value, (int, float)):
                cell.number_format = '#,##0.00'
    ws.column_dimensions["C"].width = 30
    for i in [2] + list(range(4, 26)):
        Lc = openpyxl.utils.get_column_letter(i)
        ws.column_dimensions[Lc].width = max(ws.column_dimensions[Lc].width or 8, 12)

    # ---- Resumen por clasificacion ----
    por = {}
    for d in det:
        k = d["clasificacion"] or "Sin clasificar"
        g = por.setdefault(k, {"num": 0, "neto": 0.0})
        g["num"] += 1; g["neto"] += _num(d["neto"])
    rs = wb.create_sheet("Por clasificacion")
    rs.append(["CLASIFICACION", "TRABAJADORES", "NETO PAGADO"])
    for c in rs[1]:
        c.font = Font(name="Arial", bold=True, color="FFFFFF"); c.fill = PatternFill("solid", fgColor=NAVY)
    for k, g in sorted(por.items(), key=lambda kv: kv[1]["neto"], reverse=True):
        rs.append([k, g["num"], round(g["neto"], 2)])
    rs.append(["TOTAL", sum(g["num"] for g in por.values()), round(sum(g["neto"] for g in por.values()), 2)])
    for row in range(2, rs.max_row + 1):
        rs.cell(row, 3).number_format = '#,##0.00'
    rs.column_dimensions["A"].width = 24; rs.column_dimensions["B"].width = 14; rs.column_dimensions["C"].width = 16

    # ---- Hoja: Flujo de caja (nomina real vs. lo que sale) ----
    cj = wb.create_sheet("Flujo de caja")
    cab_cj = ["NOMBRE", "NOMINA REAL", "SE QUEDA (INFONAVIT+DESC+RET)",
              "TRASPASO A 3RO", "NETO AL TRABAJADOR", "SALE DE CAJA", "DEPOSITO A"]
    cj.append(cab_cj)
    for c in cj[1]:
        c.font = Font(name="Arial", bold=True, color="FFFFFF"); c.fill = PatternFill("solid", fgColor=NAVY)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    tot_cj = {"dev": 0.0, "sq": 0.0, "tr": 0.0, "net": 0.0, "sal": 0.0}
    for d in det:
        dev = _num(d["sueldo"]) + _num(d["viaticos"]) + _num(d["bono"]) + _num(d["he_importe"])
        sq = _num(d["infonavit"]) + _num(d["desc_nomina"]) + _num(d["desc_retardos"])
        tr = _num(d["desc_otra"]); net = _num(d["neto"]); sal = round(net + tr, 2)
        cj.append([d["nombre"], round(dev, 2), round(sq, 2), tr, net, sal, d["nota"] or ""])
        tot_cj["dev"] += dev; tot_cj["sq"] += sq; tot_cj["tr"] += tr
        tot_cj["net"] += net; tot_cj["sal"] += sal
    cj.append(["TOTALES", round(tot_cj["dev"], 2), round(tot_cj["sq"], 2), round(tot_cj["tr"], 2),
               round(tot_cj["net"], 2), round(tot_cj["sal"], 2), ""])
    fila_tot = cj.max_row
    for c in cj[fila_tot]:
        c.font = Font(name="Arial", bold=True)
    # Despacho (% sobre el neto) y Total a pagar
    pct = _num(get_param(db, "porcentaje_despacho", "0"))
    despacho = round(tot_cj["net"] * pct / 100.0, 2)
    cj.append([f"DESPACHO ({pct:g}%)", "", "", "", despacho, "", ""])
    cj.append(["TOTAL A PAGAR", "", "", "", round(tot_cj["net"] + despacho, 2), "", ""])
    for rr in (cj.max_row - 1, cj.max_row):
        cj.cell(rr, 1).font = Font(name="Arial", bold=True)
        cj.cell(rr, 5).font = Font(name="Arial", bold=True)
    for row in range(2, cj.max_row + 1):
        for col in range(2, 7):
            cj.cell(row, col).number_format = '#,##0.00'
    cj.column_dimensions["A"].width = 28
    for L2 in ("B", "C", "D", "E", "F"):
        cj.column_dimensions[L2].width = 16
    cj.column_dimensions["G"].width = 22

    # ---- Hoja "Otros pagos semanales" (solo si quien descarga es Administrador
    # Y ya se genero el lote de esta semana para ESTA nomina; nunca visible para
    # superintendente/residente aunque el lote exista). ----
    if role_rank(session.get("role", "")) >= ADMIN_RANK:
        lote = db.execute(
            "SELECT * FROM otros_pagos_lotes WHERE nomina_id=?", (nomina_id,)).fetchone()
        if lote:
            items = json.loads(lote["detalle_json"] or "[]")
            op = wb.create_sheet("Otros pagos semanales")
            op.append(["CLAVE", "NOMBRE", "BANCO", "TIPO CUENTA", "No. CUENTA", "MONTO"])
            for c in op[1]:
                c.font = Font(name="Arial", bold=True, color="FFFFFF")
                c.fill = PatternFill("solid", fgColor=NAVY)
            tot_op = 0.0
            for it in items:
                op.append([it.get("clave", ""), it.get("nombre", ""), it.get("banco", ""),
                          it.get("tipo_cuenta", ""), it.get("numero_cuenta", ""),
                          _num(it.get("monto"))])
                tot_op += _num(it.get("monto"))
            op.append(["", "", "", "", "TOTAL", round(tot_op, 2)])
            for row in range(2, op.max_row + 1):
                op.cell(row, 6).number_format = '#,##0.00'
                op.cell(row, 5).number_format = "@"
            op.column_dimensions["B"].width = 26
            for L3 in ("A", "C", "D", "E", "F"):
                op.column_dimensions[L3].width = 16

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    nombre = f"Nomina_{re.sub(r'[^A-Za-z0-9]+','_',n['obra']).strip('_')}_{n['fecha_inicio']}.xlsx"
    return send_file(buf, as_attachment=True, download_name=nombre,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ---------------------------------------------------------------------------
# Otros pagos semanales (solo Administrador) - pagos que no son nomina de
# trabajadores (no llevan puesto, obra, asistencia ni contrato). Se generan
# como una hoja aparte dentro del Excel de una nomina calculada, una sola vez
# por semana, para que el contador los pueda pagar junto con la nomina real
# sin que se mezclen con los datos de los trabajadores.
# ---------------------------------------------------------------------------
@app.route("/otros-pagos", methods=["GET", "POST"])
@min_rank(ADMIN_RANK)
def otros_pagos():
    db = get_db()
    if request.method == "POST":
        nombre = request.form.get("nombre", "").strip()
        if not nombre:
            flash("Escribe el nombre.", "warning")
        else:
            db.execute(
                "INSERT INTO otros_pagos(clave, nombre, banco, tipo_cuenta, numero_cuenta, "
                "monto_semanal, creado_en) VALUES(?,?,?,?,?,?,?)",
                (siguiente_clave_otro_pago(db), titulo(nombre),
                 request.form.get("banco", "").strip(),
                 request.form.get("tipo_cuenta", "").strip(),
                 solo_digitos(request.form.get("numero_cuenta", "")),
                 float(request.form.get("monto_semanal") or 0),
                 date.today().isoformat()))
            db.commit()
            flash("Agregado.", "success")
        return redirect(url_for("otros_pagos"))
    roster = db.execute("SELECT * FROM otros_pagos ORDER BY activo DESC, clave").fetchall()
    lotes = [dict(r, num=len(json.loads(r["detalle_json"] or "[]"))) for r in db.execute(
        "SELECT l.*, o.nombre AS obra FROM otros_pagos_lotes l "
        "LEFT JOIN obras o ON o.id=l.obra_id ORDER BY l.anio DESC, l.semana_num DESC").fetchall()]
    return render_template("otros_pagos.html", roster=roster, lotes=lotes,
                           bancos=BANCOS, tipos=TIPOS_CUENTA)


@app.route("/otro-pago/<int:otro_id>/desactivar", methods=["POST"])
@min_rank(ADMIN_RANK)
def otro_pago_desactivar(otro_id):
    db = get_db()
    db.execute("UPDATE otros_pagos SET activo=0 WHERE id=?", (otro_id,))
    db.commit()
    flash("Desactivado. Ya no se incluira en los proximos lotes.", "success")
    return redirect(url_for("otros_pagos"))


@app.route("/otro-pago/<int:otro_id>/reactivar", methods=["POST"])
@min_rank(ADMIN_RANK)
def otro_pago_reactivar(otro_id):
    db = get_db()
    db.execute("UPDATE otros_pagos SET activo=1 WHERE id=?", (otro_id,))
    db.commit()
    flash("Reactivado.", "success")
    return redirect(url_for("otros_pagos"))


@app.route("/otro-pago/<int:otro_id>/editar", methods=["GET", "POST"])
@min_rank(ADMIN_RANK)
def otro_pago_editar(otro_id):
    db = get_db()
    r = db.execute("SELECT * FROM otros_pagos WHERE id=?", (otro_id,)).fetchone()
    if not r:
        abort(404)
    if request.method == "POST":
        nombre = request.form.get("nombre", "").strip()
        if not nombre:
            flash("Escribe el nombre.", "warning")
            return render_template("otro_pago_form.html", r=r, bancos=BANCOS, tipos=TIPOS_CUENTA)
        db.execute(
            "UPDATE otros_pagos SET nombre=?, banco=?, tipo_cuenta=?, numero_cuenta=?, "
            "monto_semanal=? WHERE id=?",
            (titulo(nombre), request.form.get("banco", "").strip(),
             request.form.get("tipo_cuenta", "").strip(),
             solo_digitos(request.form.get("numero_cuenta", "")),
             float(request.form.get("monto_semanal") or 0), otro_id))
        db.commit()
        flash(f"'{titulo(nombre)}' actualizado.", "success")
        return redirect(url_for("otros_pagos"))
    return render_template("otro_pago_form.html", r=r, bancos=BANCOS, tipos=TIPOS_CUENTA)


@app.route("/nomina/<int:nomina_id>/otros-pagos/generar", methods=["POST"])
@min_rank(ADMIN_RANK)
def otros_pagos_generar(nomina_id):
    db = get_db()
    n = db.execute("SELECT * FROM nominas WHERE id=?", (nomina_id,)).fetchone()
    if not n:
        abort(404)
    ya = db.execute(
        "SELECT * FROM otros_pagos_lotes WHERE anio=? AND semana_num=?",
        (n["anio"], n["semana_num"])).fetchone()
    if ya:
        flash("Ya se generaron los Otros pagos de esta semana (no se puede repetir).", "warning")
        return redirect(url_for("nomina_resultado", nomina_id=nomina_id))
    activos = db.execute("SELECT * FROM otros_pagos WHERE activo=1 ORDER BY clave").fetchall()
    if not activos:
        flash("No hay personas activas en Otros pagos semanales para generar.", "warning")
        return redirect(url_for("nomina_resultado", nomina_id=nomina_id))
    items = [{"clave": r["clave"], "nombre": r["nombre"], "banco": r["banco"],
              "tipo_cuenta": r["tipo_cuenta"], "numero_cuenta": r["numero_cuenta"],
              "monto": r["monto_semanal"]} for r in activos]
    db.execute(
        "INSERT INTO otros_pagos_lotes(anio, semana_num, nomina_id, obra_id, generado_por, "
        "generado_en, detalle_json) VALUES(?,?,?,?,?,?,?)",
        (n["anio"], n["semana_num"], nomina_id, n["obra_id"], session.get("nombre", ""),
         datetime.now().isoformat(timespec="seconds"), json.dumps(items)))
    registrar_bitacora(db, "Otros pagos semanales generados",
                       f"Semana {n['semana_num']}/{n['anio']}, nomina {nomina_id}, "
                       f"{len(items)} registro(s)")
    db.commit()
    flash(f"Listo: se agrego la hoja 'Otros pagos semanales' con {len(items)} registro(s) "
          "al Excel de esta nomina.", "success")
    return redirect(url_for("nomina_resultado", nomina_id=nomina_id))


# ---------------------------------------------------------------------------
# Cambiar contrasena
# ---------------------------------------------------------------------------
@app.route("/cuenta/password", methods=["GET", "POST"])
@login_required
def cambiar_password():
    db = get_db()
    if request.method == "POST":
        actual = request.form.get("actual", "")
        nueva = request.form.get("nueva", "")
        row = db.execute("SELECT * FROM users WHERE id=?", (session["user_id"],)).fetchone()
        if not check_password_hash(row["password_hash"], actual):
            flash("La contrasena actual no es correcta.", "danger")
        elif len(nueva) < 6:
            flash("La nueva contrasena debe tener al menos 6 caracteres.", "danger")
        else:
            db.execute("UPDATE users SET password_hash=? WHERE id=?",
                       (generate_password_hash(nueva), session["user_id"]))
            db.commit()
            flash("Contrasena actualizada.", "success")
            return redirect(url_for("dashboard"))
    return render_template("cambiar_password.html")

# ---------------------------------------------------------------------------
# Usuarios (solo el administrador puede crear cuentas y asignar roles)
# ---------------------------------------------------------------------------
ADMIN_RANK = 100  # solo administrador gestiona usuarios

@app.route("/usuarios", methods=["GET", "POST"])
@min_rank(ADMIN_RANK)
def usuarios():
    db = get_db()
    if request.method == "POST":
        u = request.form.get("username", "").strip().lower()
        nombre = request.form.get("nombre", "").strip()
        pw = request.form.get("password", "")
        role = request.form.get("role", "")

        errores = []
        if not u: errores.append("El usuario (correo) es obligatorio.")
        if not nombre: errores.append("El nombre es obligatorio.")
        if len(pw) < 6: errores.append("La contrasena debe tener al menos 6 caracteres.")
        if role not in ROLES: errores.append("Selecciona un rol valido.")
        if db.execute("SELECT 1 FROM users WHERE username=?", (u,)).fetchone():
            errores.append("Ese usuario ya existe.")

        if errores:
            for e in errores:
                flash(e, "danger")
        else:
            cur = db.execute(
                "INSERT INTO users(username, password_hash, nombre, role) VALUES(?,?,?,?)",
                (u, generate_password_hash(pw), nombre, role),
            )
            nuevo_id = cur.lastrowid
            for oid in request.form.getlist("obras"):
                try:
                    db.execute("INSERT INTO user_obras(user_id, obra_id) VALUES(?,?)", (nuevo_id, oid))
                except sqlite3.IntegrityError:
                    pass
            db.commit()
            flash("Usuario creado: " + u, "success")
        return redirect(url_for("usuarios"))

    filas = db.execute("SELECT * FROM users ORDER BY nombre").fetchall()
    obras_l = db.execute("SELECT * FROM obras ORDER BY nombre").fetchall()
    asignadas = {}
    for r in db.execute(
        "SELECT uo.user_id, o.id, o.nombre FROM user_obras uo "
        "JOIN obras o ON o.id=uo.obra_id ORDER BY o.nombre").fetchall():
        asignadas.setdefault(r["user_id"], []).append({"id": r["id"], "nombre": r["nombre"]})
    return render_template("usuarios_list.html", usuarios=filas, obras=obras_l,
                           roles=ROLES, asignadas=asignadas)


@app.route("/usuarios/<int:uid>/obras", methods=["POST"])
@min_rank(ADMIN_RANK)
def usuario_obras(uid):
    db = get_db()
    db.execute("DELETE FROM user_obras WHERE user_id=?", (uid,))
    for oid in request.form.getlist("obras"):
        try:
            db.execute("INSERT INTO user_obras(user_id, obra_id) VALUES(?,?)", (uid, oid))
        except sqlite3.IntegrityError:
            pass
    db.commit()
    flash("Obras del usuario actualizadas.", "success")
    return redirect(url_for("usuarios"))

@app.route("/usuarios/<int:uid>/rol", methods=["POST"])
@min_rank(ADMIN_RANK)
def usuario_rol(uid):
    db = get_db()
    role = request.form.get("role", "")
    if uid == session["user_id"]:
        flash("No puedes cambiar tu propio rol.", "warning")
    elif role in ROLES:
        db.execute("UPDATE users SET role=? WHERE id=?", (role, uid))
        db.commit()
        flash("Rol actualizado.", "success")
    return redirect(url_for("usuarios"))

@app.route("/usuarios/<int:uid>/username", methods=["POST"])
@min_rank(ADMIN_RANK)
def usuario_username(uid):
    db = get_db()
    nuevo = request.form.get("username", "").strip().lower()
    if not nuevo:
        flash("El correo no puede quedar vacio.", "danger")
    elif db.execute("SELECT 1 FROM users WHERE username=? AND id<>?", (nuevo, uid)).fetchone():
        flash(f"Ya existe otro usuario con el correo '{nuevo}'.", "danger")
    else:
        db.execute("UPDATE users SET username=? WHERE id=?", (nuevo, uid))
        db.commit()
        registrar_bitacora(db, "Correo de usuario actualizado", f"Usuario {uid} -> {nuevo}")
        db.commit()
        if uid == session.get("user_id"):
            session["username"] = nuevo
        flash("Correo (usuario) actualizado.", "success")
    return redirect(url_for("usuarios"))

@app.route("/usuarios/<int:uid>/nombre", methods=["POST"])
@min_rank(ADMIN_RANK)
def usuario_nombre(uid):
    db = get_db()
    nombre = request.form.get("nombre", "").strip()
    if not nombre:
        flash("El nombre no puede quedar vacio.", "danger")
    else:
        db.execute("UPDATE users SET nombre=? WHERE id=?", (nombre, uid))
        db.commit()
        if uid == session.get("user_id"):
            session["nombre"] = nombre
        flash("Nombre actualizado.", "success")
    return redirect(url_for("usuarios"))

@app.route("/usuarios/<int:uid>/password", methods=["POST"])
@min_rank(ADMIN_RANK)
def usuario_password(uid):
    db = get_db()
    nueva = request.form.get("nueva", "")
    if len(nueva) < 6:
        flash("La contrasena debe tener al menos 6 caracteres.", "danger")
    else:
        db.execute("UPDATE users SET password_hash=? WHERE id=?",
                   (generate_password_hash(nueva), uid))
        db.commit()
        flash("Contrasena restablecida.", "success")
    return redirect(url_for("usuarios"))

@app.route("/usuarios/<int:uid>/eliminar", methods=["POST"])
@min_rank(ADMIN_RANK)
def usuario_eliminar(uid):
    db = get_db()
    if uid == session["user_id"]:
        flash("No puedes eliminar tu propia cuenta.", "warning")
    else:
        db.execute("DELETE FROM users WHERE id=?", (uid,))
        db.commit()
        flash("Usuario eliminado.", "success")
    return redirect(url_for("usuarios"))

@app.route("/parametros", methods=["GET", "POST"])
@min_rank(ADMIN_RANK)
def parametros():
    db = get_db()
    if request.method == "POST":
        for clave in ("sm_general", "sm_zlfn", "propuesta_extra", "jornada_horas", "porcentaje_despacho"):
            if clave not in request.form:
                continue
            val = (request.form.get(clave, "") or "").replace(",", ".").strip()
            try:
                float(val)
                set_param(db, clave, val)
            except ValueError:
                flash(f"Valor invalido para {clave}.", "danger")
        if request.form.get("sm_vigencia", "").strip():
            set_param(db, "sm_vigencia", request.form.get("sm_vigencia", "").strip())
        if "email_contador" in request.form:
            set_param(db, "email_contador", request.form.get("email_contador", "").strip())
        # datos de la empresa (para el contrato) y API market: solo si vienen en el envio
        for clave in ("empresa_razon_social", "empresa_rfc", "empresa_representante",
                      "empresa_domicilio", "empresa_instrumento", "empresa_volumen",
                      "empresa_fecha_escritura", "empresa_notario_num", "empresa_notario_nombre",
                      "empresa_notario_ciudad", "apimarket_token", "apimarket_url_nss",
                      "apimarket_url_vigencia"):
            if clave in request.form:
                set_param(db, clave, request.form.get(clave, "").strip())
        db.commit()
        flash("Parametros actualizados.", "success")
        return redirect(url_for("parametros"))
    claves = ("sm_general", "sm_zlfn", "sm_vigencia", "propuesta_extra", "jornada_horas",
              "porcentaje_despacho", "email_contador", "empresa_razon_social", "empresa_rfc",
              "empresa_representante", "empresa_domicilio", "empresa_instrumento", "empresa_volumen",
              "empresa_fecha_escritura", "empresa_notario_num", "empresa_notario_nombre",
              "empresa_notario_ciudad", "apimarket_token", "apimarket_url_nss",
              "apimarket_url_vigencia")
    datos = {k: get_param(db, k) for k in claves}
    return render_template("parametros.html", p=datos,
                           propuesta=propuesta_salario_alta(db))

@app.errorhandler(403)
def forbidden(e):
    return render_template("403.html"), 403

if __name__ == "__main__":
    app.run(debug=True, port=5000)
